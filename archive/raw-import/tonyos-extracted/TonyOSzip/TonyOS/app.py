import os
import json
from pathlib import Path
from datetime import datetime
import logging
from contextlib import contextmanager

from flask import Flask, render_template, request, jsonify, session, Response, send_from_directory
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy.orm import DeclarativeBase
from dotenv import load_dotenv
from openai import OpenAI
from anthropic import Anthropic
import psycopg2
import requests
from bs4 import BeautifulSoup
import re
from psycopg2.extras import RealDictCursor
from werkzeug.middleware.proxy_fix import ProxyFix
from werkzeug.security import generate_password_hash, check_password_hash
from policy import enforce_policy
from promotion import should_create_candidate, candidate_id
from security import require_api_key, rate_limit, add_cors_headers, handle_cors_preflight
import time

logging.basicConfig(level=logging.DEBUG)

BASE_DIR = Path(__file__).resolve().parent
LONG_MEMORY_FILE = BASE_DIR / "long_memory.txt"
MEMORY_LOG_FILE = BASE_DIR / "memory_log.jsonl"
CHAT_HISTORY_FILE = BASE_DIR / "chat_history.json"

load_dotenv()

openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
client = openai_client

# Claude client for deep analysis (using Replit AI Integrations)
anthropic_client = Anthropic(
    api_key=os.environ.get("AI_INTEGRATIONS_ANTHROPIC_API_KEY"),
    base_url=os.environ.get("AI_INTEGRATIONS_ANTHROPIC_BASE_URL")
)

AVAILABLE_MODELS = {
    "openai": {
        "gpt-4o": "GPT-4o (Best)",
        "gpt-4o-mini": "GPT-4o Mini (Fast)",
        "gpt-4-turbo": "GPT-4 Turbo"
    }
}


class Base(DeclarativeBase):
    pass


app = Flask(__name__)
app.secret_key = os.environ.get("SESSION_SECRET")
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

from datetime import timedelta
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=30)
app.config['SESSION_COOKIE_SECURE'] = True
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

DATABASE_URL = os.getenv("DATABASE_URL")
SMART_MEMORY_ENABLED = bool(DATABASE_URL)

app.config["SQLALCHEMY_DATABASE_URI"] = DATABASE_URL
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["SQLALCHEMY_ENGINE_OPTIONS"] = {
    'pool_pre_ping': True,
    "pool_recycle": 300,
}

db = SQLAlchemy(app, model_class=Base)

with app.app_context():
    import models
    db.create_all()
    logging.info("Database tables created")


@app.after_request
def after_request(response):
    return add_cors_headers(response)


TONYOS_USERNAME = "TonyBeal40"
TONYOS_PASSWORD_HASH = generate_password_hash("BreyerJaceBeal6831!")

PUBLIC_ROUTES = ['/login', '/logout', '/static/', '/api/health', '/api/storybrand-scan', '/favicon']


@app.before_request
def before_request():
    preflight = handle_cors_preflight()
    if preflight:
        return preflight
    
    if any(request.path.startswith(route) for route in PUBLIC_ROUTES):
        return None
    
    if not session.get('authenticated'):
        if request.path.startswith('/api/'):
            return jsonify({"error": "Authentication required"}), 401
        return render_template('login.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username', '')
        password = request.form.get('password', '')
        
        if username == TONYOS_USERNAME and check_password_hash(TONYOS_PASSWORD_HASH, password):
            session['authenticated'] = True
            session['username'] = username
            session.permanent = True
            return jsonify({"success": True})
        return jsonify({"error": "Invalid credentials"}), 401
    
    return render_template('login.html')


@app.route('/logout')
def logout():
    session.clear()
    return render_template('login.html', message="Logged out successfully")


@contextmanager
def get_db_context():
    """Get database connection with context manager for safe cleanup."""
    if not DATABASE_URL:
        yield None
        return
    conn = None
    try:
        conn = psycopg2.connect(DATABASE_URL, cursor_factory=RealDictCursor)
        yield conn
    except Exception as e:
        logging.error(f"[db] Connection error: {e}")
        yield None
    finally:
        if conn:
            try:
                conn.close()
            except Exception:
                pass


def get_db():
    """Get database connection (legacy pattern - use get_db_context() for auto-cleanup)."""
    if not DATABASE_URL:
        return None
    try:
        return psycopg2.connect(DATABASE_URL, cursor_factory=RealDictCursor)
    except Exception as e:
        logging.error(f"[db] Connection error: {e}")
        return None


def embed_text(text):
    """Get embedding for text using OpenAI."""
    if not text or not isinstance(text, str):
        return None
    try:
        response = client.embeddings.create(
            model="text-embedding-3-small",
            input=text[:8000]
        )
        return response.data[0].embedding
    except Exception as e:
        logging.error(f"[embed] Error: {e}")
        return None


def retrieve_smart_memories(user_id, query_text, limit=5, min_similarity=0.75):
    """Retrieve relevant memories using vector similarity."""
    if not SMART_MEMORY_ENABLED:
        return []
    try:
        embedding = embed_text(query_text)
        if not embedding:
            return []
        with get_db_context() as conn:
            if not conn:
                return []
            cur = conn.cursor()
            cur.execute(
                "SELECT * FROM match_memories(%s, %s::vector, %s)",
                (user_id, embedding, limit)
            )
            rows = cur.fetchall()
            return [dict(r) for r in rows if r.get('similarity', 0) >= min_similarity]
    except Exception as e:
        logging.error(f"[retrieve_memories] Error: {e}")
        return []


def save_smart_memory(user_id, memory_text, importance=3):
    """Save memory with embedding to database."""
    if not SMART_MEMORY_ENABLED:
        return False
    try:
        embedding = embed_text(memory_text)
        if not embedding:
            return False
        with get_db_context() as conn:
            if not conn:
                return False
            cur = conn.cursor()
            cur.execute(
                "INSERT INTO memories (user_id, memory, importance, embedding) VALUES (%s, %s, %s, %s::vector)",
                (user_id, memory_text, importance, embedding)
            )
            conn.commit()
            logging.info(f"[memory] Saved: {memory_text[:50]}...")
            return True
    except Exception as e:
        logging.error(f"[save_memory] Error: {e}")
        return False


def load_chat_history():
    """Load chat history from file."""
    if CHAT_HISTORY_FILE.exists():
        try:
            return json.loads(CHAT_HISTORY_FILE.read_text(encoding="utf-8"))
        except:
            return []
    return []


def save_chat_history(history):
    """Save chat history to file."""
    CHAT_HISTORY_FILE.write_text(json.dumps(history, indent=2), encoding="utf-8")


chat_history = load_chat_history()

STYLE_PROFILE = """
You are TonyOS, a private AI console that lives on Tony Beal's computer.

Your job:
- Act like a strategist sitting on Tony's desk.
- Help him think, plan, write, sell, and build systems.
- Stay practical, direct, and encouraging.
- No fluff, no fake hype, no emojis.

Tone:
- Friendly but blunt when needed.
- Respectful and calm when Tony is stressed.
- Explain thinking in clear, simple language.
- Use short paragraphs and bullet points when it helps.

Behavior rules:
1. Treat this as a long term partnership with Tony.
2. You can reference what he has done in past roles.
3. Always give next action steps, not just theory.
4. If something is unclear, make a best guess and move forward.

Formatting rules:
- Use line breaks so answers are easy to read.
- Use numbered lists and bullets when giving steps or plans.
- No emojis.

Memory rules:
- You have access to a local memory file that stores key notes from past chats.
- Use that memory as context but do not dump it back verbatim.
"""

ENGINEER_PROFILE = """
You are CodeBench AI, a private execution assistant for a single user.

Mission: Help the user build, debug, and operate practical tools (dashboards, tracking systems, automations) with clean code and clear steps.

Operating rules:
1) Truth first
- Never invent facts, file contents, errors, links, results, or numbers
- If you do not know, say you do not know and ask for the exact missing input
- If a claim depends on user data, request the screenshot, snippet, or file

2) No fake actions
- Never claim you sent an email, posted on LinkedIn, changed a real system, accessed private accounts, or ran commands you cannot run
- When you provide code, treat it as a draft until the user confirms it ran successfully

3) Output style
- No emojis
- No hype
- No long speeches
- Use this format unless the user asks otherwise:
  A) Diagnosis or plan in 1 to 3 bullets
  B) Exact steps to do next (numbered)
  C) Paste-ready code only (when needed)

4) Debugging behavior
- Act like a senior engineer
- Ask for the minimum needed to resolve the issue:
  - the exact error text
  - the file name
  - the lines around the error
  - what the user expected vs what happened
- When giving fixes, change as little as possible
- Provide a single corrected block the user can paste

5) Security and privacy
- Assume the workspace may include sensitive company data
- Recommend anonymizing screenshots and demo data
- Do not encourage scraping, blasting, or bypassing access controls

6) Cost discipline
- Prefer simple solutions that minimize compute and API calls
- Default to local-first and event-driven approaches
- Avoid background polling unless requested

7) Default assumptions
- The user is building locally and may also deploy to Replit or Netlify
- Google Sheets is the system of record for dashboards unless a database is explicitly requested
- If a task can be done in Apps Script instead of a server, prefer Apps Script

When the user says "make it work," respond with:
- the exact next step they should do right now
- the exact code to paste (only if required)
- what success looks like
"""

PERSONAS = {
    "default": STYLE_PROFILE,
    "engineer": ENGINEER_PROFILE
}

SYSTEM_PROMPT_CODE = """
You are TonyOS in CODE mode.

Rules:
- You are a senior engineer + debugger.
- Output corrected code only when asked for "final code".
- When debugging: identify the root cause, then give the minimum patch.
- Never invent file contents. If missing, ask for the exact file or snippet.
- Prefer small changes over rewrites.
- If a command is needed, provide the exact command.
- Keep formatting clean and copy-paste ready.
"""

DEFAULT_MODEL_MAP = {
    "truth": ("openai", "gpt-4o"),
    "deep":  ("openai", "gpt-4o"),
    "fast":  ("openai", "gpt-4o-mini"),
    "code":  ("openai", "gpt-4o")
}


def choose_lane(user_msg: str, mode: str) -> str:
    """Route to appropriate mode based on explicit mode or auto-detection."""
    s = (user_msg or "").lower().strip()
    
    if mode in ("truth", "deep", "fast", "code"):
        return mode
    
    if any(k in s for k in ["error", "stack trace", "bug", "debug", "fix", "crash", "compile", "syntax"]):
        return "code"
    
    return "truth"


def ensure_memory_files():
    if not LONG_MEMORY_FILE.exists():
        LONG_MEMORY_FILE.write_text("", encoding="utf-8")
    if not MEMORY_LOG_FILE.exists():
        MEMORY_LOG_FILE.write_text("", encoding="utf-8")


def load_long_memory(max_chars: int = 3000) -> str:
    ensure_memory_files()
    text = LONG_MEMORY_FILE.read_text(encoding="utf-8")
    if len(text) > max_chars:
        return text[-max_chars:]
    return text


def append_to_long_memory(summary: str) -> None:
    ensure_memory_files()
    with LONG_MEMORY_FILE.open("a", encoding="utf-8") as f:
        f.write("\n" + summary.strip() + "\n")


def log_interaction(user_msg: str, assistant_msg: str) -> None:
    ensure_memory_files()
    record = {
        "timestamp": datetime.utcnow().isoformat() + "Z",
        "user": user_msg,
        "assistant": assistant_msg,
    }
    with MEMORY_LOG_FILE.open("a", encoding="utf-8") as f:
        f.write(json.dumps(record) + "\n")


def with_mood(prompt: str, mood: str) -> str:
    """Apply mood tone to prompt."""
    mood_lines = {
        "happy": "Tone: upbeat, confident, encouraging",
        "sad": "Tone: calm, empathetic, slightly heavy",
        "mad": "Tone: blunt, impatient with excuses, still helpful",
        "neutral": "Tone: neutral, crisp"
    }
    mood_line = mood_lines.get(mood, mood_lines["neutral"])
    return f"{mood_line}\n\n{prompt}"


def truthify_prompt(prompt: str) -> str:
    """Apply truth policy to prompt."""
    rules = " ".join([
        "If you cannot verify something, say so explicitly.",
        "Cite sources when making factual claims.",
        "Distinguish between opinion and fact.",
        "When uncertain, admit uncertainty."
    ])
    return f"[TRUTH POLICY: {rules}]\n\n{prompt}"


def extract_memory_if_worth_saving(user_message: str) -> dict:
    """Extract stable, useful facts from user messages to save to memory."""
    system_prompt = """You are a memory extraction engine for a personal assistant.
Only save stable, useful facts that will matter later.

Save examples:
- preferences (tone, tools, writing style)
- long-term goals (income targets, projects)
- ongoing projects (TonyOS, Natoli reachout)
- constraints (no emojis, no em dashes)
- recurring context (location/timezone)

Do NOT save:
- temporary emotions
- one-time tasks
- random details
- private secrets like API keys
- medical/legal/financial sensitive details unless explicitly requested

Return JSON only:
{
  "should_save": true/false,
  "memory": "one clean sentence",
  "importance": 1-5
}"""

    try:
        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message}
            ],
            temperature=0.2
        )
        txt = completion.choices[0].message.content.strip()
        
        start = txt.find("{")
        end = txt.rfind("}")
        if start == -1 or end == -1:
            return {"should_save": False}
        
        obj = json.loads(txt[start:end + 1])
        if not obj or obj.get("should_save") != True:
            return {"should_save": False}
        if not obj.get("memory") or not isinstance(obj.get("memory"), str):
            return {"should_save": False}
        if "api key" in obj.get("memory", "").lower():
            return {"should_save": False}
        
        importance = max(1, min(5, int(obj.get("importance", 3))))
        return {
            "should_save": True,
            "memory": obj["memory"][:220],
            "importance": importance
        }
    except Exception:
        return {"should_save": False}


def call_tonyos_model(user_message: str, mode: str = "normal", mood: str = "neutral", truth: bool = False, persona: str = "default") -> str:
    lane = choose_lane(user_message, mode)
    _, model_name = DEFAULT_MODEL_MAP.get(lane, ("openai", "gpt-4o"))
    
    print(f"[lane] Mode: {mode} -> Lane: {lane} -> Model: {model_name}")

    long_memory = load_long_memory()
    
    smart_memories = retrieve_smart_memories("tony", user_message, limit=5)
    smart_memory_text = ""
    if smart_memories:
        smart_memory_text = "\n".join([f"- {m['memory']}" for m in smart_memories])
        print(f"[smart_memory] Retrieved {len(smart_memories)} relevant memories")
    
    processed_message = user_message
    if mood and mood != "neutral":
        processed_message = with_mood(processed_message, mood)
    if truth:
        processed_message = truthify_prompt(processed_message)
    
    if lane == "code":
        active_persona = SYSTEM_PROMPT_CODE
    else:
        active_persona = PERSONAS.get(persona, STYLE_PROFILE)
    messages = [
        {"role": "system", "content": active_persona},
        {
            "role": "system",
            "content": (
                "Here is a snapshot of long term memory. "
                "Use it as background only.\n\n"
                + long_memory
            ),
        },
    ]
    
    if smart_memory_text:
        messages.append({
            "role": "system",
            "content": f"Relevant context from past conversations:\n{smart_memory_text}"
        })

    for item in chat_history[-10:]:
        messages.append(item)

    messages.append({"role": "user", "content": processed_message})

    completion = client.chat.completions.create(
        model=model_name,
        messages=messages,
        temperature=0.4,
    )

    return completion.choices[0].message.content.strip()


def call_multi_provider(user_message: str, provider: str = "openai", model: str = None, 
                        mode: str = "normal", mood: str = "neutral", truth: bool = False, 
                        persona: str = "default") -> dict:
    """Call AI with support for multiple providers: openai, claude, gemini."""
    
    lane = choose_lane(user_message, mode)
    long_memory = load_long_memory()
    
    smart_memories = retrieve_smart_memories("tony", user_message, limit=5)
    smart_memory_text = ""
    if smart_memories:
        smart_memory_text = "\n".join([f"- {m['memory']}" for m in smart_memories])
        logging.info(f"[smart_memory] Retrieved {len(smart_memories)} relevant memories")
    
    processed_message = user_message
    if mood and mood != "neutral":
        processed_message = with_mood(processed_message, mood)
    if truth:
        processed_message = truthify_prompt(processed_message)
    
    if lane == "code":
        system_prompt = SYSTEM_PROMPT_CODE
    else:
        system_prompt = PERSONAS.get(persona, STYLE_PROFILE)
    
    memory_context = f"Here is a snapshot of long term memory. Use it as background only.\n\n{long_memory}"
    if smart_memory_text:
        memory_context += f"\n\nRelevant context from past conversations:\n{smart_memory_text}"
    
    full_system = f"{system_prompt}\n\n{memory_context}"
    
    messages_for_openai = [
        {"role": "system", "content": full_system}
    ]
    for item in chat_history[-10:]:
        messages_for_openai.append(item)
    messages_for_openai.append({"role": "user", "content": processed_message})
    
    logging.info(f"[multi-provider] Provider: {provider}, Model: {model}, Lane: {lane}")
    
    try:
        if provider == "openai":
            chosen_model = model or "gpt-4o"
            completion = openai_client.chat.completions.create(
                model=chosen_model,
                messages=messages_for_openai,
                temperature=0.4,
            )
            text = completion.choices[0].message.content.strip()
            return {"text": text, "provider": "openai", "model": chosen_model}
        
        else:
            return {"error": f"Unknown provider: {provider}", "text": None}
    
    except Exception as e:
        logging.error(f"[multi-provider] Error with {provider}: {e}")
        return {"error": str(e), "text": None, "provider": provider}


@app.route("/api/models", methods=["GET"])
def api_models():
    """Get available AI models."""
    available = {}
    if os.getenv("OPENAI_API_KEY"):
        available["openai"] = AVAILABLE_MODELS["openai"]
    return jsonify({"models": available})


@app.route("/playbook-code")
def playbook_code():
    """Serve the revenue playbook HTML as plain text for easy copying."""
    playbook_path = BASE_DIR / "static" / "revenue-playbook.html"
    if playbook_path.exists():
        content = playbook_path.read_text()
        return Response(content, mimetype='text/plain')
    return "File not found", 404


@app.route("/whx-dubai")
def whx_dubai():
    """Serve the WHX Dubai trade show evaluation page."""
    file_path = BASE_DIR / "static" / "whx-dubai.html"
    if file_path.exists():
        content = file_path.read_text()
        return Response(content, mimetype='text/html')
    return "File not found", 404


@app.route("/whx-dubai-code")
def whx_dubai_code():
    """Serve the WHX Dubai evaluation HTML as plain text for easy copying."""
    file_path = BASE_DIR / "static" / "whx-dubai.html"
    if file_path.exists():
        content = file_path.read_text()
        return Response(content, mimetype='text/plain')
    return "File not found", 404


@app.route("/media-scorecard")
def media_scorecard():
    """Serve the Media Scorecard page."""
    file_path = BASE_DIR / "static" / "media-scorecard.html"
    if file_path.exists():
        content = file_path.read_text()
        return Response(content, mimetype='text/html')
    return "File not found", 404


@app.route("/media-scorecard-code")
def media_scorecard_code():
    """Serve the Media Scorecard HTML as plain text for easy copying."""
    file_path = BASE_DIR / "static" / "media-scorecard.html"
    if file_path.exists():
        content = file_path.read_text()
        return Response(content, mimetype='text/plain')
    return "File not found", 404


@app.route("/apollo-netlify")
def apollo_netlify():
    """Serve the Apollo Netlify Export page."""
    file_path = BASE_DIR / "static" / "apollo-netlify-export.html"
    if file_path.exists():
        content = file_path.read_text()
        return Response(content, mimetype='text/html')
    return "File not found", 404


@app.route("/apollo-netlify-code")
def apollo_netlify_code():
    """Serve the Apollo Netlify Export HTML as plain text for easy copying."""
    file_path = BASE_DIR / "static" / "apollo-netlify-export.html"
    if file_path.exists():
        content = file_path.read_text()
        return Response(content, mimetype='text/plain')
    return "File not found", 404


@app.route("/")
def index():
    try:
        from flask_login import current_user
        user = current_user if current_user.is_authenticated else None
    except:
        user = None
    return render_template("index.html", user=user)


@app.route("/chat")
def chat():
    return render_template("console.html")


@app.route("/tracker")
def tracker():
    return render_template("tracker.html")


@app.route("/research")
def research():
    return render_template("research.html")


@app.route("/storybrand-scanner")
def storybrand_scanner():
    return render_template("storybrand-scanner.html")


@app.route("/crm-code")
def crm_code():
    """Show the full CRM portal code for Netlify deployment"""
    import os
    code_path = os.path.join(app.static_folder, "internal/territory-rep-portal/index.html")
    with open(code_path, "r") as f:
        code = f.read()
    
    json_path = os.path.join(app.static_folder, "internal/territory-rep-portal/accounts.json")
    with open(json_path, "r") as f:
        json_code = f.read()
    
    return render_template("crm-code.html", html_code=code, json_code=json_code)


@app.route("/api/storybrand-scan", methods=["POST"])
def storybrand_scan():
    """StoryBrand 7-Part Framework Website Scanner"""
    try:
        data = request.get_json()
        url = data.get("url")
        
        if not url:
            return jsonify({"error": "URL required"}), 400
        
        if not url.startswith("http"):
            url = "https://" + url
        
        # Fetch website content
        try:
            response = requests.get(url, timeout=10, headers={
                "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
            })
            soup = BeautifulSoup(response.text, "html.parser")
            
            # Extract text content
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()
            
            text = soup.get_text(separator=" ", strip=True)[:8000]
            title = soup.title.string if soup.title else "Unknown"
            
        except Exception as e:
            return jsonify({"error": f"Could not fetch website: {str(e)}"}), 400
        
        # Analyze with OpenAI
        prompt = f"""Analyze this website through Donald Miller's 7-part StoryBrand framework.

Website: {url}
Title: {title}
Content: {text}

For each of the 7 parts, evaluate what's present and what's missing:

1. CHARACTER (Hero) - Who is the customer? Is it clear?
2. PROBLEM - External, internal, philosophical problems stated?
3. GUIDE - Does the brand show empathy + authority?
4. PLAN - Are there clear simple steps?
5. CALL TO ACTION - Is there a clear, direct CTA?
6. STAKES (Failure) - What happens if they don't act?
7. SUCCESS - What positive outcomes are promised?

Return a JSON object with this exact structure:
{{
  "company_name": "Company Name from website",
  "score": 0-100 clarity score,
  "summary": "2-3 sentence overall assessment",
  "framework": {{
    "character": {{"found": "what you found", "status": "pass|warning|fail", "recommendation": "specific action"}},
    "problem": {{"found": "what you found", "status": "pass|warning|fail", "recommendation": "specific action"}},
    "guide": {{"found": "what you found", "status": "pass|warning|fail", "recommendation": "specific action"}},
    "plan": {{"found": "what you found", "status": "pass|warning|fail", "recommendation": "specific action"}},
    "cta": {{"found": "what you found", "status": "pass|warning|fail", "recommendation": "specific action"}},
    "stakes": {{"found": "what you found", "status": "pass|warning|fail", "recommendation": "specific action"}},
    "success": {{"found": "what you found", "status": "pass|warning|fail", "recommendation": "specific action"}}
  }},
  "actions": [
    {{"priority": "high|medium|low", "text": "specific action to take"}},
    {{"priority": "high|medium|low", "text": "specific action to take"}},
    {{"priority": "high|medium|low", "text": "specific action to take"}}
  ]
}}

Be specific and actionable. Return ONLY valid JSON."""

        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            response_format={"type": "json_object"}
        )
        
        result = json.loads(response.choices[0].message.content)
        return jsonify(result)
        
    except Exception as e:
        logging.error(f"StoryBrand scan error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/scan", methods=["POST"])
def scan_website():
    """McKinsey-Grade Deep Business Intelligence Scanner using Claude Opus"""
    try:
        data = request.get_json()
        url = data.get("url")
        deep_scan = data.get("deep_scan", True)
        
        if not url:
            return jsonify({"error": "URL required"}), 400
        
        # Ensure URL has protocol
        if not url.startswith("http"):
            url = "https://" + url
        
        # Deep scan - comprehensive page list for McKinsey-grade analysis
        paths = [
            "",  # Homepage
            "/about", "/about-us", "/company", "/our-story",
            "/team", "/leadership", "/management", "/executives",
            "/products", "/services", "/solutions", "/offerings",
            "/pricing", "/plans", "/packages",
            "/customers", "/clients", "/case-studies", "/success-stories",
            "/careers", "/jobs", "/work-with-us",
            "/partners", "/integrations", "/ecosystem",
            "/blog", "/news", "/press", "/media",
            "/investors", "/investor-relations",
            "/contact", "/demo", "/get-started",
            "/features", "/capabilities", "/platform",
            "/industries", "/sectors", "/markets",
            "/resources", "/whitepapers", "/ebooks",
            "/faq", "/support", "/help",
            "/privacy", "/terms", "/legal"
        ]
        
        combined_text = ""
        pages_fetched = 0
        
        for path in paths:
            try:
                response = requests.get(url + path, timeout=6, headers={
                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
                })
                if response.status_code == 200:
                    soup = BeautifulSoup(response.text, "html.parser")
                    # Remove non-content elements
                    for el in soup(["script", "style", "nav", "footer", "header", "aside", "noscript"]):
                        el.decompose()
                    
                    page_text = soup.get_text(separator=" ", strip=True)
                    if len(page_text) > 100:  # Only include pages with real content
                        combined_text += f"\n\n=== PAGE: {path or '/'} ===\n{page_text}"
                        pages_fetched += 1
            except:
                continue
        
        # Allow more content for Claude's larger context
        combined_text = combined_text[:50000]
        
        if not combined_text.strip():
            return jsonify({"error": "Could not fetch website content"}), 400
        
        # Load the master prompt
        prompt_path = BASE_DIR / "scanner" / "prompts" / "masterScanPrompt.txt"
        with open(prompt_path, "r") as f:
            system_prompt = f.read()
        
        # Use OpenAI GPT-4o for McKinsey-grade deep analysis
        completion = openai_client.chat.completions.create(
            model="gpt-4o",
            temperature=0,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Perform a comprehensive McKinsey-grade strategic analysis of this company based on their website content. Be thorough, specific, and actionable.\n\nWEBSITE URL: {url}\n\nCONTENT FROM {pages_fetched} PAGES:\n{combined_text}"}
            ]
        )
        
        result_text = completion.choices[0].message.content
        
        # Clean up response - remove markdown code blocks if present
        if result_text.startswith("```"):
            result_text = result_text.split("```")[1]
            if result_text.startswith("json"):
                result_text = result_text[4:]
        result_text = result_text.strip()
        
        scan_result = json.loads(result_text)
        scan_result["_meta"] = {
            "scanned_url": url,
            "scanned_at": datetime.now().isoformat(),
            "pages_analyzed": pages_fetched,
            "analysis_model": "gpt-4o",
            "scan_type": "deep" if deep_scan else "standard"
        }
        
        # Add post-discovery actions
        scan_result["post_discovery_actions"] = {
            "must_change": [
                "Messaging must align to the actual decision owner identified in the scan",
                "Sales angle must reflect the primary operational risk surfaced",
                "Website must clarify who it is for and why within the first screen"
            ],
            "must_build": [
                "A single decision-owner CTA aligned to the sales strategy",
                "Proof tied to outcomes instead of capabilities",
                "A clear execution narrative across site and outreach"
            ],
            "must_stop": [
                "Targeting non-buyers despite engagement",
                "Explaining features before consequences",
                "Adding complexity instead of clarity"
            ],
            "execution_paths": {
                "internal": {
                    "what_it_requires": "Strong internal alignment and disciplined execution",
                    "risk_if_done_poorly": "Reinforcing the same problems with more effort"
                },
                "assisted": {
                    "where_help_accelerates": "Decision framing and execution alignment",
                    "time_to_impact": "Weeks instead of months"
                },
                "do_nothing": {
                    "what_stays_broken": "Unclear positioning and slow conversion",
                    "compounding_risk": "More activity with flat results"
                }
            }
        }
        
        return jsonify(scan_result)
        
    except json.JSONDecodeError as e:
        return jsonify({"error": f"Failed to parse AI response: {str(e)}"}), 500
    except Exception as e:
        logging.error(f"Scan error: {str(e)}")
        return jsonify({"error": str(e)}), 500


@app.route("/memory")
def memory():
    return render_template("memory.html")


@app.route("/tools")
def tools():
    return render_template("tools.html")


@app.route("/jobs")
def jobs():
    return render_template("jobs.html")


@app.route("/builder")
def builder():
    return render_template("builder.html")


@app.route("/storybrand")
def storybrand():
    return render_template("storybrand.html")


@app.route("/daily")
def daily():
    return render_template("daily.html")


@app.route("/crm")
def crm():
    return render_template("crm.html")


@app.route("/apollo-crm", methods=["GET", "POST"])
@app.route("/apollo-crm/<path:subpath>", methods=["GET", "POST"])
def apollo_crm_proxy(subpath=""):
    """Proxy Apollo CRM requests to the Node.js server on port 3001."""
    import requests as req
    target_url = f"http://localhost:3001/{subpath}"
    if request.query_string:
        target_url += f"?{request.query_string.decode()}"
    try:
        if request.method == "GET":
            resp = req.get(target_url, timeout=30)
        else:
            resp = req.post(target_url, json=request.get_json(), timeout=60)
        
        content_type = resp.headers.get("Content-Type", "text/html")
        return Response(resp.content, status=resp.status_code, content_type=content_type)
    except Exception as e:
        return jsonify({"error": f"Apollo CRM unavailable: {str(e)}"}), 503


@app.route("/contact")
def contact():
    return render_template("contact.html")


@app.route("/outbound")
def outbound():
    return render_template("outbound.html")


@app.route("/resume")
def resume():
    return render_template("resume.html")


# ============ RESUME BUILDER API ============
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document as DocxDocument
from docx.shared import Pt
from models import Resume, ResumeVersion

FLUFF_PHRASES = [
    "responsible for", "worked on", "helped with", "assisted",
    "participated in", "tasked with", "involved in"
]

STOP_WORDS = set([
    "the","and","or","a","an","to","of","in","for","with","on","at","by","from","as",
    "is","are","was","were","be","been","being","this","that","these","those",
    "you","your","we","our","they","their","i","me","my","it","its",
    "will","can","may","should","must","not","no","yes",
    "job","role","work","team","teams","experience","years","year",
    "responsible","responsibilities","required","preferred","skills"
])

def split_keywords(text):
    import re
    raw = re.sub(r'[^a-z0-9\s+#/.-]', ' ', (text or '').lower())
    raw = re.sub(r'\s+', ' ', raw).strip()
    tokens = [t for t in raw.split(' ') if len(t) >= 3 and t not in STOP_WORDS]
    freq = {}
    for t in tokens:
        freq[t] = freq.get(t, 0) + 1
    return [t for t, _ in sorted(freq.items(), key=lambda x: -x[1])[:40]]

def bullet_score(bullet):
    import re
    s = (bullet or '').strip()
    if not s:
        return 0
    score = 50
    lower = s.lower()
    for p in FLUFF_PHRASES:
        if p in lower:
            score -= 15
    if re.search(r'\b\d+(\.\d+)?%?\b', s):
        score += 20
    if re.search(r'\b(\$|m\b|mm\b|k\b|b\b|arr\b|roi\b|cagr\b|sql\b|mql\b)\b', s, re.I):
        score += 10
    if len(s) > 160:
        score -= 10
    if len(s) > 220:
        score -= 20
    if re.match(r'^(Led|Built|Owned|Drove|Improved|Reduced|Increased|Launched|Implemented|Automated|Designed|Managed|Developed|Scaled|Optimized|Standardized|Rebuilt|Created)\b', s, re.I):
        score += 10
    return max(0, min(100, score))

def generate_bullets_from_answers(answers):
    company = (answers.get('company') or '').strip()
    title = (answers.get('title') or '').strip()
    goal = (answers.get('businessGoal') or '').strip()
    change = (answers.get('whatChanged') or '').strip()
    metric = (answers.get('metricMoved') or '').strip()
    tools = (answers.get('tools') or '').strip()
    scope = (answers.get('scope') or '').strip()
    
    metric_line = metric if metric else "Improved key performance metrics with measurable operational impact"
    bullets = []
    
    if change:
        bullets.append(f"Led improvements at {company} by {change}{', resulting in ' + metric if metric else ''}")
    elif goal:
        bullets.append(f"Drove initiatives aligned to {goal}{', delivering ' + metric if metric else ''}")
    else:
        bullets.append(f"Delivered impact in the {title} role with focus on execution and measurable outcomes")
    
    if goal:
        bullets.append(f"Aligned cross functional work to {goal} and maintained stakeholder visibility on progress and results")
    bullets.append("Implemented repeatable workflows and standards to increase consistency, speed, and quality across the function")
    if tools:
        bullets.append(f"Used {tools} to streamline reporting, execution, and decision making for internal teams")
    if scope:
        bullets.append(f"Operated at {scope} scope to support scale and reduce risk in day to day operations")
    bullets.append(metric_line)
    
    seen = set()
    unique = []
    for b in bullets:
        key = b.lower()
        if key not in seen:
            seen.add(key)
            unique.append(b)
    return [{"text": t, "score": bullet_score(t)} for t in unique[:5]]

def compute_match_score(resume_data):
    jd = (resume_data.get('jobDescription') or '').strip()
    resume_text = json.dumps(resume_data).lower()
    if not jd:
        return {"score": 0, "keywords": [], "missing": [], "note": "No job description provided"}
    keywords = split_keywords(jd)
    hits = 0
    missing = []
    for k in keywords:
        if k in resume_text:
            hits += 1
        else:
            missing.append(k)
    score = round((hits / max(1, len(keywords))) * 100)
    return {"score": score, "keywords": keywords, "missing": missing[:15], "note": ""}

@app.route("/api/resumes", methods=["GET"])
def api_list_resumes():
    resumes = Resume.query.order_by(Resume.id.desc()).limit(50).all()
    return jsonify({"items": [{"id": r.id, "title": r.title, "created_at": r.created_at.isoformat()} for r in resumes]})

@app.route("/api/resumes", methods=["POST"])
def api_create_resume():
    data = request.get_json() or {}
    title = (data.get('title') or 'Untitled Resume').strip()
    resume_data = data.get('data', {})
    resume = Resume(title=title, data_json=json.dumps(resume_data))
    db.session.add(resume)
    db.session.commit()
    return jsonify({"id": resume.id})

@app.route("/api/resumes/<int:resume_id>", methods=["GET"])
def api_get_resume(resume_id):
    resume = Resume.query.get_or_404(resume_id)
    return jsonify({
        "id": resume.id,
        "title": resume.title,
        "created_at": resume.created_at.isoformat(),
        "data": json.loads(resume.data_json)
    })

@app.route("/api/resumes/<int:resume_id>", methods=["PUT"])
def api_update_resume(resume_id):
    resume = Resume.query.get_or_404(resume_id)
    data = request.get_json() or {}
    if data.get('title'):
        resume.title = data['title'].strip()
    if 'data' in data:
        resume.data_json = json.dumps(data['data'])
    db.session.commit()
    return jsonify({"ok": True})

@app.route("/api/resumes/<int:resume_id>/versions", methods=["GET"])
def api_list_versions(resume_id):
    versions = ResumeVersion.query.filter_by(resume_id=resume_id).order_by(ResumeVersion.id.desc()).limit(50).all()
    return jsonify({"items": [{"id": v.id, "label": v.label, "created_at": v.created_at.isoformat()} for v in versions]})

@app.route("/api/resumes/<int:resume_id>/versions", methods=["POST"])
def api_create_version(resume_id):
    data = request.get_json() or {}
    label = (data.get('label') or 'Snapshot').strip()
    resume_data = data.get('data', {})
    version = ResumeVersion(resume_id=resume_id, label=label, data_json=json.dumps(resume_data))
    db.session.add(version)
    db.session.commit()
    return jsonify({"id": version.id})

@app.route("/api/versions/<int:version_id>", methods=["GET"])
def api_get_version(version_id):
    version = ResumeVersion.query.get_or_404(version_id)
    return jsonify({
        "id": version.id,
        "resume_id": version.resume_id,
        "label": version.label,
        "created_at": version.created_at.isoformat(),
        "data": json.loads(version.data_json)
    })

@app.route("/api/bullets/generate", methods=["POST"])
def api_generate_bullets():
    data = request.get_json() or {}
    answers = data.get('answers', {})
    bullets = generate_bullets_from_answers(answers)
    return jsonify({"bullets": bullets})

@app.route("/api/match/score", methods=["POST"])
def api_match_score():
    data = request.get_json() or {}
    resume_data = data.get('data', {})
    result = compute_match_score(resume_data)
    return jsonify(result)

@app.route("/api/job-recommendations", methods=["POST"])
def api_job_recommendations():
    """AI-powered job search recommendations using OpenAI"""
    data = request.get_json() or {}
    focus = data.get('focus', 'general')
    
    tony_profile = """
    Tony Beal - RevOps & Growth Leader
    - 20+ years in B2B manufacturing, industrial sales, and revenue operations
    - Expert in CRM systems (HubSpot, Salesforce), pipeline management, sales enablement
    - StoryBrand Certified Guide, Challenger Sales trained
    - Strong background in account-based selling, territory management, data-driven forecasting
    - Experience with pharmaceutical, nutraceutical, industrial tooling markets
    - Skills: Revenue operations, GTM strategy, demand generation, sales process optimization
    - Leadership: Built and managed sales teams, implemented CRM transformations
    - Location preference: Remote, Hybrid, or St. Louis/Kansas City area
    """
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": f"""You are a career strategist helping Tony find his next role.

{tony_profile}

Generate personalized job search recommendations. For each recommendation, provide:
1. A specific job title to search for
2. Why this role fits Tony's background
3. Companies or industries to target
4. The exact search query to use

Format your response as JSON with this structure:
{{
  "strategy": "Brief overall strategy (1-2 sentences)",
  "recommendations": [
    {{
      "title": "Job Title",
      "why_it_fits": "Why this matches Tony",
      "target_companies": "Types of companies to look for",
      "search_query": "Exact search string to use",
      "priority": "high/medium/low"
    }}
  ],
  "pro_tips": ["Tip 1", "Tip 2", "Tip 3"]
}}

Generate 6-8 specific, actionable recommendations."""
                },
                {
                    "role": "user",
                    "content": f"Generate job search recommendations for Tony. Focus area: {focus}"
                }
            ],
            max_tokens=1500,
            temperature=0.7
        )
        
        content = response.choices[0].message.content
        
        try:
            import re
            json_match = re.search(r'\{[\s\S]*\}', content)
            if json_match:
                recommendations = json.loads(json_match.group())
            else:
                recommendations = {"raw": content}
        except json.JSONDecodeError:
            recommendations = {"raw": content}
        
        return jsonify(recommendations)
        
    except Exception as e:
        logging.error(f"Job recommendations error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/job-tailor", methods=["POST"])
def api_job_tailor():
    """Generate tailored resume summary, bullets, cover letter, and LinkedIn message"""
    data = request.get_json() or {}
    job_description = data.get('description', '')
    
    if not job_description or len(job_description) < 100:
        return jsonify({"error": "Please paste the full job description (at least 100 characters)"}), 400
    
    tony_experience = """
    TONY BEAL - Revenue Operations & Growth Leader
    
    CORE POSITIONING:
    Revenue Operations & Growth Leader who builds repeatable revenue systems for complex B2B companies
    
    KEY ACHIEVEMENTS:
    - Built $6M+ pipeline in manufacturing through strategic territory development
    - Owned full sales cycle from lead generation to close in complex B2B environments
    - Led territory & account expansion initiatives driving 40%+ growth
    - Managed channel & distributor partnerships across multiple regions
    
    REVOPS & SYSTEMS:
    - Expert in CRM platforms (HubSpot, Salesforce, Apollo.io)
    - Built persona scoring & tiering systems for pipeline prioritization
    - Created dashboards, forecasting models, and attribution frameworks
    - Implemented workflow automation and AI-powered prospecting tools
    
    MARKET & BRAND:
    - Grew LinkedIn presence generating consistent inbound leads
    - StoryBrand Certified Guide - developed customer-centric messaging
    - Created outreach scripts, prospecting systems, and GTM frameworks
    - Challenger Sales trained - positioned solutions vs. pitching products
    
    INDUSTRIES: Pharmaceutical, Nutraceutical, Industrial Manufacturing, Tooling, B2B Technology
    
    CERTIFICATIONS: StoryBrand Certified Guide, Challenger Sales, HubSpot CRM
    """
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": f"""You are Tony's career strategist. Your job is to help him be the FIRST and BEST applicant for every role.

Tony's Experience:
{tony_experience}

When given a job description, generate FOUR things:

1. SUMMARY (3-4 sentences)
   - Mirror the language of the job posting
   - Lead with their most important requirement
   - Quantify achievements that match their needs
   - End with what makes Tony uniquely valuable for THIS role

2. TAILORED BULLETS (5-6 bullets)
   - Each bullet should directly address a requirement from the job posting
   - Use their exact terminology and keywords
   - Lead with action verbs and quantified results
   - Show progression and impact

3. COVER LETTER (3 paragraphs, 150-200 words)
   - Opening: Hook them with why THIS role at THIS company
   - Middle: 2-3 specific achievements that prove you can do their job
   - Close: Express genuine interest and request conversation

4. LINKEDIN MESSAGE (3 lines max)
   - Direct, personal, references the specific role
   - Shows you understand what they need
   - Ends with soft call to action

Format your response as JSON:
{{
  "summary": "...",
  "bullets": ["bullet 1", "bullet 2", ...],
  "cover_letter": "...",
  "linkedin_message": "..."
}}

Be specific. Mirror their language. Make Tony sound like the obvious choice."""
                },
                {
                    "role": "user",
                    "content": f"Tailor Tony's application for this job:\n\n{job_description[:5000]}"
                }
            ],
            max_tokens=2000,
            temperature=0.7
        )
        
        content = response.choices[0].message.content
        
        try:
            import re
            json_match = re.search(r'\{[\s\S]*\}', content)
            if json_match:
                result = json.loads(json_match.group())
            else:
                result = {"raw": content}
        except json.JSONDecodeError:
            result = {"raw": content}
        
        return jsonify(result)
        
    except Exception as e:
        logging.error(f"Job tailor error: {e}")
        return jsonify({"error": str(e)}), 500

@app.route("/api/job-analyze", methods=["POST"])
def api_job_analyze():
    """AI-powered job description analysis for Tony's Job Engine"""
    data = request.get_json() or {}
    job_description = data.get('description', '')
    
    if not job_description or len(job_description) < 50:
        return jsonify({"error": "Job description too short"}), 400
    
    tony_profile = """
    Tony Beal - RevOps & Growth Leader
    - 20+ years in B2B manufacturing, industrial sales, and revenue operations
    - Expert in CRM systems (HubSpot, Salesforce), pipeline management, sales enablement
    - StoryBrand Certified Guide, Challenger Sales trained
    - Strong background in account-based selling, territory management, data-driven forecasting
    - Experience with pharmaceutical, nutraceutical, industrial tooling markets
    - Skills: Revenue operations, GTM strategy, demand generation, sales process optimization
    - Leadership: Built and managed sales teams, implemented CRM transformations
    """
    
    keywords = [
        "revenue operations", "revops", "sales operations", "business development",
        "head of growth", "gtm", "go-to-market", "crm", "hubspot", "salesforce",
        "manufacturing", "b2b", "pipeline", "demand generation", "sales enablement",
        "territory", "account executive", "customer success", "storybrand", "challenger"
    ]
    
    matched_keywords = [kw for kw in keywords if kw.lower() in job_description.lower()]
    keyword_score = len(matched_keywords)
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "system",
                    "content": f"""You are a career strategist helping Tony analyze job opportunities.
                    
Tony's Profile:
{tony_profile}

Analyze job descriptions and provide:
1. FIT SCORE (1-10): How well does this role match Tony's experience?
2. STRENGTHS: 3-4 specific ways Tony's background aligns with this role
3. GAPS: Any requirements Tony may need to address or reframe
4. POSITIONING: How should Tony tailor his resume/cover letter for this role?
5. TALKING POINTS: 2-3 key achievements from Tony's background to highlight
6. RED FLAGS: Any concerns about the role or company (optional)

Be direct, strategic, and actionable. Focus on how to win this opportunity."""
                },
                {
                    "role": "user",
                    "content": f"Analyze this job posting for Tony:\n\n{job_description[:4000]}"
                }
            ],
            max_tokens=1000,
            temperature=0.7
        )
        
        analysis = response.choices[0].message.content
        
        return jsonify({
            "analysis": analysis,
            "keyword_score": keyword_score,
            "matched_keywords": matched_keywords,
            "total_keywords": len(keywords)
        })
        
    except Exception as e:
        logging.error(f"Job analysis error: {e}")
        return jsonify({
            "error": "AI analysis failed",
            "keyword_score": keyword_score,
            "matched_keywords": matched_keywords
        }), 500

@app.route("/api/export/pdf", methods=["POST"])
def api_export_pdf():
    data = request.get_json() or {}
    resume_data = data.get('data', {})
    profile = resume_data.get('profile', {})
    experiences = resume_data.get('experience', [])
    skills = resume_data.get('skills', [])
    summary = (resume_data.get('summary') or '').strip()
    
    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y = height - 48
    
    c.setFont("Helvetica-Bold", 18)
    c.drawString(48, y, profile.get('name', 'Your Name'))
    y -= 20
    
    c.setFont("Helvetica", 10)
    contact_parts = [profile.get('location', ''), profile.get('email', ''), profile.get('phone', ''), profile.get('linkedin', ''), profile.get('website', '')]
    contact_line = ' | '.join([p for p in contact_parts if p])
    c.drawString(48, y, contact_line)
    y -= 30
    
    if summary:
        c.setFont("Helvetica-Bold", 11)
        c.drawString(48, y, "SUMMARY")
        y -= 15
        c.setFont("Helvetica", 10)
        for line in summary.split('\n')[:3]:
            c.drawString(48, y, line[:90])
            y -= 14
        y -= 10
    
    if skills:
        c.setFont("Helvetica-Bold", 11)
        c.drawString(48, y, "SKILLS")
        y -= 15
        c.setFont("Helvetica", 10)
        c.drawString(48, y, ', '.join(skills)[:100])
        y -= 25
    
    if experiences:
        c.setFont("Helvetica-Bold", 11)
        c.drawString(48, y, "EXPERIENCE")
        y -= 20
        
        for exp in experiences[:5]:
            if y < 100:
                c.showPage()
                y = height - 48
            
            c.setFont("Helvetica-Bold", 10)
            header = f"{exp.get('title', '')} | {exp.get('company', '')}"
            c.drawString(48, y, header[:80])
            y -= 14
            
            c.setFont("Helvetica", 9)
            dates = f"{exp.get('start', '')} to {exp.get('end', '')}"
            if exp.get('location'):
                dates += f" | {exp.get('location')}"
            c.drawString(48, y, dates)
            y -= 14
            
            bullets = exp.get('bullets', [])
            for b in bullets[:5]:
                text = b if isinstance(b, str) else b.get('text', '')
                c.drawString(58, y, f"• {text[:85]}")
                y -= 12
            y -= 10
    
    c.save()
    buffer.seek(0)
    return Response(buffer.getvalue(), mimetype='application/pdf', headers={'Content-Disposition': 'attachment; filename=resume.pdf'})

@app.route("/api/export/docx", methods=["POST"])
def api_export_docx():
    data = request.get_json() or {}
    resume_data = data.get('data', {})
    profile = resume_data.get('profile', {})
    experiences = resume_data.get('experience', [])
    skills = resume_data.get('skills', [])
    summary = (resume_data.get('summary') or '').strip()
    
    doc = DocxDocument()
    
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(profile.get('name', 'Your Name'))
    name_run.bold = True
    name_run.font.size = Pt(16)
    
    contact_parts = [profile.get('location', ''), profile.get('email', ''), profile.get('phone', ''), profile.get('linkedin', ''), profile.get('website', '')]
    contact_line = ' | '.join([p for p in contact_parts if p])
    if contact_line:
        doc.add_paragraph(contact_line)
    
    if summary:
        heading = doc.add_paragraph()
        heading.add_run("SUMMARY").bold = True
        doc.add_paragraph(summary)
    
    if skills:
        heading = doc.add_paragraph()
        heading.add_run("SKILLS").bold = True
        doc.add_paragraph(', '.join(skills))
    
    if experiences:
        heading = doc.add_paragraph()
        heading.add_run("EXPERIENCE").bold = True
        
        for exp in experiences:
            exp_para = doc.add_paragraph()
            exp_para.add_run(f"{exp.get('title', '')} | {exp.get('company', '')}").bold = True
            
            dates = f"{exp.get('start', '')} to {exp.get('end', '')}"
            if exp.get('location'):
                dates += f" | {exp.get('location')}"
            doc.add_paragraph(dates)
            
            bullets = exp.get('bullets', [])
            for b in bullets:
                text = b if isinstance(b, str) else b.get('text', '')
                doc.add_paragraph(f"• {text}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return Response(buffer.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers={'Content-Disposition': 'attachment; filename=resume.docx'})

# ============ END RESUME BUILDER API ============


@app.route("/journal")
def journal():
    return render_template("journal.html")


@app.route("/visual-builder")
def visual_builder():
    return render_template("visual-builder.html")


@app.route("/api/visual-builder/generate", methods=["POST"])
@rate_limit(20)
@require_api_key
def api_visual_builder_generate():
    """Generate LinkedIn post content from a topic."""
    try:
        if not os.getenv("OPENAI_API_KEY"):
            return jsonify({"error": "OpenAI API key not configured"}), 500
        
        data = request.get_json() or {}
        topic = data.get("topic", "").strip()
        
        if not topic:
            return jsonify({"error": "Missing topic"}), 400
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": """You are a LinkedIn content expert. Generate scroll-stopping visual post content.
Return JSON only with these fields:
- headline: A bold, punchy headline (under 12 words)
- body: 3-5 supporting points, one per line
- template: One of: statement, list, framework, metric, beforeAfter, steps, quote, doNotDo
- theme: One of: midnight, neonSlate, steel, cleanLight, warmDark, ocean"""},
                {"role": "user", "content": f"Create LinkedIn visual post content about: {topic}"}
            ],
            response_format={"type": "json_object"},
            temperature=0.7
        )
        
        result = json.loads(response.choices[0].message.content)
        return jsonify(result)
    except Exception as e:
        logging.error(f"[visual-builder/generate] Error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/visual-builder/rewrite", methods=["POST"])
@rate_limit(20)
@require_api_key
def api_visual_builder_rewrite():
    """Rewrite existing copy for LinkedIn posts."""
    try:
        if not os.getenv("OPENAI_API_KEY"):
            return jsonify({"error": "OpenAI API key not configured"}), 500
        
        data = request.get_json() or {}
        headline = data.get("headline", "").strip()
        body = data.get("body", "").strip()
        
        if not headline and not body:
            return jsonify({"error": "Missing headline or body"}), 400
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": """You are a LinkedIn copywriting expert. Rewrite the given content to be more engaging and scroll-stopping.
Return JSON only with these fields:
- headline: A punchier, more attention-grabbing headline (under 12 words)
- body: Improved supporting points, one per line"""},
                {"role": "user", "content": f"Rewrite this LinkedIn post:\nHeadline: {headline}\nBody: {body}"}
            ],
            response_format={"type": "json_object"},
            temperature=0.7
        )
        
        result = json.loads(response.choices[0].message.content)
        return jsonify(result)
    except Exception as e:
        logging.error(f"[visual-builder/rewrite] Error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/visual-builder/suggest-layout", methods=["POST"])
@rate_limit(20)
@require_api_key
def api_visual_builder_suggest_layout():
    """Suggest optimal layout for content."""
    try:
        if not os.getenv("OPENAI_API_KEY"):
            return jsonify({"error": "OpenAI API key not configured"}), 500
        
        data = request.get_json() or {}
        headline = data.get("headline", "").strip()
        body = data.get("body", "").strip()
        
        if not headline:
            return jsonify({"error": "Missing headline"}), 400
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": """You are a visual design expert for LinkedIn content. Analyze the content and suggest the best layout.
Return JSON only with these fields:
- template: Best template from: statement, list, framework, metric, beforeAfter, steps, quote, doNotDo
- theme: Best theme from: midnight, neonSlate, steel, cleanLight, warmDark, ocean
- accent: Hex color code that fits the content mood
- reason: One sentence explaining why this layout works"""},
                {"role": "user", "content": f"Suggest the best layout for:\nHeadline: {headline}\nBody: {body}"}
            ],
            response_format={"type": "json_object"},
            temperature=0.5
        )
        
        result = json.loads(response.choices[0].message.content)
        return jsonify(result)
    except Exception as e:
        logging.error(f"[visual-builder/suggest-layout] Error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/settings")
def settings():
    return render_template("settings.html")


@app.route("/opportunity-engine")
def opportunity_engine():
    return render_template("opportunity-engine.html")


@app.route("/natoli-revops")
def natoli_revops():
    return render_template("natoli-revops/index.html")


@app.route("/natoli-revops/system")
def natoli_revops_system():
    return render_template("natoli-revops/system.html")


@app.route("/natoli-revops/targeting")
def natoli_revops_targeting():
    return render_template("natoli-revops/targeting.html")


@app.route("/natoli-revops/persona")
def natoli_revops_persona():
    return render_template("natoli-revops/persona.html")


@app.route("/natoli-revops/results")
def natoli_revops_results():
    return render_template("natoli-revops/results.html")


@app.route("/natoli-revops/downloads")
def natoli_revops_downloads():
    return render_template("natoli-revops/downloads.html")


@app.route("/natoli-revops/formal-intro")
def natoli_revops_formal_intro():
    return redirect("/natoli-revops/message-center")


@app.route("/natoli-revops/message-center")
def natoli_revops_message_center():
    return render_template("natoli-revops/message-center.html")


@app.route("/simple-chat")
def simple_chat():
    return render_template("simple_chat.html")


@app.route("/api/simple-chat", methods=["POST"])
@rate_limit(30)
@require_api_key
def api_simple_chat():
    """Simple chat endpoint - minimal OpenAI integration."""
    try:
        data = request.get_json() or {}
        message = data.get("message", "").strip()
        
        if not message:
            return jsonify({"error": "Missing message"}), 400
        
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            return jsonify({"error": "OpenAI API key not configured"}), 500
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": message}],
            max_tokens=2000
        )
        
        text = response.choices[0].message.content if response.choices else "No response"
        return jsonify({"text": text})
    except Exception as e:
        logging.error(f"[simple-chat] Error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/health", methods=["GET"])
def api_health():
    """Health check endpoint for connection status."""
    health = {
        "status": "ok",
        "timestamp": datetime.now().isoformat(),
        "services": {}
    }
    
    if DATABASE_URL:
        try:
            with get_db_context() as conn:
                if conn:
                    cur = conn.cursor()
                    cur.execute("SELECT 1")
                    health["services"]["database"] = "connected"
                else:
                    health["services"]["database"] = "failed"
        except Exception as e:
            health["services"]["database"] = f"error: {str(e)[:50]}"
    else:
        health["services"]["database"] = "not configured"
    
    health["services"]["smart_memory"] = "enabled" if SMART_MEMORY_ENABLED else "disabled"
    
    if os.getenv("OPENAI_API_KEY"):
        health["services"]["openai"] = "configured"
    else:
        health["services"]["openai"] = "not configured"
    
    all_ok = all(v in ["connected", "enabled", "configured", "not configured", "disabled"] 
                 for v in health["services"].values())
    health["status"] = "ok" if all_ok else "degraded"
    
    return jsonify(health)


@app.route("/api/export", methods=["GET"])
@require_api_key
@rate_limit(max_per_minute=5)
def api_export():
    """Export all database data as JSON for backup/download."""
    if not DATABASE_URL:
        return jsonify({"error": "Database not configured"}), 500
    
    export_data = {
        "exported_at": datetime.now().isoformat(),
        "crm_contacts": [],
        "crm_activities": [],
        "journal_entries": [],
        "contact_submissions": []
    }
    
    try:
        with get_db_context() as conn:
            if not conn:
                return jsonify({"error": "Database connection failed"}), 500
            
            cur = conn.cursor()
            
            cur.execute("SELECT * FROM crm_connections ORDER BY created_at DESC")
            rows = cur.fetchall()
            for row in rows:
                record = dict(row)
                for key, val in record.items():
                    if isinstance(val, datetime):
                        record[key] = val.isoformat()
                    elif hasattr(val, 'isoformat'):
                        record[key] = val.isoformat()
                export_data["crm_contacts"].append(record)
            
            cur.execute("SELECT * FROM crm_activity ORDER BY date DESC")
            rows = cur.fetchall()
            for row in rows:
                record = dict(row)
                for key, val in record.items():
                    if isinstance(val, datetime):
                        record[key] = val.isoformat()
                    elif hasattr(val, 'isoformat'):
                        record[key] = val.isoformat()
                export_data["crm_activities"].append(record)
            
            cur.execute("SELECT * FROM journal_entries ORDER BY date DESC")
            rows = cur.fetchall()
            for row in rows:
                record = dict(row)
                for key, val in record.items():
                    if isinstance(val, datetime):
                        record[key] = val.isoformat()
                    elif hasattr(val, 'isoformat'):
                        record[key] = val.isoformat()
                export_data["journal_entries"].append(record)
            
            cur.execute("SELECT * FROM contact_submissions ORDER BY created_at DESC")
            rows = cur.fetchall()
            for row in rows:
                record = dict(row)
                for key, val in record.items():
                    if isinstance(val, datetime):
                        record[key] = val.isoformat()
                    elif hasattr(val, 'isoformat'):
                        record[key] = val.isoformat()
                export_data["contact_submissions"].append(record)
        
        response = jsonify(export_data)
        response.headers["Content-Disposition"] = f"attachment; filename=tonyos_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        return response
    
    except Exception as e:
        logging.error(f"[export] Error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/agent/journal-insight", methods=["POST"])
def proxy_agent_journal_insight():
    """Proxy endpoint to forward requests to the internal agent server."""
    try:
        data = request.get_json() or {}
        resp = requests.post(
            "http://localhost:6000/agent/journal-insight",
            json=data,
            timeout=30
        )
        return jsonify(resp.json()), resp.status_code
    except requests.exceptions.ConnectionError:
        return jsonify({"error": "Agent server not available"}), 503
    except Exception as e:
        logging.error(f"[agent-proxy] Error: {e}")
        return jsonify({"error": str(e)}), 500


def init_journal_table():
    """Initialize journal entries table."""
    if not DATABASE_URL:
        return
    try:
        conn = get_db()
        if not conn:
            return
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS journal_entries (
                id SERIAL PRIMARY KEY,
                user_id VARCHAR(100) DEFAULT 'tony',
                date DATE NOT NULL,
                mood VARCHAR(20),
                content TEXT,
                gratitude JSONB DEFAULT '[]',
                focus TEXT,
                reflection TEXT,
                tags JSONB DEFAULT '[]',
                created_at TIMESTAMP DEFAULT NOW(),
                updated_at TIMESTAMP DEFAULT NOW(),
                UNIQUE(user_id, date)
            )
        """)
        conn.commit()
        conn.close()
        print("[journal] Table initialized")
    except Exception as e:
        print(f"[journal] Table init error: {e}")


@app.route("/api/journal", methods=["GET"])
def get_journal_entries():
    """Get all journal entries."""
    if not DATABASE_URL:
        return jsonify({"entries": []})
    try:
        conn = get_db()
        if not conn:
            return jsonify({"entries": []})
        cur = conn.cursor()
        cur.execute("""
            SELECT id, date, mood, content, gratitude, focus, reflection, tags, updated_at
            FROM journal_entries
            WHERE user_id = 'tony'
            ORDER BY date DESC
        """)
        rows = cur.fetchall()
        conn.close()
        entries = []
        for row in rows:
            entries.append({
                "id": row["id"],
                "date": row["date"].isoformat() if row["date"] else None,
                "mood": row["mood"],
                "content": row["content"],
                "gratitude": row["gratitude"] or [],
                "focus": row["focus"],
                "reflection": row["reflection"],
                "tags": row["tags"] or [],
                "updated_at": row["updated_at"].isoformat() if row["updated_at"] else None
            })
        return jsonify({"entries": entries})
    except Exception as e:
        print(f"[journal] Get error: {e}")
        return jsonify({"entries": [], "error": str(e)})


@app.route("/api/journal", methods=["POST"])
def save_journal_entry():
    """Save or update a journal entry."""
    data = request.get_json(force=True) or {}
    date_str = data.get("date")
    
    if not date_str:
        return jsonify({"error": "Date required"}), 400
    
    if not DATABASE_URL:
        return jsonify({"error": "Database not configured"}), 500
    
    try:
        conn = get_db()
        if not conn:
            return jsonify({"error": "Database connection failed"}), 500
        cur = conn.cursor()
        
        cur.execute("""
            INSERT INTO journal_entries (user_id, date, mood, content, gratitude, focus, reflection, tags, updated_at)
            VALUES ('tony', %s, %s, %s, %s, %s, %s, %s, NOW())
            ON CONFLICT (user_id, date)
            DO UPDATE SET
                mood = EXCLUDED.mood,
                content = EXCLUDED.content,
                gratitude = EXCLUDED.gratitude,
                focus = EXCLUDED.focus,
                reflection = EXCLUDED.reflection,
                tags = EXCLUDED.tags,
                updated_at = NOW()
            RETURNING id
        """, (
            date_str,
            data.get("mood"),
            data.get("content"),
            json.dumps(data.get("gratitude", [])),
            data.get("focus"),
            data.get("reflection"),
            json.dumps(data.get("tags", []))
        ))
        
        result = cur.fetchone()
        entry_id = result["id"] if result else None
        conn.commit()
        conn.close()
        
        return jsonify({"success": True, "id": entry_id})
    except Exception as e:
        print(f"[journal] Save error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/journal/<date_key>", methods=["DELETE"])
def delete_journal_entry(date_key):
    """Delete a journal entry by date."""
    if not DATABASE_URL:
        return jsonify({"error": "Database not configured"}), 500
    
    try:
        conn = get_db()
        if not conn:
            return jsonify({"error": "Database connection failed"}), 500
        cur = conn.cursor()
        cur.execute("""
            DELETE FROM journal_entries
            WHERE user_id = 'tony' AND date = %s
            RETURNING id
        """, (date_key,))
        result = cur.fetchone()
        conn.commit()
        conn.close()
        
        if result:
            return jsonify({"success": True, "deleted_id": result["id"]})
        return jsonify({"error": "Entry not found"}), 404
    except Exception as e:
        print(f"[journal] Delete error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/journal/prompt", methods=["POST"])
def generate_journal_prompt():
    """Generate an AI writing prompt based on recent entries."""
    data = request.get_json(force=True) or {}
    recent_entries = data.get("recent_entries", [])
    
    context = ""
    if recent_entries:
        context = f"Recent journal themes: {', '.join([e[:100] for e in recent_entries if e])}"
    
    prompt = f"""Generate a single thought-provoking journal writing prompt.
{context}

Rules:
- One sentence only
- Make it personal and introspective
- Avoid cliches
- No emojis
- Focus on self-discovery, growth, or reflection

Return only the prompt text, nothing else."""

    try:
        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a thoughtful journaling coach."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.8
        )
        prompt_text = completion.choices[0].message.content.strip().strip('"')
        return jsonify({"prompt": prompt_text})
    except Exception as e:
        return jsonify({"prompt": "What would you do differently if you could start today over?", "error": str(e)})


init_journal_table()


def init_contact_table():
    """Initialize contact submissions table."""
    if not DATABASE_URL:
        return
    try:
        conn = get_db()
        if not conn:
            return
        cur = conn.cursor()
        cur.execute("""
            CREATE TABLE IF NOT EXISTS contact_submissions (
                id SERIAL PRIMARY KEY,
                name VARCHAR(200) NOT NULL,
                email VARCHAR(200) NOT NULL,
                company VARCHAR(200),
                phone VARCHAR(50),
                message TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT NOW()
            )
        """)
        conn.commit()
        conn.close()
        print("[contact] Table initialized")
    except Exception as e:
        print(f"[contact] Table init error: {e}")


init_contact_table()


@app.route("/api/contact", methods=["POST"])
def api_contact_submit():
    """Handle contact form submissions."""
    data = request.get_json(force=True) or {}
    
    name = (data.get("name") or "").strip()
    email = (data.get("email") or "").strip()
    company = (data.get("company") or "").strip()
    phone = (data.get("phone") or "").strip()
    message = (data.get("message") or "").strip()
    
    if not name or not email or not message:
        return jsonify({"success": False, "error": "Name, email, and message are required"}), 400
    
    if not DATABASE_URL:
        return jsonify({"success": False, "error": "Database not configured"}), 500
    
    try:
        conn = get_db()
        if not conn:
            return jsonify({"success": False, "error": "Database connection failed"}), 500
        cur = conn.cursor()
        
        cur.execute("""
            INSERT INTO contact_submissions (name, email, company, phone, message)
            VALUES (%s, %s, %s, %s, %s)
            RETURNING id, created_at
        """, (name, email, company, phone, message))
        
        result = cur.fetchone()
        submission_id = result["id"] if result else None
        created_at = result["created_at"].isoformat() if result and result.get("created_at") else None
        conn.commit()
        conn.close()
        
        logging.info(f"[contact] New submission from {email}")
        
        return jsonify({
            "success": True,
            "id": submission_id,
            "created_at": created_at
        })
    except Exception as e:
        logging.error(f"[contact] Submit error: {e}")
        return jsonify({"success": False, "error": str(e)}), 500


LOG_DIR = Path("logs")
LOG_DIR.mkdir(exist_ok=True)
DAILY_LOG_FILE = LOG_DIR / "daily.jsonl"


def write_daily_log(record: dict) -> None:
    record["timestamp"] = datetime.now().isoformat(timespec="seconds")
    with open(DAILY_LOG_FILE, "a", encoding="utf-8") as f:
        f.write(json.dumps(record, ensure_ascii=False) + "\n")


def build_daily_prompt(one_win: str, pressure: str, next_step: str) -> str:
    return f"""
You are TonyOS, a practical daily command center.

Input:
1) Today's One Win: {one_win}
2) Top Pressure: {pressure}
3) Next Action: {next_step}

Output EXACTLY these 3 sections, each 1 to 3 sentences max.
No bullets unless needed.
No emojis.
No hype.

1) One Sentence Plan:
2) Blocker Breaker:
3) 10 Minute Sprint:
""".strip()


@app.route("/api/daily", methods=["POST"])
def api_daily():
    data = request.get_json(force=True) or {}
    one_win = (data.get("one_win") or "").strip()
    pressure = (data.get("pressure") or "").strip()
    next_step = (data.get("next_step") or "").strip()

    if not one_win or not pressure or not next_step:
        return jsonify({"error": "Fill out all 3 fields"}), 400

    prompt = build_daily_prompt(one_win, pressure, next_step)

    try:
        completion = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are TonyOS, a practical daily command center."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4
        )
        text = completion.choices[0].message.content.strip()
    except Exception as e:
        return jsonify({"error": f"OpenAI call failed: {str(e)}"}), 500

    record = {
        "one_win": one_win,
        "pressure": pressure,
        "next_step": next_step,
        "result": text
    }
    write_daily_log(record)

    return jsonify({"result": text})


def clean_phone(s):
    """Extract phone number from text."""
    if not s:
        return ""
    match = re.search(r'(\+?\d[\d\s().\-]{7,}\d)', str(s))
    return match.group(1).strip() if match else ""


def pick_best_address(soup):
    """Try to find address from page."""
    addr_tag = soup.find("address")
    if addr_tag:
        addr = " ".join(addr_tag.get_text().split())
        if len(addr) > 10:
            return addr
    
    body_text = " ".join(soup.get_text().split())
    match = re.search(r'\d{1,6}\s+[A-Za-z0-9.\- ]+,\s*[A-Za-z.\- ]+,\s*[A-Z]{2}\s*\d{5}(-\d{4})?', body_text)
    return match.group(0).strip() if match else ""


def extract_from_jsonld(soup):
    """Extract company info from JSON-LD schema."""
    name = ""
    phone = ""
    address = ""
    
    for script in soup.find_all("script", type="application/ld+json"):
        try:
            data = json.loads(script.string)
            if isinstance(data, list):
                for item in data:
                    if isinstance(item, dict):
                        if not name and item.get("name"):
                            name = str(item["name"])
                        if not phone and (item.get("telephone") or item.get("phone")):
                            phone = str(item.get("telephone") or item.get("phone"))
                        addr_obj = item.get("address")
                        if not address and addr_obj:
                            if isinstance(addr_obj, str):
                                address = addr_obj
                            elif isinstance(addr_obj, dict):
                                parts = [
                                    addr_obj.get("streetAddress"),
                                    addr_obj.get("addressLocality"),
                                    addr_obj.get("addressRegion"),
                                    addr_obj.get("postalCode")
                                ]
                                address = ", ".join(p for p in parts if p)
            elif isinstance(data, dict):
                if not name and data.get("name"):
                    name = str(data["name"])
                if not phone and (data.get("telephone") or data.get("phone")):
                    phone = str(data.get("telephone") or data.get("phone"))
                addr_obj = data.get("address")
                if not address and addr_obj:
                    if isinstance(addr_obj, str):
                        address = addr_obj
                    elif isinstance(addr_obj, dict):
                        parts = [
                            addr_obj.get("streetAddress"),
                            addr_obj.get("addressLocality"),
                            addr_obj.get("addressRegion"),
                            addr_obj.get("postalCode")
                        ]
                        address = ", ".join(p for p in parts if p)
        except Exception:
            pass
    
    return {"name": name, "phone": phone, "address": address}


@app.route("/api/enrich", methods=["POST"])
def api_enrich():
    """Scrape website for company info."""
    data = request.get_json(force=True) or {}
    url = (data.get("url") or "").strip()
    
    if not url:
        return jsonify({"error": "url required"}), 400
    
    if not url.startswith(("http://", "https://")):
        url = "https://" + url
    
    try:
        headers = {"User-Agent": "TonyOS/1.0 (Lead Research Tool)"}
        resp = requests.get(url, headers=headers, timeout=10)
        resp.raise_for_status()
        
        soup = BeautifulSoup(resp.text, "html.parser")
        
        from_schema = extract_from_jsonld(soup)
        
        phone = from_schema["phone"]
        if not phone:
            tel_link = soup.find("a", href=re.compile(r'^tel:'))
            if tel_link:
                phone = tel_link.get("href", "").replace("tel:", "").strip()
        phone = clean_phone(phone)
        
        address = from_schema["address"] or pick_best_address(soup)
        
        name = from_schema["name"]
        if not name:
            title_tag = soup.find("title")
            if title_tag:
                title_text = title_tag.get_text().strip()
                name = title_text.split("|")[0].split(" - ")[0].strip()
        
        return jsonify({
            "name": name or "",
            "phone": phone or "",
            "address": address or "",
            "source": "html+schema"
        })
    
    except requests.RequestException as e:
        return jsonify({"error": f"Failed to fetch website: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/storybrand", methods=["POST"])
@rate_limit(20)
@require_api_key
def api_storybrand():
    data = request.get_json(force=True)
    url = data.get("url", "")
    notes = data.get("notes", "")

    if not url:
        return jsonify({"error": "URL is required"}), 400

    website_content = ""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        }
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        
        for script in soup(["script", "style", "nav", "footer", "header"]):
            script.decompose()
        
        title = soup.find("title")
        title_text = title.get_text().strip() if title else ""
        
        meta_desc = soup.find("meta", attrs={"name": "description"})
        meta_text = meta_desc.get("content", "") if meta_desc else ""
        
        headings = []
        for h in soup.find_all(["h1", "h2", "h3"])[:10]:
            headings.append(h.get_text().strip())
        
        paragraphs = []
        for p in soup.find_all("p")[:20]:
            text = p.get_text().strip()
            if len(text) > 30:
                paragraphs.append(text[:500])
        
        website_content = f"""
Title: {title_text}
Meta Description: {meta_text}
Headings: {' | '.join(headings)}
Content:
{chr(10).join(paragraphs[:10])}
"""
        logging.info(f"[storybrand] Fetched {len(website_content)} chars from {url}")
    except requests.RequestException as e:
        logging.warning(f"[storybrand] Could not fetch {url}: {e}")
        website_content = f"Could not fetch website content. URL provided: {url}"

    prompt = f"""Analyze this website and create StoryBrand messaging.

WEBSITE CONTENT:
{website_content[:8000]}

ADDITIONAL NOTES: {notes}

Return a JSON object with these exact keys. Keep each value to 1-2 sentences:
{{
  "character": "Who is the target customer",
  "want": "What do they want to achieve",
  "external_problem": "What visible problem do they face",
  "internal_problem": "How does it make them feel",
  "philosophical": "Why is this problem wrong or unfair",
  "empathy": "How you understand their struggle",
  "authority": "What credentials or proof they have",
  "step1": "First step of the plan",
  "step2": "Second step of the plan",
  "step3": "Third step of the plan",
  "cta_primary": "Primary call to action",
  "cta_secondary": "Transitional call to action",
  "failure": "What happens if they don't act",
  "success": "What do they achieve if they act"
}}

Return ONLY the JSON object, no markdown, no explanation."""

    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a StoryBrand certified messaging expert. Extract messaging directly from the website content provided. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4
        )
        response_text = completion.choices[0].message.content.strip()
        
        start = response_text.find("{")
        end = response_text.rfind("}") + 1
        if start != -1 and end > start:
            json_str = response_text[start:end]
            try:
                parsed = json.loads(json_str)
                return jsonify({
                    "success": True,
                    "data": parsed,
                    "url": url
                })
            except json.JSONDecodeError:
                pass
        
        return jsonify({
            "success": False,
            "story_script": response_text,
            "url": url,
            "error": "Could not parse structured response"
        })
    except Exception as e:
        logging.error(f"[storybrand] OpenAI error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/crm/scan-company", methods=["POST"])
@rate_limit(20)
@require_api_key
def api_crm_scan_company():
    """Scan a company website and extract CRM-relevant intel."""
    data = request.get_json(force=True)
    url = data.get("url", "").strip()

    if not url:
        return jsonify({"error": "URL is required"}), 400

    website_content = ""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        }
        resp = requests.get(url, headers=headers, timeout=15)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.text, "html.parser")
        
        for script in soup(["script", "style", "nav", "footer"]):
            script.decompose()
        
        title = soup.find("title")
        title_text = title.get_text().strip() if title else ""
        
        meta_desc = soup.find("meta", attrs={"name": "description"})
        meta_text = meta_desc.get("content", "") if meta_desc else ""
        
        headings = []
        for h in soup.find_all(["h1", "h2", "h3"])[:15]:
            headings.append(h.get_text().strip())
        
        paragraphs = []
        for p in soup.find_all("p")[:25]:
            text = p.get_text().strip()
            if len(text) > 30:
                paragraphs.append(text[:600])
        
        about_text = ""
        about_section = soup.find(id=re.compile("about|company|who", re.I)) or soup.find(class_=re.compile("about|company", re.I))
        if about_section:
            about_text = about_section.get_text()[:1000]
        
        website_content = f"""
Title: {title_text}
Meta Description: {meta_text}
Headings: {' | '.join(headings)}
About Section: {about_text[:500] if about_text else 'Not found'}
Content:
{chr(10).join(paragraphs[:15])}
"""
        logging.info(f"[crm-scan] Fetched {len(website_content)} chars from {url}")
    except requests.RequestException as e:
        logging.warning(f"[crm-scan] Could not fetch {url}: {e}")
        website_content = f"Could not fetch website content. URL provided: {url}"

    prompt = f"""Analyze this company website for CRM/sales intelligence.

WEBSITE CONTENT:
{website_content[:8000]}

Return a JSON object with these exact keys. Be specific and actionable for sales outreach:
{{
  "company_name": "The company name",
  "industry": "Their industry or sector",
  "company_description": "What they do in 1-2 sentences",
  "products_services": "Main products or services offered",
  "target_market": "Who they sell to",
  "pain_points": "3 likely pain points or challenges they face",
  "value_props": "Their key value propositions or differentiators",
  "recent_news": "Any recent announcements, funding, or growth signals",
  "key_people": "Any executives or key contacts mentioned",
  "outreach_angle": "Best angle for initial outreach - what problem can you solve for them",
  "talking_points": "3 conversation starters based on their website"
}}

Return ONLY the JSON object, no markdown, no explanation."""

    natoli_data = None
    apollo_data = None
    
    # Extract domain from URL for Apollo enrichment
    try:
        from urllib.parse import urlparse
        parsed = urlparse(url)
        domain = parsed.netloc.replace("www.", "")
    except:
        domain = None
    
    # Call Natoli Scanner for revops decision
    try:
        scanner_resp = requests.post(
            "http://localhost:3000/scan",
            json={"url": url},
            timeout=10
        )
        if scanner_resp.status_code == 200:
            natoli_data = scanner_resp.json()
            decision = natoli_data.get("revops_decision", {})
            logging.info(f"[crm-scan] Natoli gate: {decision.get('market_gate')}, tier: {decision.get('environment_tier')}")
    except Exception as e:
        logging.warning(f"[crm-scan] Natoli scanner unavailable: {e}")
    
    # Call Apollo for enrichment (division fit, pain points, why Natoli)
    if domain:
        try:
            apollo_resp = requests.post(
                "http://localhost:3001/analyze",
                json={"domain": domain},
                timeout=30
            )
            if apollo_resp.status_code == 200:
                apollo_data = apollo_resp.json()
                logging.info(f"[crm-scan] Apollo: {apollo_data.get('company')} - {apollo_data.get('division_fit')}")
        except Exception as e:
            logging.warning(f"[crm-scan] Apollo enrichment unavailable: {e}")

    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a sales intelligence analyst. Extract actionable company intel for CRM and outreach purposes. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4
        )
        response_text = completion.choices[0].message.content.strip()
        
        start = response_text.find("{")
        end = response_text.rfind("}") + 1
        if start != -1 and end > start:
            json_str = response_text[start:end]
            try:
                parsed = json.loads(json_str)
                result = {
                    "success": True,
                    "data": parsed,
                    "url": url
                }
                if natoli_data and natoli_data.get("revops_decision"):
                    result["natoli"] = natoli_data.get("revops_decision")
                if apollo_data:
                    result["apollo"] = {
                        "company": apollo_data.get("company"),
                        "division_fit": apollo_data.get("division_fit"),
                        "is_primary_market": apollo_data.get("is_primary_market"),
                        "is_qualified": apollo_data.get("is_qualified"),
                        "tier": apollo_data.get("tier"),
                        "tier_reason": apollo_data.get("tier_reason"),
                        "pain_points": apollo_data.get("pain_points", []),
                        "why_natoli": apollo_data.get("why_natoli"),
                        "detected_industry": apollo_data.get("detected_industry"),
                        "employees": apollo_data.get("employees"),
                        "revenue": apollo_data.get("revenue"),
                        "account_id": apollo_data.get("account_id")
                    }
                return jsonify(result)
            except json.JSONDecodeError:
                pass
        
        return jsonify({
            "success": False,
            "raw_analysis": response_text,
            "url": url,
            "error": "Could not parse structured response"
        })
    except Exception as e:
        logging.error(f"[crm-scan] OpenAI error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/crm/apollo-enrich", methods=["POST"])
@rate_limit(20)
@require_api_key
def api_crm_apollo_enrich():
    """Enrich company data using Apollo.io API."""
    data = request.get_json(force=True)
    domain = data.get("domain", "").strip()

    if not domain:
        return jsonify({"error": "Domain is required"}), 400

    try:
        apollo_resp = requests.post(
            "http://localhost:3001/api/company/enrich",
            json={"domain": domain},
            timeout=30
        )
        
        if apollo_resp.status_code == 404:
            return jsonify({"error": "Company not found in Apollo", "domain": domain}), 404
        
        if not apollo_resp.ok:
            return jsonify({"error": f"Apollo server error: {apollo_resp.status_code}"}), 500
        
        apollo_data = apollo_resp.json()
        logging.info(f"[apollo-enrich] Enriched {domain}: {apollo_data.get('name', 'Unknown')}")
        
        return jsonify({
            "success": True,
            "data": apollo_data,
            "domain": domain
        })
    except requests.exceptions.ConnectionError:
        return jsonify({"error": "Apollo server not available. Make sure the Apollo Server workflow is running."}), 503
    except Exception as e:
        logging.error(f"[apollo-enrich] Error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/crm/apollo-quick", methods=["POST"])
@rate_limit(30)
@require_api_key
def api_crm_apollo_quick():
    """Quick check if company is a Natoli fit using Apollo."""
    data = request.get_json(force=True)
    domain = data.get("domain", "").strip()

    if not domain:
        return jsonify({"error": "Domain is required"}), 400

    try:
        apollo_resp = requests.post(
            "http://localhost:3001/api/company/quick-check",
            json={"domain": domain},
            timeout=15
        )
        
        if not apollo_resp.ok:
            return jsonify({"error": "Quick check failed"}), 500
        
        return jsonify(apollo_resp.json())
    except Exception as e:
        logging.error(f"[apollo-quick] Error: {e}")
        return jsonify({"error": str(e)}), 500


# ===== CRM GOVERNOR ROUTES (Judgment Layer) =====

@app.route("/api/crm/governor/assumptions", methods=["GET", "POST"])
@rate_limit(60)
@require_api_key
def api_crm_assumptions():
    """Log and retrieve assumptions about accounts."""
    from crm_governor.assumptions import log_assumption, get_assumptions, get_account_assumptions
    
    if request.method == "POST":
        data = request.get_json(force=True)
        entry = log_assumption(
            account_id=data.get("account_id"),
            account_name=data.get("account_name", "Unknown"),
            assumption=data.get("assumption", ""),
            confidence=data.get("confidence", "medium"),
            context=data.get("context")
        )
        return jsonify({"success": True, "entry": entry})
    
    account_id = request.args.get("account_id")
    if account_id:
        return jsonify({"assumptions": get_account_assumptions(account_id)})
    return jsonify({"assumptions": get_assumptions(limit=50)})


@app.route("/api/crm/governor/outcomes", methods=["GET", "POST"])
@rate_limit(60)
@require_api_key
def api_crm_outcomes():
    """Score and retrieve outcomes for account actions."""
    from crm_governor.outcomes import score_outcome, get_outcomes, get_account_outcomes, get_outcome_summary
    
    if request.method == "POST":
        data = request.get_json(force=True)
        entry = score_outcome(
            account_id=data.get("account_id"),
            account_name=data.get("account_name", "Unknown"),
            action=data.get("action", ""),
            outcome=data.get("outcome", "neutral"),
            notes=data.get("notes", ""),
            learnings=data.get("learnings")
        )
        return jsonify({"success": True, "entry": entry})
    
    account_id = request.args.get("account_id")
    if account_id:
        return jsonify({"outcomes": get_account_outcomes(account_id)})
    
    summary = request.args.get("summary")
    if summary:
        return jsonify(get_outcome_summary())
    
    return jsonify({"outcomes": get_outcomes(limit=50)})


@app.route("/api/crm/governor/health", methods=["GET"])
@rate_limit(60)
@require_api_key
def api_crm_health():
    """Get pipeline health analysis."""
    from crm_governor.outcomes import get_outcomes
    from crm_governor.signal_summary import get_pipeline_health
    
    days = int(request.args.get("days", 30))
    outcomes = get_outcomes(limit=200)
    health = get_pipeline_health(outcomes, days=days)
    return jsonify(health)


@app.route("/api/crm/governor/stale", methods=["GET"])
@rate_limit(60)
@require_api_key
def api_crm_stale():
    """Get stale assumptions that need re-evaluation."""
    from crm_governor.assumptions import get_assumptions
    from crm_governor.memory_decay import flag_stale_assumptions
    
    days = int(request.args.get("days", 60))
    assumptions = get_assumptions(limit=200)
    stale = flag_stale_assumptions(assumptions, days=days)
    return jsonify({"stale_assumptions": stale, "count": len(stale)})


@app.route("/api/crm/governor/pre-action", methods=["POST"])
@rate_limit(60)
@require_api_key
def api_crm_pre_action():
    """Pre-action gate - checks momentum and identity before action."""
    from crm_governor.cognitive import pre_action_gate, get_second_order_prompt
    
    data = request.get_json(force=True)
    action = data.get("action", "")
    aligns_with_identity = data.get("aligns_with_identity", True)
    
    result = pre_action_gate(action, aligns_with_identity)
    if result.get('allowed'):
        result['second_order_prompt'] = get_second_order_prompt(action)
    
    return jsonify(result)


@app.route("/api/crm/governor/session", methods=["GET", "POST"])
@rate_limit(60)
@require_api_key
def api_crm_session():
    """Close a work session or get recent sessions."""
    from crm_governor.cognitive import close_session, get_recent_sessions
    
    if request.method == "POST":
        data = request.get_json(force=True)
        session = close_session(
            summary=data.get("summary", ""),
            wins=data.get("wins", []),
            learnings=data.get("learnings", []),
            next_focus=data.get("next_focus")
        )
        return jsonify({"success": True, "session": session})
    
    return jsonify({"sessions": get_recent_sessions(limit=10)})


@app.route("/api/crm/governor/confidence", methods=["POST"])
@rate_limit(60)
@require_api_key
def api_crm_confidence():
    """Tag a statement with confidence level."""
    from crm_governor.cognitive import tag_confidence, get_confidence_warning
    
    data = request.get_json(force=True)
    tagged = tag_confidence(
        statement=data.get("statement", ""),
        level=data.get("level", "informed_guess"),
        evidence=data.get("evidence")
    )
    tagged['warning'] = get_confidence_warning(data.get("level", "informed_guess"))
    return jsonify(tagged)


@app.route("/api/job_analysis", methods=["POST"])
def api_job_analysis():
    data = request.get_json(force=True)
    desc = data.get("description", "")
    resume = data.get("resume", "")
    proof = data.get("proof", "")

    prompt = f"""Act as an elite job analysis agent for a senior operator.

INPUTS:
Job Description:
{desc}

Resume Summary:
{resume}

Proof Log (wins, metrics, systems):
{proof}

TASKS:
1. Analyze role fit honestly. Score 1-10 and explain why.
2. Identify where I am overqualified, underqualified, or misaligned.
3. Identify the top 3 problems this role likely has.
4. Show how my background uniquely solves those problems.
5. Generate interview talking points with proof.
6. Identify red flags or reasons to avoid this role.
7. Recommend: GO / PROCEED CAUTIOUSLY / NO-GO

RULES:
- Be blunt, not polite
- Optimize for leverage, not volume
- Do not rewrite resume
- Focus on decision-making and positioning"""

    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a blunt career strategist who helps operators make smart job decisions."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4
        )
        return jsonify({"analysis": completion.choices[0].message.content.strip()})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/web-answer", methods=["POST"])
@rate_limit(20)
@require_api_key
def api_web_answer():
    """Trust-first web answer endpoint - returns verified info with sources."""
    data = request.get_json(force=True)
    question = data.get("question", "").strip()

    if not question:
        return jsonify({"error": "question required"}), 400

    system_prompt = """You answer ONLY with verifiable info.
If you do not have sources, say "I cannot verify that."
Return your response in this format:
1) A short, direct answer
2) List any sources if available
3) State your confidence level (low/medium/high)

Be honest about uncertainty. No em dashes."""

    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": question}
            ],
            temperature=0.3
        )
        answer = completion.choices[0].message.content.strip()
        return jsonify({
            "answer": answer,
            "sources": [],
            "confidence": "medium"
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/chat", methods=["POST"])
@rate_limit(30)
@require_api_key
def api_chat():
    data = request.get_json(force=True)
    user_message = data.get("message", "").strip()
    mode = data.get("mode", "truth")
    mood = data.get("mood", "neutral")
    truth = data.get("truth", False)
    persona = data.get("persona", "default")
    provider = data.get("provider", "openai")
    model = data.get("model")
    
    source_app = data.get("source_app", "AI Chat")
    intent = data.get("intent", "chat")
    context = data.get("context", {})
    research_urls = data.get("research_urls")

    if not user_message:
        return jsonify({"error": "Empty message"}), 400

    decision = enforce_policy(source_app, mode, intent)
    
    if not decision.ok:
        return jsonify({
            "reply": decision.reason,
            "citations": [],
            "confidence": 1.0,
            "policy": {
                "app": decision.app,
                "mode": decision.mode,
                "capabilities": decision.capabilities,
                "allowed_apps_for_memory": decision.allowed_apps_for_memory
            }
        })

    print(f"[policy] App: {decision.app}, Mode: {mode}, Intent: {intent}, Provider: {provider}, Model: {model}")

    chat_history.append({"role": "user", "content": user_message})

    try:
        result = call_multi_provider(user_message, provider=provider, model=model, mode=mode, mood=mood, truth=truth, persona=persona)
        if result.get("error"):
            assistant_reply = f"Error: {result['error']}"
        else:
            assistant_reply = result.get("text", "No response received.")
    except Exception as e:
        print(f"[chat] Error calling {provider}: {e}")
        assistant_reply = f"I hit an error talking to {provider}. Error: {str(e)}"

    chat_history.append({"role": "assistant", "content": assistant_reply})
    save_chat_history(chat_history)
    log_interaction(user_message, assistant_reply)
    
    promo = should_create_candidate(user_message)
    memory_info = {
        "candidate_created": promo.should_candidate,
        "type": promo.memory_type,
        "confidence": promo.confidence,
        "reason": promo.reason
    }
    
    if promo.should_candidate:
        try:
            cid = candidate_id(data.get("session_id", "default"), promo.cleaned_text)
            with get_db_context() as conn:
                if conn:
                    cur = conn.cursor()
                    cur.execute("""
                        INSERT INTO memory_candidates (id, created_at, session_id, source_app, memory_type, confidence, text, reason, status)
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                        ON CONFLICT (id) DO NOTHING
                    """, (cid, int(time.time()), data.get("session_id", "default"), decision.app,
                          promo.memory_type, promo.confidence, promo.cleaned_text, promo.reason, "candidate"))
                    conn.commit()
                    memory_info["candidate_id"] = cid
        except Exception as e:
            logging.error(f"[memory_candidate] Error: {e}")
    
    try:
        memory_result = extract_memory_if_worth_saving(user_message)
        if memory_result.get("should_save"):
            append_to_long_memory(memory_result["memory"])
            save_smart_memory("tony", memory_result["memory"], memory_result.get("importance", 3))
    except Exception:
        pass

    return jsonify({
        "reply": assistant_reply,
        "memory": memory_info,
        "policy": {
            "app": decision.app,
            "mode": decision.mode,
            "capabilities": decision.capabilities,
            "allowed_apps_for_memory": decision.allowed_apps_for_memory
        }
    })


@app.route("/api/memory/confirm", methods=["POST"])
def api_memory_confirm():
    data = request.get_json(force=True)
    cid = data.get("candidate_id")
    session_id = data.get("session_id", "default")
    
    if not cid:
        return jsonify({"ok": False, "error": "Missing candidate_id"}), 400
    
    try:
        with get_db_context() as conn:
            if not conn:
                return jsonify({"ok": False, "error": "Database unavailable"}), 500
            cur = conn.cursor()
            cur.execute("""
                SELECT id, session_id, source_app, memory_type, confidence, text
                FROM memory_candidates
                WHERE id=%s AND session_id=%s AND status='candidate'
            """, (cid, session_id))
            row = cur.fetchone()
            
            if not row:
                return jsonify({"ok": False, "error": "Candidate not found"})
            
            cur.execute("UPDATE memory_candidates SET status='confirmed' WHERE id=%s", (cid,))
            append_to_long_memory(row['text'])
            save_smart_memory("tony", row['text'], int(row['confidence'] * 5))
            conn.commit()
            return jsonify({"ok": True, "confirmed": row['text']})
    except Exception as e:
        logging.error(f"[memory_confirm] Error: {e}")
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/memory", methods=["GET"])
def api_memory_get():
    memory = load_long_memory(max_chars=10000)
    return jsonify({"memory": memory})


@app.route("/api/memory", methods=["POST"])
def api_memory_post():
    data = request.get_json(force=True)
    note = data.get("note", "").strip()
    if note:
        append_to_long_memory(note)
    return jsonify({"ok": True})


@app.route("/smart_research", methods=["POST"])
def smart_research():
    data = request.get_json(force=True)
    query = data.get("query", "").strip()
    urls = data.get("urls", [])

    if not query:
        return jsonify({"error": "No query provided"}), 400

    prompt = f"Research topic: {query}\n"
    if urls:
        prompt += f"Source URLs to consider: {', '.join(urls)}\n"
    prompt += "\nProvide a structured briefing with key findings, risks/gaps, and recommended actions."

    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a research assistant. Provide structured briefings."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3
        )
        summary = completion.choices[0].message.content.strip()
        return jsonify({"summary": summary, "sources": urls})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/offer_builder", methods=["POST"])
@rate_limit(20)
@require_api_key
def api_offer_builder():
    data = request.get_json(force=True)
    audience = data.get("audience", "")
    problem = data.get("problem", "")
    outcome = data.get("outcome", "")
    proof = data.get("proof", "")

    prompt = f"""Build an offer based on:
Audience: {audience}
Problem: {problem}
Desired outcome: {outcome}
Proof/assets: {proof}

Return: offer name, pitch, bullet points, and who it is for."""

    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a sales and offer strategist."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4
        )
        return jsonify({"text": completion.choices[0].message.content.strip()})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/outreach_builder", methods=["POST"])
@rate_limit(20)
@require_api_key
def api_outreach_builder():
    data = request.get_json(force=True)
    persona = data.get("persona", "")
    context = data.get("context", "")
    angle = data.get("angle", "")
    channel = data.get("channel", "email + LinkedIn")

    prompt = f"""Create outreach for:
Persona: {persona}
Context: {context}
Angle: {angle}
Channels: {channel}

Return: short email, LinkedIn note, and follow-up line."""

    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an outreach copywriter."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4
        )
        return jsonify({"text": completion.choices[0].message.content.strip()})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/content_ideas", methods=["POST"])
@rate_limit(20)
@require_api_key
def api_content_ideas():
    data = request.get_json(force=True)
    niche = data.get("niche", "")
    platform = data.get("platform", "LinkedIn")
    themes = data.get("themes", "")

    prompt = f"""Generate content ideas for:
Niche: {niche}
Platform: {platform}
Themes: {themes}

Return: 5-7 content ideas that build demand and credibility."""

    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a content strategist."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.5
        )
        return jsonify({"text": completion.choices[0].message.content.strip()})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/company_intel", methods=["POST"])
@rate_limit(20)
@require_api_key
def api_company_intel():
    data = request.get_json(force=True)
    url = data.get("url", "")
    goal = data.get("goal", "")

    prompt = f"""Analyze this company:
URL: {url}
Goal: {goal}

Return: company overview, positioning angles, and talking points."""

    try:
        completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a business analyst."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.4
        )
        return jsonify({"text": completion.choices[0].message.content.strip()})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============ CRM Database API ============

@app.route("/api/db/connections", methods=["GET"])
def api_db_connections_list():
    """Get all CRM connections from database."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM crm_connections ORDER BY created_at DESC")
        rows = cur.fetchall()
        conn.close()
        return jsonify({"connections": [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/connections", methods=["POST"])
def api_db_connections_create():
    """Create a new CRM connection."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO crm_connections 
            (first_name, last_name, full_name, title, company, email, phone, industry, linkedin_url, stage, status, next_step, notes, sent, sent_date, customer_tier)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
        """, (
            data.get("firstName", ""),
            data.get("lastName", ""),
            data.get("fullName", ""),
            data.get("title", ""),
            data.get("company", ""),
            data.get("email", ""),
            data.get("phone", ""),
            data.get("industry", ""),
            data.get("linkedinUrl", ""),
            data.get("stage", "New"),
            data.get("status", "New"),
            data.get("nextStep", ""),
            data.get("notes", ""),
            data.get("sent", False),
            data.get("sentDate"),
            data.get("customerTier")
        ))
        row = cur.fetchone()
        conn.commit()
        conn.close()
        return jsonify({"ok": True, "id": row["id"]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/connections/<int:id>", methods=["PUT"])
def api_db_connections_update(id):
    """Update a CRM connection."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            UPDATE crm_connections SET
                first_name=%s, last_name=%s, full_name=%s, title=%s, company=%s,
                email=%s, phone=%s, industry=%s, linkedin_url=%s, stage=%s, status=%s,
                next_step=%s, notes=%s, sent=%s, sent_date=%s, customer_tier=%s, updated_at=CURRENT_TIMESTAMP
            WHERE id=%s
        """, (
            data.get("firstName", ""),
            data.get("lastName", ""),
            data.get("fullName", ""),
            data.get("title", ""),
            data.get("company", ""),
            data.get("email", ""),
            data.get("phone", ""),
            data.get("industry", ""),
            data.get("linkedinUrl", ""),
            data.get("stage", "New"),
            data.get("status", "New"),
            data.get("nextStep", ""),
            data.get("notes", ""),
            data.get("sent", False),
            data.get("sentDate"),
            data.get("customerTier"),
            id
        ))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/connections/<int:id>", methods=["DELETE"])
def api_db_connections_delete(id):
    """Delete a CRM connection."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM crm_connections WHERE id=%s", (id,))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/activities", methods=["GET"])
def api_db_activities_list():
    """Get all CRM activities."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM crm_activities ORDER BY activity_date DESC")
        rows = cur.fetchall()
        conn.close()
        return jsonify({"activities": [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/activities", methods=["POST"])
def api_db_activities_create():
    """Create a CRM activity."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO crm_activities 
            (connection_id, contact_name, company, channel, activity_type, outcome, next_step, next_step_due, notes)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
        """, (
            data.get("connectionId"),
            data.get("contactName", ""),
            data.get("company", ""),
            data.get("channel", ""),
            data.get("type", ""),
            data.get("outcome", ""),
            data.get("nextStep", ""),
            data.get("nextStepDue"),
            data.get("notes", "")
        ))
        row = cur.fetchone()
        conn.commit()
        conn.close()
        return jsonify({"ok": True, "id": row["id"]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/activities/<int:id>", methods=["DELETE"])
def api_db_activities_delete(id):
    """Delete a CRM activity."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM crm_activities WHERE id=%s", (id,))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/deals", methods=["GET"])
def api_db_deals_list():
    """Get all CRM deals."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM crm_deals ORDER BY created_at DESC")
        rows = cur.fetchall()
        conn.close()
        return jsonify({"deals": [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/deals", methods=["POST"])
def api_db_deals_create():
    """Create a CRM deal."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO crm_deals 
            (connection_id, name, company, stage, est_value, probability, next_step, next_step_due, notes)
            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
            RETURNING id
        """, (
            data.get("connectionId"),
            data.get("name", ""),
            data.get("company", ""),
            data.get("stage", "Awareness"),
            data.get("estValue", 0),
            data.get("probability", 50),
            data.get("nextStep", ""),
            data.get("nextStepDue"),
            data.get("notes", "")
        ))
        row = cur.fetchone()
        conn.commit()
        conn.close()
        return jsonify({"ok": True, "id": row["id"]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/deals/<int:id>", methods=["PUT"])
def api_db_deals_update(id):
    """Update a CRM deal."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            UPDATE crm_deals SET stage=%s, last_activity_date=CURRENT_TIMESTAMP WHERE id=%s
        """, (data.get("stage"), id))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/deals/<int:id>", methods=["DELETE"])
def api_db_deals_delete(id):
    """Delete a CRM deal."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM crm_deals WHERE id=%s", (id,))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ============ CRM Targets Database API ============

@app.route("/api/db/targets", methods=["GET"])
def api_db_targets_get():
    """Get CRM targets."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("SELECT target_type, target_value FROM crm_targets")
        rows = cur.fetchall()
        conn.close()
        targets = {}
        for row in rows:
            targets[row["target_type"]] = row["target_value"]
        return jsonify({
            "targets": {
                "outreach_target": targets.get("outreach", 150),
                "connections_target": targets.get("connections", 50)
            }
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/targets", methods=["POST"])
def api_db_targets_save():
    """Save CRM targets."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        for target_type in ["outreach", "connections"]:
            value = data.get(f"{target_type}_target", data.get(target_type, 150 if target_type == "outreach" else 50))
            cur.execute("SELECT id FROM crm_targets WHERE target_type=%s", (target_type,))
            existing = cur.fetchone()
            if existing:
                cur.execute("UPDATE crm_targets SET target_value=%s, updated_at=CURRENT_TIMESTAMP WHERE target_type=%s", (value, target_type))
            else:
                cur.execute("INSERT INTO crm_targets (target_type, target_value) VALUES (%s, %s)", (target_type, value))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# ============ Budget Database API ============

@app.route("/api/db/budget/settings", methods=["GET"])
def api_db_budget_settings_get():
    """Get budget settings."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM budget_settings ORDER BY id LIMIT 1")
        row = cur.fetchone()
        cur.execute("SELECT * FROM budget_savings_accounts ORDER BY id")
        savings = cur.fetchall()
        conn.close()
        return jsonify({
            "settings": dict(row) if row else {"weekly_budget": None, "checking_balance": None},
            "savingsAccounts": [dict(s) for s in savings]
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/budget/settings", methods=["POST"])
def api_db_budget_settings_save():
    """Save budget settings."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("SELECT id FROM budget_settings LIMIT 1")
        existing = cur.fetchone()
        if existing:
            cur.execute("""
                UPDATE budget_settings SET weekly_budget=%s, checking_balance=%s, updated_at=CURRENT_TIMESTAMP WHERE id=%s
            """, (data.get("weeklyBudget"), data.get("checkingBalance"), existing["id"]))
        else:
            cur.execute("""
                INSERT INTO budget_settings (weekly_budget, checking_balance) VALUES (%s, %s)
            """, (data.get("weeklyBudget"), data.get("checkingBalance")))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/budget/savings", methods=["GET"])
def api_db_budget_savings_list():
    """Get all savings accounts."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM budget_savings_accounts ORDER BY id")
        rows = cur.fetchall()
        conn.close()
        return jsonify({"accounts": [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/budget/savings", methods=["POST"])
def api_db_budget_savings_create():
    """Create a savings account."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO budget_savings_accounts (name, balance) VALUES (%s, %s) RETURNING id
        """, (data.get("name", ""), data.get("balance")))
        row = cur.fetchone()
        conn.commit()
        conn.close()
        return jsonify({"ok": True, "id": row["id"]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/budget/savings/<int:id>", methods=["PUT"])
def api_db_budget_savings_update(id):
    """Update a savings account."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            UPDATE budget_savings_accounts SET name=%s, balance=%s, updated_at=CURRENT_TIMESTAMP WHERE id=%s
        """, (data.get("name"), data.get("balance"), id))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/budget/savings/<int:id>", methods=["DELETE"])
def api_db_budget_savings_delete(id):
    """Delete a savings account."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM budget_savings_accounts WHERE id=%s", (id,))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/budget/transactions", methods=["GET"])
def api_db_budget_transactions_list():
    """Get all budget transactions."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM budget_transactions ORDER BY transaction_date DESC LIMIT 100")
        rows = cur.fetchall()
        conn.close()
        return jsonify({"transactions": [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/budget/transactions", methods=["POST"])
def api_db_budget_transactions_create():
    """Create a budget transaction."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO budget_transactions (transaction_type, amount, category, merchant, note)
            VALUES (%s, %s, %s, %s, %s) RETURNING id
        """, (
            data.get("type", "expense"),
            data.get("amount", 0),
            data.get("category", ""),
            data.get("merchant", ""),
            data.get("note", "")
        ))
        row = cur.fetchone()
        conn.commit()
        conn.close()
        return jsonify({"ok": True, "id": row["id"]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/budget/transactions/<int:id>", methods=["DELETE"])
def api_db_budget_transactions_delete(id):
    """Delete a budget transaction."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM budget_transactions WHERE id=%s", (id,))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============ Timer Sessions ============
@app.route("/api/db/timer-sessions", methods=["GET"])
def api_db_timer_sessions_list():
    """List all timer sessions."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM crm_timer_sessions ORDER BY start_time DESC LIMIT 50")
        rows = cur.fetchall()
        conn.close()
        return jsonify({"sessions": [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/timer-sessions", methods=["POST"])
def api_db_timer_sessions_create():
    """Create a new timer session."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO crm_timer_sessions (duration_seconds, focus_type, notes, completed)
            VALUES (%s, %s, %s, %s) RETURNING id, start_time
        """, (
            data.get("duration_seconds", 600),
            data.get("focus_type", "Outreach Sprint"),
            data.get("notes", ""),
            data.get("completed", True)
        ))
        row = cur.fetchone()
        conn.commit()
        conn.close()
        return jsonify({"ok": True, "id": row["id"], "start_time": row["start_time"].isoformat() if row["start_time"] else None})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/timer-sessions/<int:id>", methods=["DELETE"])
def api_db_timer_sessions_delete(id):
    """Delete a timer session."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM crm_timer_sessions WHERE id=%s", (id,))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/timer-sessions/stats", methods=["GET"])
def api_db_timer_sessions_stats():
    """Get timer session statistics."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT 
                COUNT(*) as total_sessions,
                COALESCE(SUM(duration_seconds), 0) as total_seconds,
                COUNT(CASE WHEN DATE(start_time) = CURRENT_DATE THEN 1 END) as today_sessions,
                COALESCE(SUM(CASE WHEN DATE(start_time) = CURRENT_DATE THEN duration_seconds ELSE 0 END), 0) as today_seconds,
                COUNT(CASE WHEN start_time >= CURRENT_DATE - INTERVAL '7 days' THEN 1 END) as week_sessions,
                COALESCE(SUM(CASE WHEN start_time >= CURRENT_DATE - INTERVAL '7 days' THEN duration_seconds ELSE 0 END), 0) as week_seconds
            FROM crm_timer_sessions WHERE completed = TRUE
        """)
        row = cur.fetchone()
        conn.close()
        return jsonify(dict(row) if row else {})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============ LinkedIn Activities ============
@app.route("/api/db/linkedin-activities", methods=["GET"])
def api_db_linkedin_activities_list():
    """List all LinkedIn activities."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM linkedin_activities ORDER BY imported_at DESC LIMIT 100")
        rows = cur.fetchall()
        conn.close()
        return jsonify({"activities": [dict(r) for r in rows]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/linkedin-activities", methods=["POST"])
def api_db_linkedin_activities_create():
    """Create a LinkedIn activity."""
    data = request.get_json(force=True)
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            INSERT INTO linkedin_activities (activity_type, person_name, company, title, content, activity_date, linkedin_url)
            VALUES (%s, %s, %s, %s, %s, %s, %s) RETURNING id
        """, (
            data.get("activity_type", "connection"),
            data.get("person_name", ""),
            data.get("company", ""),
            data.get("title", ""),
            data.get("content", ""),
            data.get("activity_date"),
            data.get("linkedin_url", "")
        ))
        row = cur.fetchone()
        conn.commit()
        conn.close()
        return jsonify({"ok": True, "id": row["id"]})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/linkedin-activities/bulk", methods=["POST"])
def api_db_linkedin_activities_bulk():
    """Bulk import LinkedIn activities."""
    data = request.get_json(force=True)
    activities = data.get("activities", [])
    if not activities:
        return jsonify({"error": "No activities provided"}), 400
    
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        imported = 0
        for act in activities:
            cur.execute("""
                INSERT INTO linkedin_activities (activity_type, person_name, company, title, content, activity_date, linkedin_url)
                VALUES (%s, %s, %s, %s, %s, %s, %s)
            """, (
                act.get("activity_type", "connection"),
                act.get("person_name", ""),
                act.get("company", ""),
                act.get("title", ""),
                act.get("content", ""),
                act.get("activity_date"),
                act.get("linkedin_url", "")
            ))
            imported += 1
        conn.commit()
        conn.close()
        return jsonify({"ok": True, "imported": imported})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/linkedin-activities/<int:id>", methods=["DELETE"])
def api_db_linkedin_activities_delete(id):
    """Delete a LinkedIn activity."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM linkedin_activities WHERE id=%s", (id,))
        conn.commit()
        conn.close()
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/db/linkedin-activities/stats", methods=["GET"])
def api_db_linkedin_activities_stats():
    """Get LinkedIn activity statistics."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        cur.execute("""
            SELECT 
                COUNT(*) as total,
                COUNT(CASE WHEN activity_type = 'connection' THEN 1 END) as connections,
                COUNT(CASE WHEN activity_type = 'message' THEN 1 END) as messages,
                COUNT(CASE WHEN activity_type = 'post' THEN 1 END) as posts,
                COUNT(CASE WHEN activity_type = 'comment' THEN 1 END) as comments,
                COUNT(CASE WHEN imported_at >= CURRENT_DATE - INTERVAL '7 days' THEN 1 END) as week_total
            FROM linkedin_activities
        """)
        row = cur.fetchone()
        conn.close()
        return jsonify(dict(row) if row else {})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============ Analytics API ============
@app.route("/api/db/analytics", methods=["GET"])
def api_db_analytics():
    """Get CRM analytics data for charts."""
    conn = get_db()
    if not conn:
        return jsonify({"error": "Database not available"}), 500
    try:
        cur = conn.cursor()
        
        cur.execute("""
            SELECT DATE(created_at) as date, COUNT(*) as count
            FROM crm_activities
            WHERE created_at >= CURRENT_DATE - INTERVAL '30 days'
            GROUP BY DATE(created_at)
            ORDER BY date
        """)
        daily_activities = [{"date": str(r["date"]), "count": r["count"]} for r in cur.fetchall()]
        
        cur.execute("""
            SELECT outcome, COUNT(*) as count
            FROM crm_activities
            WHERE created_at >= CURRENT_DATE - INTERVAL '30 days'
            GROUP BY outcome
        """)
        outcomes = {r["outcome"]: r["count"] for r in cur.fetchall()}
        
        cur.execute("""
            SELECT channel, COUNT(*) as count
            FROM crm_activities
            WHERE created_at >= CURRENT_DATE - INTERVAL '30 days'
            GROUP BY channel
        """)
        channels = {r["channel"]: r["count"] for r in cur.fetchall()}
        
        cur.execute("""
            SELECT stage, COUNT(*) as count
            FROM crm_connections
            GROUP BY stage
        """)
        stages = {r["stage"]: r["count"] for r in cur.fetchall()}
        
        cur.execute("""
            SELECT 
                COUNT(CASE WHEN created_at >= CURRENT_DATE - INTERVAL '7 days' THEN 1 END) as week_activities,
                COUNT(CASE WHEN created_at >= CURRENT_DATE - INTERVAL '30 days' THEN 1 END) as month_activities,
                COUNT(CASE WHEN outcome = 'Reply' AND created_at >= CURRENT_DATE - INTERVAL '30 days' THEN 1 END) as month_replies,
                COUNT(CASE WHEN outcome = 'Meeting' AND created_at >= CURRENT_DATE - INTERVAL '30 days' THEN 1 END) as month_meetings
            FROM crm_activities
        """)
        summary = dict(cur.fetchone())
        
        cur.execute("SELECT COUNT(*) as total FROM crm_connections")
        summary["total_connections"] = cur.fetchone()["total"]
        
        cur.execute("SELECT COUNT(*) as active FROM crm_deals WHERE stage NOT IN ('Won', 'Lost')")
        summary["active_deals"] = cur.fetchone()["active"]
        
        cur.execute("SELECT COUNT(*) as won FROM crm_deals WHERE stage = 'Won'")
        summary["won_deals"] = cur.fetchone()["won"]
        
        conn.close()
        return jsonify({
            "daily_activities": daily_activities,
            "outcomes": outcomes,
            "channels": channels,
            "stages": stages,
            "summary": summary
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ============ Google Sheets Integration ============
from googleapiclient.discovery import build
from google.oauth2.credentials import Credentials

def get_sheets_credentials():
    """Get Google Sheets access token from Replit connector."""
    hostname = os.getenv("REPLIT_CONNECTORS_HOSTNAME")
    repl_identity = os.getenv("REPL_IDENTITY")
    web_repl_renewal = os.getenv("WEB_REPL_RENEWAL")
    
    if repl_identity:
        x_replit_token = f"repl {repl_identity}"
    elif web_repl_renewal:
        x_replit_token = f"depl {web_repl_renewal}"
    else:
        return None
    
    if not hostname:
        return None
    
    try:
        resp = requests.get(
            f"https://{hostname}/api/v2/connection?include_secrets=true&connector_names=google-sheet",
            headers={
                "Accept": "application/json",
                "X_REPLIT_TOKEN": x_replit_token
            },
            timeout=10
        )
        data = resp.json()
        items = data.get("items", [])
        if not items:
            return None
        
        settings = items[0].get("settings", {})
        access_token = settings.get("access_token") or settings.get("oauth", {}).get("credentials", {}).get("access_token")
        
        if not access_token:
            return None
        
        return Credentials(token=access_token)
    except Exception as e:
        print(f"[sheets] Error getting credentials: {e}")
        return None

def get_sheets_service():
    """Get Google Sheets API service."""
    creds = get_sheets_credentials()
    if not creds:
        return None
    return build("sheets", "v4", credentials=creds, cache_discovery=False)


@app.route("/api/sheets/status", methods=["GET"])
def api_sheets_status():
    """Check if Google Sheets is connected."""
    service = get_sheets_service()
    if service:
        return jsonify({"connected": True})
    return jsonify({"connected": False})


@app.route("/api/sheets/list", methods=["GET"])
def api_sheets_list():
    """List rows from a Google Sheet."""
    sheet_id = request.args.get("sheet_id", "")
    range_name = request.args.get("range", "Sheet1!A:G")
    
    if not sheet_id:
        return jsonify({"error": "sheet_id required"}), 400
    
    service = get_sheets_service()
    if not service:
        return jsonify({"error": "Google Sheets not connected"}), 401
    
    try:
        result = service.spreadsheets().values().get(
            spreadsheetId=sheet_id,
            range=range_name
        ).execute()
        
        values = result.get("values", [])
        if not values:
            return jsonify({"rows": [], "headers": []})
        
        headers = values[0] if values else []
        rows = []
        for row in values[1:]:
            row_dict = {}
            for i, header in enumerate(headers):
                row_dict[header.lower().replace(" ", "_")] = row[i] if i < len(row) else ""
            rows.append(row_dict)
        
        return jsonify({"rows": rows, "headers": headers})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/sheets/append", methods=["POST"])
def api_sheets_append():
    """Append a row to a Google Sheet."""
    data = request.get_json(force=True)
    sheet_id = data.get("sheet_id", "")
    range_name = data.get("range", "Sheet1!A:G")
    values = data.get("values", [])
    
    if not sheet_id:
        return jsonify({"error": "sheet_id required"}), 400
    if not values:
        return jsonify({"error": "values required"}), 400
    
    service = get_sheets_service()
    if not service:
        return jsonify({"error": "Google Sheets not connected"}), 401
    
    try:
        body = {"values": [values]}
        result = service.spreadsheets().values().append(
            spreadsheetId=sheet_id,
            range=range_name,
            valueInputOption="USER_ENTERED",
            insertDataOption="INSERT_ROWS",
            body=body
        ).execute()
        
        return jsonify({"ok": True, "updates": result.get("updates", {})})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/sheets/create", methods=["POST"])
def api_sheets_create():
    """Create a new Google Sheet."""
    data = request.get_json(force=True)
    title = data.get("title", "TonyOS Budget")
    
    service = get_sheets_service()
    if not service:
        return jsonify({"error": "Google Sheets not connected"}), 401
    
    try:
        spreadsheet = {
            "properties": {"title": title},
            "sheets": [{
                "properties": {"title": "Transactions"}
            }]
        }
        result = service.spreadsheets().create(body=spreadsheet).execute()
        sheet_id = result.get("spreadsheetId")
        
        headers = [["Date", "Type", "Category", "Merchant", "Amount", "Note"]]
        service.spreadsheets().values().update(
            spreadsheetId=sheet_id,
            range="Transactions!A1:F1",
            valueInputOption="RAW",
            body={"values": headers}
        ).execute()
        
        return jsonify({"ok": True, "sheet_id": sheet_id, "url": f"https://docs.google.com/spreadsheets/d/{sheet_id}"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


TONYOS_TABS = {
    "Connections": ["First Name", "Last Name", "Title", "Company", "Email", "Phone", "Industry", "LinkedIn", "Stage", "Next Step", "Notes", "Sent", "Created"],
    "Outbound": ["Date", "Contact", "Company", "Channel", "Type", "Outcome", "Next Step", "Notes"],
    "Budget": ["Date", "Type", "Category", "Merchant", "Amount", "Note"],
    "Deals": ["Name", "Company", "Value", "Stage", "Created", "Closed", "Notes"]
}


@app.route("/api/sheets/create-master", methods=["POST"])
def api_sheets_create_master():
    """Create the master TonyOS Data spreadsheet with all tabs."""
    service = get_sheets_service()
    if not service:
        return jsonify({"error": "Google Sheets not connected"}), 401
    
    try:
        sheets_list = [{"properties": {"title": tab_name}} for tab_name in TONYOS_TABS.keys()]
        spreadsheet = {
            "properties": {"title": "TonyOS Data"},
            "sheets": sheets_list
        }
        result = service.spreadsheets().create(body=spreadsheet).execute()
        sheet_id = result.get("spreadsheetId")
        
        for tab_name, headers in TONYOS_TABS.items():
            col_letter = chr(ord('A') + len(headers) - 1)
            service.spreadsheets().values().update(
                spreadsheetId=sheet_id,
                range=f"{tab_name}!A1:{col_letter}1",
                valueInputOption="RAW",
                body={"values": [headers]}
            ).execute()
        
        return jsonify({
            "ok": True, 
            "sheet_id": sheet_id, 
            "url": f"https://docs.google.com/spreadsheets/d/{sheet_id}",
            "tabs": list(TONYOS_TABS.keys())
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/sheets/tabs", methods=["GET"])
def api_sheets_tabs():
    """Get list of tabs in a spreadsheet."""
    sheet_id = request.args.get("sheet_id", "")
    
    if not sheet_id:
        return jsonify({"error": "sheet_id required"}), 400
    
    service = get_sheets_service()
    if not service:
        return jsonify({"error": "Google Sheets not connected"}), 401
    
    try:
        result = service.spreadsheets().get(spreadsheetId=sheet_id).execute()
        tabs = [sheet["properties"]["title"] for sheet in result.get("sheets", [])]
        return jsonify({"tabs": tabs})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/sheets/add-tab", methods=["POST"])
def api_sheets_add_tab():
    """Add a new tab to an existing spreadsheet."""
    data = request.get_json(force=True)
    sheet_id = data.get("sheet_id", "")
    tab_name = data.get("tab_name", "")
    headers = data.get("headers", [])
    
    if not sheet_id:
        return jsonify({"error": "sheet_id required"}), 400
    if not tab_name:
        return jsonify({"error": "tab_name required"}), 400
    
    service = get_sheets_service()
    if not service:
        return jsonify({"error": "Google Sheets not connected"}), 401
    
    try:
        request_body = {
            "requests": [{
                "addSheet": {
                    "properties": {"title": tab_name}
                }
            }]
        }
        service.spreadsheets().batchUpdate(spreadsheetId=sheet_id, body=request_body).execute()
        
        if headers:
            col_letter = chr(ord('A') + len(headers) - 1)
            service.spreadsheets().values().update(
                spreadsheetId=sheet_id,
                range=f"{tab_name}!A1:{col_letter}1",
                valueInputOption="RAW",
                body={"values": [headers]}
            ).execute()
        
        return jsonify({"ok": True, "tab_name": tab_name})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-discovery/search", methods=["POST"])
def api_market_discovery_search():
    """Search for companies using Bing Web Search API - legal market intelligence."""
    data = request.get_json(force=True)
    query = data.get("query", "")
    bing_key = data.get("bingKey", "")
    count = data.get("count", 10)
    
    if not query:
        return jsonify({"error": "Query required"}), 400
    if not bing_key:
        return jsonify({"error": "Bing API key required"}), 400
    
    try:
        import requests as req
        headers = {"Ocp-Apim-Subscription-Key": bing_key}
        params = {"q": query, "count": min(count, 50), "mkt": "en-US"}
        
        response = req.get(
            "https://api.bing.microsoft.com/v7.0/search",
            headers=headers,
            params=params,
            timeout=10
        )
        response.raise_for_status()
        
        results = response.json().get("webPages", {}).get("value", [])
        
        return jsonify([{
            "name": r.get("name", ""),
            "url": r.get("url", ""),
            "snippet": r.get("snippet", "")
        } for r in results])
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-discovery/openai", methods=["POST"])
def api_market_discovery_openai():
    """Discover companies using OpenAI API - AI-powered market intelligence."""
    data = request.get_json(force=True)
    industry = data.get("industry", "")
    count = min(data.get("count", 10), 20)
    
    if not industry:
        return jsonify({"error": "Industry required"}), 400
    
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return jsonify({"error": "OpenAI API key not configured"}), 500
    
    prompt = f"""You are a market intelligence analyst for Natoli Engineering Company, which sells tablet presses, tooling, and formulation services to manufacturers.

Find {count} real US-based companies in this industry: "{industry}"

For each company, provide ALL of these fields:
- company: Full company name
- city: City in USA
- state: State abbreviation (e.g., CA, TX, NJ)
- website: Company website URL (required - look it up)
- phone: Company main phone number (format: +1-XXX-XXX-XXXX, or empty string if unknown)
- linkedin: Company LinkedIn URL (format: https://linkedin.com/company/..., or empty string if unknown)
- fit: "Engineering" (makes tablets/needs presses), "Scientific" (formulation/R&D), or "Both"
- reason: One sentence explaining why they would need Natoli products
- confidence: 0-100 score for how confident you are this is accurate

Return ONLY a valid JSON array of objects. No explanation text.
Focus on real companies that actually manufacture tablets, capsules, pellets, or compressed products.

Example format:
[
  {{"company": "Catalent Pharma Solutions", "city": "Somerset", "state": "NJ", "website": "https://catalent.com", "phone": "+1-732-537-6200", "linkedin": "https://linkedin.com/company/catalent", "fit": "Both", "reason": "Large CDMO with tablet manufacturing and formulation development", "confidence": 95}},
  {{"company": "Lonza Group", "city": "Greenwood", "state": "SC", "website": "https://lonza.com", "phone": "+1-864-223-7831", "linkedin": "https://linkedin.com/company/lonza", "fit": "Both", "reason": "Global CDMO with solid oral dosage capabilities", "confidence": 90}}
]"""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a market research analyst. Return only valid JSON arrays."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        content = response.choices[0].message.content.strip()
        
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]
        
        import json
        companies = json.loads(content)
        
        return jsonify({"companies": companies})
        
    except json.JSONDecodeError as e:
        return jsonify({"error": f"Failed to parse AI response: {str(e)}", "raw": content}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


NATOLI_COMPETITORS = [
    "fette compacting", "fette", "korsch", "romaco", "ima", "ima group",
    "bosch packaging", "syntegon", "gea", "glatt", "manesty", "kikusui",
    "stokes", "cadmach", "sejong", "fluidpack", "chamunda", "karnavati",
    "riddhi", "shakti", "pharmachine", "tablet press co", "patterson kelley",
    "elizabeth companies", "elizabeth hata", "hata", "kg pharma", "killian",
    "key international", "riva", "piccola", "dott bonapace", "mini press",
    "capplus", "clit", "zhejiang", "shanghai tianhe", "tianhe"
]

def is_natoli_competitor(company_name):
    """Check if company is a known Natoli competitor."""
    name_lower = company_name.lower()
    for competitor in NATOLI_COMPETITORS:
        if competitor in name_lower:
            return True, competitor
    return False, None


@app.route("/api/manufacturing-intel", methods=["POST"])
def api_manufacturing_intel():
    """AI-powered manufacturing intelligence - determine if company manufactures in-house or uses 3rd party."""
    data = request.get_json(force=True)
    company = data.get("company", "").strip()
    product_type = data.get("productType", "Tablet")
    claimed_location = data.get("claimedLocation", "")
    
    if not company:
        return jsonify({"error": "Company name required"}), 400
    
    is_competitor, competitor_match = is_natoli_competitor(company)
    if is_competitor:
        return jsonify({
            "success": True, 
            "result": {
                "company": company,
                "productType": "Competitor",
                "manufacturingModel": "Competitor",
                "manufacturingExplanation": f"This company ({competitor_match}) is a direct Natoli competitor - they manufacture tablet presses or tooling.",
                "likelyLocation": "N/A",
                "natoliFit": "No Fit",
                "natoliFitReason": "Direct competitor - do not pursue",
                "division": "No Fit",
                "evidenceSignals": [{"signal": "Competitor Detection", "detected": True, "notes": f"Matched: {competitor_match}"}],
                "confidence": "High",
                "confidenceScore": 100,
                "keyContacts": "N/A",
                "nextSteps": "SKIP - This is a Natoli competitor",
                "isCompetitor": True
            }
        })
    
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return jsonify({"error": "OpenAI API key not configured"}), 500
    
    product_instruction = ""
    if product_type == "Auto":
        product_instruction = """First determine what type of compressed/pressed products this company makes. Focus on:
- Pharmaceutical tablets (Rx, OTC)
- Nutraceuticals/Vitamins/Supplements (sold at Amazon, Walmart, CVS, Walgreens, GNC, Vitamin Shoppe)
- Veterinary tablets/chews (pet supplements, medications)
- Consumer pressed products (dishwasher tablets, bath bombs, cleaning tablets, laundry pods)
- Cosmetic pressed powders (eyeshadow, foundation, blush)
- Food/confectionery (candy tablets, mints, bouillon cubes)
- Industrial pressed products (catalyst pellets, abrasives, battery materials)"""
    else:
        product_instruction = f"Product focus: {product_type}"
    
    prompt = f"""You are a manufacturing intelligence analyst for Natoli Engineering Company, which sells tablet presses, tooling, and compression equipment to companies that MAKE pressed/compressed products.

CRITICAL: We are looking for CUSTOMERS, not competitors. If this company sells tablet presses, punches, dies, or compression tooling - they are a COMPETITOR and not a fit.

Research this company: "{company}"
{product_instruction}
{f'Claimed location: {claimed_location}' if claimed_location else ''}

SHELF PRODUCT ANALYSIS:
Check if this company sells products at major retailers (Amazon, Walmart, Target, Home Depot, Costco, CVS, Walgreens) that involve tablet compression:
- Vitamins, supplements, nutraceuticals (pill/tablet form)
- OTC medications (tablets, caplets)
- Dishwasher tablets, laundry pods
- Bath bombs, shower steamers
- Cleaning tablets (toilet, surface)
- Candy tablets, mints, pressed confections
- Pet supplements, chews
- Any other compressed/pressed consumer product

Analyze and determine:

1. **Manufacturing Model**: Do they manufacture in-house (own tablet presses), use CDMOs, or hybrid?
2. **Retail Presence**: Are their products sold at major retailers? What products?
3. **Natoli Fit**: 
   - Engineering Division: In-house manufacturing with tablet presses/tooling needs
   - Scientific Division: R&D, formulation, method development needs
4. **Evidence Signals**: What signals support in-house manufacturing?

Consider these signal types:
- FDA Registration (own manufacturing facility)
- Job Postings (tablet press operators, compression technicians, packaging operators)
- LinkedIn Manufacturing Density (production staff at facilities)
- "Manufactured For" Language (outsourced = CDMO users, not direct Natoli customers)
- cGMP Facility Claims (own manufacturing capabilities)
- Retail Listings (products on Amazon, Walmart, etc.)
- Factory Images (do they show production equipment?)

Return a JSON object with these exact fields:
{{
  "company": "Full company name",
  "productType": "Detected product type (Pharmaceutical, Nutraceutical, Veterinary, Consumer/Household, Cosmetic, Food/Confection, Industrial, Other)",
  "retailPresence": "Where products are sold (Amazon, Walmart, etc.) or 'Not Found'",
  "shelfProducts": "Specific products found on shelves that use compression (e.g., 'Vitamin D tablets, Probiotic capsules')",
  "manufacturingModel": "In-house" | "Outsourced (CDMO)" | "Hybrid" | "Unable to Determine",
  "manufacturingExplanation": "2-3 sentence explanation with evidence of how they manufacture",
  "headquarters": "City, State/Country of corporate headquarters",
  "manufacturingLocations": ["List of known manufacturing facility locations (City, State/Country)", "Include plant names if known"],
  "likelyLocation": "Primary manufacturing location if known, or 'Unknown'",
  "natoliFit": "Strong" | "Medium" | "Weak" | "No Fit",
  "natoliFitReason": "Why they are or aren't a good fit for Natoli products",
  "division": "Engineering" | "Scientific" | "Both" | "No Fit",
  "evidenceSignals": [
    {{"signal": "signal name", "detected": true/false, "notes": "brief note"}}
  ],
  "confidence": "High" | "Medium" | "Low",
  "confidenceScore": 0-100,
  "keyContacts": "Types of people to reach out to (e.g., Director of Manufacturing, VP Operations)",
  "nextSteps": "Recommended action for sales outreach",
  "isCompetitor": false
}}

IMPORTANT: If this company MAKES tablet presses or tooling (not uses them), set natoliFit to "No Fit" and isCompetitor to true.

Return ONLY valid JSON. No markdown, no explanation text."""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a manufacturing intelligence analyst. Research companies that make compressed/pressed products (tablets, supplements, etc.) and determine their manufacturing model. Focus on retail/shelf products. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=2000
        )
        
        content = response.choices[0].message.content.strip()
        
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]
        content = content.strip()
        
        import json
        result = json.loads(content)
        
        if result.get("isCompetitor"):
            result["natoliFit"] = "No Fit"
            result["division"] = "No Fit"
            result["nextSteps"] = "SKIP - This is a Natoli competitor"
        
        return jsonify({"success": True, "result": result})
        
    except json.JSONDecodeError as e:
        return jsonify({"error": f"Failed to parse AI response: {str(e)}", "raw": content}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


def detect_natoli_industry(company_name, website=""):
    """Detect industry based on company name and website using Natoli's 25+ validated markets."""
    text = (company_name + " " + website).lower()
    
    industry_keywords = {
        'Nuclear Fuel/SMR': ['nuclear', 'uranium', 'reactor', 'fuel pellet', 'smr', 'fission', 'bwx', 'framatome', 'areva', 'gnf', 'centrus', 'urenco', 'nuclear fuel', 'enrichment'],
        'Hydrogen Storage': ['hydrogen', 'fuel cell', 'h2 ', 'electrolyzer', 'plug power', 'nel ', 'hydrogenics', 'proton onsite', 'nuvera', 'powercell', 'giner', 'mcphy', 'hexagon lincoln', 'hyon', 'stratosfuel'],
        'Carbon Capture': ['carbon capture', 'carbon clean', 'direct air', 'dac ', 'co2 capture', 'climeworks', 'carboncure', 'lanzatech', 'global thermostat', 'svante', 'air capture', 'carbonfree', 'charm industrial', 'heirloom', 'ebb carbon', 'noya', 'remora', 'carbonbuilt', 'mission zero', 'cleano2', 'opus 12'],
        'Catalyst Manufacturing': ['catalyst', 'catalytic', 'grace', 'uop', 'basf catalyst', 'johnson matthey', 'shell catalyst', 'albemarle'],
        'Advanced Ceramics': ['ceramic', 'ceradyne', 'coorstek', 'corning', 'blasch', 'rauschert', 'sintering', 'alumina', 'zirconia', 'advanced ceramic'],
        'Battery Manufacturing': ['battery', 'lithium', 'sodium-ion', 'solid-state', 'cathode', 'anode', 'cell manufacturing', 'quantumscape', 'solid power'],
        'Rare Earth/Magnets': ['rare earth', 'magnet', 'neodymium', 'permanent magnet', 'mp materials'],
        'Semiconductor': ['semiconductor', 'wafer', 'chip', 'silicon', 'intel', 'tsmc', 'micron'],
        'Explosives/Defense': ['explosive', 'propellant', 'ammunition', 'ballistic', 'defense', 'munition', 'northrop', 'raytheon', 'lockheed'],
        'Pharmaceutical': ['pharma', 'pharmaceutical', 'drug', 'medication', 'rx ', 'therapeutics', 'pfizer', 'merck', 'novartis', 'lilly', 'abbvie', 'bristol-myers'],
        'Generic Pharma': ['generic', 'generics', 'teva', 'mylan', 'sandoz', 'amneal', 'lupin'],
        'Nutraceutical': ['nutraceutical', 'dietary', 'supplement', 'vitamin', 'nutrition', 'health product', 'gnc', 'nature made'],
        'Veterinary Pharma': ['veterinary', 'animal health', 'vet ', 'pet med', 'zoetis', 'elanco', 'boehringer animal'],
        'Cannabis': ['cannabis', 'marijuana', 'thc', 'cbd', 'dispensary'],
        '3D Printing/Additive': ['3d print', 'additive', 'metal powder', 'feedstock', 'binder jet', 'desktop metal', 'markforged'],
        'Abrasives': ['abrasive', 'grinding', 'polishing compound', 'saint-gobain abrasive', 'norton'],
        'Cosmetics': ['cosmetic', 'pressed powder', 'makeup', 'beauty', 'loreal', 'estee lauder'],
        'Medical Devices': ['medical device', 'implant', 'orthopedic', 'dental implant', 'stryker', 'medtronic', 'zimmer'],
        'Animal Feed': ['animal feed', 'feed supplement', 'livestock', 'poultry feed', 'cargill feed', 'purina'],
        'Chemicals': ['chemical', 'dow ', 'basf', 'dupont', 'lyondell', 'chevron phillips', 'ineos', 'flint hills', 'american elements', 'praxair', 'linde'],
        'Energy/Utilities': ['energy', 'utility', 'power company', 'electric', 'pge', 'pseg', 'dominion', 'southern company', 'energy harbor', 'terrestrial energy', 'x-energy']
    }
    
    for industry, keywords in industry_keywords.items():
        for kw in keywords:
            if kw in text:
                return industry
    return 'Manufacturing'

@app.route("/api/compaction-intelligence", methods=["POST"])
def api_compaction_intelligence():
    """AI-powered compaction intelligence for sales rep enablement."""
    data = request.get_json(force=True)
    company = data.get("company", "").strip()
    title = data.get("title", "").strip()
    function = data.get("function", "").strip()
    website = data.get("website", "").strip()
    
    if not company:
        return jsonify({"error": "Company name required"}), 400
    
    detected_industry = detect_natoli_industry(company, website)
    
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return jsonify({"error": "OpenAI API key not configured"}), 500
    
    system_prompt = """You are a manufacturing reasoning assistant for Natoli Engineering Company.
Do NOT invent specifications or make up numbers.
Use probabilistic language only (e.g., "likely", "probably", "typically").
Focus on compaction/tableting context - Natoli sells tablet presses, tooling, and compression equipment.
Infer likely tooling class, press OEM exposure, production stage, probable tablet defects, and a safe Natoli discussion angle."""

    prompt = f"""Analyze this manufacturing contact:

Company: {company}
Title: {title}
Function: {function}
Industry: {detected_industry}

Based on the company type, job title, and function, infer:
1. Tooling_Class (Likely) - B, D, DB, EU, or custom based on company scale/type
2. Press_OEM (Likely) - Which tablet press brands they probably use (Fette, Korsch, Manesty, IMA, Stokes, etc.)
3. Production_Stage - R&D, Scale-up, or Production
4. Probable_Defects - Common compaction issues they likely face (capping, lamination, sticking, picking, weight variation, etc.)
5. Natoli_Angle - A safe, relevant discussion topic for a Natoli rep to open with
6. Email_Subject - A compelling email subject line (under 50 chars)
7. LinkedIn_Opening - A personalized LinkedIn connection request opener (under 200 chars)

Return ONLY valid JSON with these exact keys:
{{
  "Tooling_Class": "...",
  "Press_OEM": "...",
  "Production_Stage": "...",
  "Probable_Defects": "...",
  "Natoli_Angle": "...",
  "Email_Subject": "...",
  "LinkedIn_Opening": "..."
}}"""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2,
            max_tokens=800
        )
        
        content = response.choices[0].message.content.strip()
        
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]
        content = content.strip()
        
        import json
        result = json.loads(content)
        result["Industry"] = detected_industry
        
        return jsonify({"success": True, "result": result})
        
    except json.JSONDecodeError as e:
        return jsonify({"error": f"Failed to parse AI response: {str(e)}"}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


batch_jobs = {}

@app.route("/api/manufacturing-intel/batch", methods=["POST"])
def api_manufacturing_intel_batch_start():
    """Start a background batch scan job."""
    import uuid
    import threading
    
    data = request.get_json(force=True)
    companies = data.get("companies", [])
    product_type = data.get("productType", "Auto")
    
    if not companies:
        return jsonify({"error": "No companies provided"}), 400
    
    job_id = str(uuid.uuid4())[:8]
    
    batch_jobs[job_id] = {
        "status": "running",
        "total": len(companies),
        "completed": 0,
        "current": "",
        "results": [],
        "started_at": datetime.now().isoformat()
    }
    
    def run_batch():
        import time
        try:
            for company in companies:
                if job_id not in batch_jobs:
                    break
                if batch_jobs[job_id]["status"] == "cancelled":
                    break
                
                batch_jobs[job_id]["current"] = company
                
                is_competitor, competitor_match = is_natoli_competitor(company)
                if is_competitor:
                    batch_jobs[job_id]["results"].append({
                        "company": company,
                        "productType": "Competitor",
                        "manufacturingModel": "Competitor",
                        "natoliFit": "No Fit",
                        "division": "No Fit",
                        "confidence": "High",
                        "confidenceScore": 100,
                        "isCompetitor": True,
                        "nextSteps": "SKIP - Natoli competitor"
                    })
                    batch_jobs[job_id]["completed"] += 1
                    continue
                
                try:
                    with app.app_context():
                        result = analyze_company_intel(company, product_type)
                        batch_jobs[job_id]["results"].append(result)
                except Exception as e:
                    batch_jobs[job_id]["results"].append({
                        "company": company,
                        "error": str(e),
                        "natoliFit": "Error"
                    })
                
                batch_jobs[job_id]["completed"] += 1
                time.sleep(0.5)
            
            if job_id in batch_jobs:
                batch_jobs[job_id]["status"] = "completed"
                batch_jobs[job_id]["current"] = ""
        except Exception as e:
            if job_id in batch_jobs:
                batch_jobs[job_id]["status"] = "error"
                batch_jobs[job_id]["error"] = str(e)
    
    thread = threading.Thread(target=run_batch)
    thread.daemon = True
    thread.start()
    
    return jsonify({"success": True, "jobId": job_id, "total": len(companies)})


@app.route("/api/manufacturing-intel/batch/<job_id>", methods=["GET"])
def api_manufacturing_intel_batch_status(job_id):
    """Get status of a batch scan job."""
    if job_id not in batch_jobs:
        return jsonify({"error": "Job not found"}), 404
    
    job = batch_jobs[job_id]
    return jsonify({
        "jobId": job_id,
        "status": job["status"],
        "total": job["total"],
        "completed": job["completed"],
        "current": job["current"],
        "results": job["results"],
        "startedAt": job["started_at"]
    })


@app.route("/api/manufacturing-intel/batch/<job_id>/cancel", methods=["POST"])
def api_manufacturing_intel_batch_cancel(job_id):
    """Cancel a batch scan job."""
    if job_id not in batch_jobs:
        return jsonify({"error": "Job not found"}), 404
    
    batch_jobs[job_id]["status"] = "cancelled"
    return jsonify({"success": True, "message": "Job cancelled"})


territory_jobs = {}

TERRITORY_MAP = {
    'CA': 'West Coast', 'OR': 'West Coast', 'WA': 'West Coast', 'HI': 'West Coast', 'AK': 'West Coast',
    'CO': 'Rocky Mountain', 'AZ': 'Rocky Mountain', 'UT': 'Rocky Mountain', 'NV': 'Rocky Mountain', 
    'NM': 'Rocky Mountain', 'WY': 'Rocky Mountain', 'MT': 'Rocky Mountain', 'ID': 'Rocky Mountain',
    'IL': 'Midwest', 'OH': 'Midwest', 'MI': 'Midwest', 'TX': 'Midwest', 'MN': 'Midwest', 
    'WI': 'Midwest', 'IN': 'Midwest', 'MO': 'Midwest', 'IA': 'Midwest', 'KS': 'Midwest',
    'NE': 'Midwest', 'OK': 'Midwest', 'ND': 'Midwest', 'SD': 'Midwest', 'AR': 'Midwest', 'LA': 'Midwest',
    'FL': 'Southeast', 'GA': 'Southeast', 'TN': 'Southeast', 'NC': 'Southeast', 'SC': 'Southeast',
    'AL': 'Southeast', 'MS': 'Southeast', 'KY': 'Southeast', 'VA': 'Southeast', 'WV': 'Southeast',
    'NY': 'Northeast', 'PA': 'Northeast', 'MA': 'Northeast', 'NJ': 'Northeast', 'CT': 'Northeast',
    'NH': 'Northeast', 'VT': 'Northeast', 'ME': 'Northeast', 'RI': 'Northeast', 'MD': 'Northeast', 'DE': 'Northeast', 'DC': 'Northeast'
}

ENGINEERING_KEYWORDS = ['tablet press', 'tooling', 'die', 'punch', 'compression', 'powder metallurgy', 
                         'ceramic', 'pellet', 'pressing', 'compaction', 'manufacturing', 'production',
                         'operations', 'plant', 'facility', 'equipment']
SCIENTIFIC_KEYWORDS = ['r&d', 'research', 'formulation', 'analytical', 'pilot', 'scale-up', 
                        'material science', 'lab', 'development', 'scientist', 'laboratory']
EXCLUDE_KEYWORDS = ['clinic', 'dealer', 'distributor', 'retail', 'service', 'repair', 'hospital',
                     'veterinary clinic', 'pharmacy', 'drugstore', 'staffing', 'recruiting']

TIER1_TITLES = ['ceo', 'coo', 'cfo', 'cto', 'cso', 'cmo', 'chief', 'president', 'owner', 'founder',
                'vp', 'vice president', 'director', 'head of', 'senior director', 'executive']
TIER2_TITLES = ['manager', 'senior manager', 'lead', 'supervisor', 'principal']
TIER3_TITLES = ['engineer', 'scientist', 'analyst', 'specialist', 'coordinator', 'associate']


def classify_territory_row(row):
    """Classify a single row from territory CSV."""
    company = row.get('Company Name', row.get('company', row.get('Organization Name', ''))).strip()
    title = row.get('Title', row.get('title', row.get('Job Title', ''))).strip().lower()
    industry = row.get('Industry', row.get('industry', '')).strip()
    state = row.get('State', row.get('state', row.get('Company State', ''))).strip().upper()
    
    for kw in EXCLUDE_KEYWORDS:
        if kw in company.lower() or kw in title:
            return None
    
    fit = None
    title_industry = f"{title} {industry}".lower()
    
    for kw in ENGINEERING_KEYWORDS:
        if kw in title_industry:
            fit = 'Natoli Engineering'
            break
    
    for kw in SCIENTIFIC_KEYWORDS:
        if kw in title_industry:
            if fit == 'Natoli Engineering':
                fit = 'Both'
            else:
                fit = 'Natoli Scientific'
            break
    
    if not fit:
        return None
    
    territory = TERRITORY_MAP.get(state[:2] if state else '', 'Unknown')
    
    tier = 'Tier 3'
    for t1 in TIER1_TITLES:
        if t1 in title:
            tier = 'Tier 1'
            break
    if tier != 'Tier 1':
        for t2 in TIER2_TITLES:
            if t2 in title:
                tier = 'Tier 2'
                break
    
    return {
        **row,
        'Natoli_Fit': fit,
        'Territory': territory,
        'Contact_Tier': tier,
        'Industry_Clean': industry or 'Other'
    }


@app.route("/api/territory-intel/process", methods=["POST"])
def api_territory_intel_process():
    """Start background territory intelligence processing."""
    import uuid
    import threading
    import csv
    import io
    
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    if not file.filename.endswith('.csv'):
        return jsonify({"error": "File must be a CSV"}), 400
    
    content = file.read().decode('utf-8')
    
    job_id = str(uuid.uuid4())[:8]
    
    territory_jobs[job_id] = {
        "status": "running",
        "total": 0,
        "completed": 0,
        "qualified": 0,
        "excluded": 0,
        "results": [],
        "stats": {},
        "started_at": datetime.now().isoformat()
    }
    
    def process_csv():
        import time
        reader = csv.DictReader(io.StringIO(content))
        rows = list(reader)
        
        territory_jobs[job_id]["total"] = len(rows)
        
        results = []
        excluded = 0
        
        for i, row in enumerate(rows):
            if territory_jobs[job_id]["status"] == "cancelled":
                break
            
            territory_jobs[job_id]["completed"] = i + 1
            
            classified = classify_territory_row(row)
            if classified:
                results.append(classified)
            else:
                excluded += 1
            
            if i % 100 == 0:
                territory_jobs[job_id]["results"] = results
                territory_jobs[job_id]["qualified"] = len(results)
                territory_jobs[job_id]["excluded"] = excluded
                time.sleep(0.01)
        
        eng_count = len([r for r in results if r.get('Natoli_Fit') == 'Natoli Engineering'])
        sci_count = len([r for r in results if r.get('Natoli_Fit') == 'Natoli Scientific'])
        both_count = len([r for r in results if r.get('Natoli_Fit') == 'Both'])
        
        territories = {}
        for r in results:
            t = r.get('Territory', 'Unknown')
            if t not in territories:
                territories[t] = 0
            territories[t] += 1
        
        territory_jobs[job_id]["results"] = results
        territory_jobs[job_id]["qualified"] = len(results)
        territory_jobs[job_id]["excluded"] = excluded
        territory_jobs[job_id]["stats"] = {
            "engineering": eng_count,
            "scientific": sci_count,
            "both": both_count,
            "territories": territories,
            "uniqueCompanies": len(set(r.get('Company Name', r.get('company', '')).lower() for r in results))
        }
        territory_jobs[job_id]["status"] = "completed"
    
    thread = threading.Thread(target=process_csv)
    thread.daemon = True
    thread.start()
    
    return jsonify({"success": True, "jobId": job_id})


@app.route("/api/territory-intel/status/<job_id>", methods=["GET"])
def api_territory_intel_status(job_id):
    """Get status of territory processing job."""
    if job_id not in territory_jobs:
        return jsonify({"error": "Job not found"}), 404
    
    job = territory_jobs[job_id]
    return jsonify({
        "jobId": job_id,
        "status": job["status"],
        "total": job["total"],
        "completed": job["completed"],
        "qualified": job["qualified"],
        "excluded": job["excluded"],
        "stats": job["stats"],
        "resultCount": len(job["results"]),
        "startedAt": job["started_at"]
    })


@app.route("/api/territory-intel/results/<job_id>", methods=["GET"])
def api_territory_intel_results(job_id):
    """Get full results of territory processing job."""
    if job_id not in territory_jobs:
        return jsonify({"error": "Job not found"}), 404
    
    job = territory_jobs[job_id]
    return jsonify({
        "jobId": job_id,
        "status": job["status"],
        "results": job["results"],
        "stats": job["stats"]
    })


@app.route("/api/territory-intel/cancel/<job_id>", methods=["POST"])
def api_territory_intel_cancel(job_id):
    """Cancel territory processing job."""
    if job_id not in territory_jobs:
        return jsonify({"error": "Job not found"}), 404
    
    territory_jobs[job_id]["status"] = "cancelled"
    return jsonify({"success": True})


@app.route("/api/territory-intel/jobs", methods=["GET"])
def api_territory_intel_jobs():
    """List all territory processing jobs."""
    jobs = []
    for job_id, job in territory_jobs.items():
        jobs.append({
            "jobId": job_id,
            "status": job["status"],
            "total": job["total"],
            "qualified": job["qualified"],
            "startedAt": job["started_at"]
        })
    return jsonify({"jobs": sorted(jobs, key=lambda x: x["startedAt"], reverse=True)})


NATOLI_INDUSTRIES = {
    "core": [
        "Pharmaceutical (Rx tablets, OTC medications)",
        "Generic Pharma (ANDA manufacturers)",
        "Nutraceutical / Dietary Supplements (vitamins, minerals, herbals)",
        "Veterinary Pharma (animal health tablets, chews)",
        "Cannabis / CBD (edibles, pressed products)"
    ],
    "energy_materials": [
        "Nuclear Fuel / SMR (uranium pellets, fuel rods)",
        "Battery Manufacturing (solid-state, lithium, sodium-ion electrodes)",
        "Hydrogen Storage (metal hydride pellets)",
        "Catalyst Manufacturing (catalytic pellets, zeolites)",
        "Carbon Capture Sorbents (CO2 capture pellets)"
    ],
    "advanced_manufacturing": [
        "Advanced Ceramics (aerospace, defense, technical ceramics)",
        "3D Printing Feedstock (metal/ceramic powders, binder jetting)",
        "Rare Earth Magnets (neodymium, permanent magnets)",
        "Semiconductor Materials (wafer processing, substrates)",
        "Space Materials (satellite components, aerospace composites)"
    ],
    "defense": [
        "Explosives / Propellants (energetic materials, pyrotechnics)",
        "Ammunition / Ballistics (projectiles, casings)"
    ],
    "industrial_specialty": [
        "Abrasives (grinding wheels, polishing compounds)",
        "Medical Implants (orthopedic, dental ceramics)",
        "Cosmetics Pressed Powders (makeup, skincare compacts)",
        "Agricultural Micronutrient Pellets (fertilizer tablets)",
        "Animal Feed Supplements (feed additives, blocks)",
        "Forensic Standards (reference materials, calibration)",
        "Recycling / Reclaimed Materials (recycled pellets)",
        "Art Conservation / Pigments (pressed pigments, restoration)",
        "Consumer Pressed Products (dishwasher tablets, bath bombs, cleaning tablets)",
        "Food / Confectionery (candy tablets, mints, breath strips)"
    ]
}

ALL_NATOLI_INDUSTRIES = []
for category in NATOLI_INDUSTRIES.values():
    ALL_NATOLI_INDUSTRIES.extend(category)


INDUSTRY_KEYWORDS = {
    "Pharmaceutical": ["pharma", "drug", "medication", "rx", "prescription", "fda", "therapeutic", "api", "active pharmaceutical"],
    "Generic Pharma": ["generic", "anda", "abbreviated new drug"],
    "Nutraceutical": ["vitamin", "supplement", "nutraceutical", "dietary", "mineral", "herbal", "probiotic", "omega", "gnc", "nature made", "solgar"],
    "Veterinary": ["veterinary", "animal health", "pet", "livestock", "chewable", "zoetis"],
    "Cannabis": ["cannabis", "cbd", "thc", "hemp", "marijuana", "edible"],
    "Nuclear": ["nuclear", "uranium", "fuel rod", "reactor", "enrichment", "smr", "small modular"],
    "Battery": ["battery", "electrode", "lithium", "sodium-ion", "solid-state", "anode", "cathode", "ev battery"],
    "Hydrogen": ["hydrogen", "metal hydride", "fuel cell", "h2 storage"],
    "Catalyst": ["catalyst", "zeolite", "petrochemical", "refinery catalyst"],
    "Carbon Capture": ["carbon capture", "co2", "sorbent", "sequestration", "direct air capture"],
    "Ceramics": ["ceramic", "alumina", "zirconia", "technical ceramic", "sintering", "refractory"],
    "3D Printing": ["additive", "3d print", "powder bed", "binder jet", "metal powder"],
    "Magnets": ["magnet", "neodymium", "rare earth", "permanent magnet", "ndfeb"],
    "Semiconductor": ["semiconductor", "wafer", "substrate", "silicon", "chip fab"],
    "Space": ["aerospace", "satellite", "space", "rocket", "propulsion", "nasa"],
    "Explosives": ["explosive", "propellant", "pyrotechnic", "energetic", "munition", "detonator"],
    "Ammunition": ["ammunition", "bullet", "projectile", "cartridge", "ballistic", "defense contractor"],
    "Abrasives": ["abrasive", "grinding wheel", "polishing", "sandpaper", "cutting wheel"],
    "Medical Implants": ["implant", "orthopedic", "dental", "biomedical", "prosthetic", "medical device"],
    "Cosmetics": ["cosmetic", "makeup", "pressed powder", "compact", "foundation", "beauty"],
    "Agricultural": ["fertilizer", "micronutrient", "agricultural pellet", "crop nutrient"],
    "Animal Feed": ["animal feed", "feed additive", "animal nutrition", "feed block", "livestock supplement"],
    "Consumer": ["dishwasher tablet", "bath bomb", "cleaning tablet", "detergent pod", "effervescent"],
    "Food": ["candy", "mint", "confectionery", "pressed candy", "tablet candy", "breath mint", "lozenge"]
}

def detect_industry(company_name, additional_text=""):
    """Detect industry from company name and context."""
    text = f"{company_name} {additional_text}".lower()
    
    for industry, keywords in INDUSTRY_KEYWORDS.items():
        for kw in keywords:
            if kw.lower() in text:
                return industry
    return "Other Industrial"


def analyze_company_intel(company, product_type):
    """Analyze a single company for manufacturing intelligence."""
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return {"company": company, "error": "OpenAI API key not configured"}
    
    prompt = f"""Analyze this company for Natoli Engineering: "{company}"

Natoli sells tablet presses, tooling (dies/punches), and provides formulation R&D services.
Target customers: ANY company that compresses powders into tablets, pellets, or compacts.

Return a SHORT JSON (keep it brief to avoid truncation):
{{
  "company": "{company}",
  "industry": "Pharma|Nutra|Vet|Cannabis|Nuclear|Battery|Hydrogen|Catalyst|Ceramic|3DPrint|Magnet|Semiconductor|Space|Explosive|Ammo|Abrasive|Implant|Cosmetic|AgFeed|Consumer|Food|Other",
  "product": "What they compress/pelletize",
  "retail": "Amazon,Walmart,CVS,Target or None",
  "model": "In-house|CDMO|Hybrid|Unknown",
  "fit": "Strong|Medium|Weak|NoFit",
  "division": "Engineering|Scientific|Both|NoFit",
  "reason": "Brief reason",
  "confidence": 85
}}

ONLY return valid JSON, no markdown."""

    for attempt in range(3):
        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": "Return only valid JSON. Keep responses short."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.2,
                max_tokens=500
            )
            
            content = response.choices[0].message.content.strip()
            
            if content.startswith("```"):
                lines = content.split("\n")
                content = "\n".join(lines[1:-1] if lines[-1] == "```" else lines[1:])
            if content.startswith("json"):
                content = content[4:].strip()
            
            content = content.strip()
            if not content.endswith("}"):
                last_brace = content.rfind("}")
                if last_brace > 0:
                    content = content[:last_brace + 1]
            
            import json
            result = json.loads(content)
            
            industry_full = {
                "Pharma": "Pharmaceutical", "Nutra": "Nutraceutical", "Vet": "Veterinary",
                "Cannabis": "Cannabis / CBD", "Nuclear": "Nuclear Fuel", "Battery": "Battery Manufacturing",
                "Hydrogen": "Hydrogen Storage", "Catalyst": "Catalyst Manufacturing", "Ceramic": "Advanced Ceramics",
                "3DPrint": "3D Printing", "Magnet": "Rare Earth Magnets", "Semiconductor": "Semiconductor",
                "Space": "Space Materials", "Explosive": "Explosives", "Ammo": "Ammunition",
                "Abrasive": "Abrasives", "Implant": "Medical Implants", "Cosmetic": "Cosmetics",
                "AgFeed": "Agricultural / Animal Feed", "Consumer": "Consumer Products",
                "Food": "Food / Confectionery", "Other": "Other Industrial"
            }
            
            industry_short = result.get("industry", "Other")
            result["industry"] = industry_full.get(industry_short, industry_short)
            result["productType"] = result.pop("product", "Unknown")
            result["retailPresence"] = result.pop("retail", "Unknown")
            result["manufacturingModel"] = result.pop("model", "Unknown")
            result["natoliFit"] = result.pop("fit", "Unknown")
            result["natoliFitReason"] = result.pop("reason", "")
            result["confidenceScore"] = result.pop("confidence", 50)
            result["confidence"] = "High" if result["confidenceScore"] >= 80 else "Medium" if result["confidenceScore"] >= 50 else "Low"
            result["isCompetitor"] = False
            result["nextSteps"] = "Research and outreach" if result["natoliFit"] in ["Strong", "Medium"] else "Low priority"
            
            return result
            
        except json.JSONDecodeError as e:
            if attempt < 2:
                continue
            return {
                "company": company,
                "industry": detect_industry(company, product_type),
                "productType": product_type,
                "natoliFit": "Medium",
                "division": "Both",
                "confidence": "Low",
                "confidenceScore": 30,
                "natoliFitReason": "AI analysis failed - needs manual review",
                "isCompetitor": False,
                "error": f"JSON parse error: {str(e)}"
            }
        except Exception as e:
            return {"company": company, "error": str(e), "natoliFit": "Error"}

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Manufacturing intelligence analyst. Return only valid JSON."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=1500
        )
        
        content = response.choices[0].message.content.strip()
        if content.startswith("```"):
            content = content.split("```")[1]
            if content.startswith("json"):
                content = content[4:]
        content = content.strip()
        
        import json
        result = json.loads(content)
        return result
        
    except Exception as e:
        return {"company": company, "error": str(e), "natoliFit": "Error"}


@app.route("/api/market-scans", methods=["GET"])
def api_get_market_scans():
    """Get all saved market scans from database."""
    try:
        from models import MarketScan
        scans = MarketScan.query.order_by(MarketScan.created_at.desc()).limit(50).all()
        return jsonify({
            "scans": [{
                "id": s.id,
                "name": s.name,
                "status": s.status,
                "current_industry": s.current_industry,
                "total_industries": s.total_industries,
                "total_companies": s.total_companies,
                "created_at": s.created_at.isoformat() if s.created_at else None,
                "updated_at": s.updated_at.isoformat() if s.updated_at else None
            } for s in scans]
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-scans", methods=["POST"])
def api_create_market_scan():
    """Create a new market scan job."""
    try:
        from models import MarketScan
        data = request.get_json(force=True)
        
        scan = MarketScan(
            name=data.get("name", "Industry Discovery"),
            status="pending",
            industries_json=json.dumps(data.get("industries", [])),
            companies_json=json.dumps([]),
            current_industry=0,
            total_industries=len(data.get("industries", [])),
            total_companies=0,
            options_json=json.dumps(data.get("options", {}))
        )
        db.session.add(scan)
        db.session.commit()
        
        return jsonify({"id": scan.id, "status": "created"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-scans/<int:scan_id>", methods=["GET"])
def api_get_market_scan(scan_id):
    """Get a specific market scan with full data."""
    try:
        from models import MarketScan
        scan = MarketScan.query.get(scan_id)
        if not scan:
            return jsonify({"error": "Scan not found"}), 404
        
        return jsonify({
            "id": scan.id,
            "name": scan.name,
            "status": scan.status,
            "industries": json.loads(scan.industries_json) if scan.industries_json else [],
            "companies": json.loads(scan.companies_json) if scan.companies_json else [],
            "current_industry": scan.current_industry,
            "total_industries": scan.total_industries,
            "total_companies": scan.total_companies,
            "options": json.loads(scan.options_json) if scan.options_json else {},
            "error_message": scan.error_message,
            "created_at": scan.created_at.isoformat() if scan.created_at else None,
            "updated_at": scan.updated_at.isoformat() if scan.updated_at else None
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-scans/<int:scan_id>", methods=["PATCH"])
def api_update_market_scan(scan_id):
    """Update a market scan (add companies, update progress)."""
    try:
        from models import MarketScan
        scan = MarketScan.query.get(scan_id)
        if not scan:
            return jsonify({"error": "Scan not found"}), 404
        
        data = request.get_json(force=True)
        
        if "status" in data:
            scan.status = data["status"]
        if "current_industry" in data:
            scan.current_industry = data["current_industry"]
        if "companies" in data:
            scan.companies_json = json.dumps(data["companies"])
            scan.total_companies = len(data["companies"])
        if "error_message" in data:
            scan.error_message = data["error_message"]
        if "name" in data:
            scan.name = data["name"]
        
        db.session.commit()
        
        return jsonify({"status": "updated", "id": scan.id})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-scans/<int:scan_id>", methods=["DELETE"])
def api_delete_market_scan(scan_id):
    """Delete a market scan."""
    try:
        from models import MarketScan
        scan = MarketScan.query.get(scan_id)
        if not scan:
            return jsonify({"error": "Scan not found"}), 404
        
        db.session.delete(scan)
        db.session.commit()
        
        return jsonify({"status": "deleted"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 500


def run_background_scan(app_context, scan_id):
    """Background worker to run market discovery scan."""
    import threading
    from openai import OpenAI
    
    with app_context:
        from models import MarketScan
        
        try:
            scan = MarketScan.query.get(scan_id)
            if not scan:
                return
            
            scan.status = "running"
            db.session.commit()
            
            industries = json.loads(scan.industries_json) if scan.industries_json else []
            options = json.loads(scan.options_json) if scan.options_json else {}
            count_per_industry = options.get("count", 5)
            
            api_key = os.environ.get("OPENAI_API_KEY")
            if not api_key:
                scan.status = "error"
                scan.error_message = "OpenAI API key not configured"
                db.session.commit()
                return
            
            client = OpenAI(api_key=api_key)
            all_companies = []
            
            for i, industry in enumerate(industries):
                scan.current_industry = i + 1
                db.session.commit()
                
                try:
                    prompt = f"""You are a market research assistant for Natoli Engineering Company.
Find {count_per_industry} real companies in this industry: {industry}

For each company provide:
- company: Full company name
- city: City location
- state: State/Province (2-letter code for US)
- phone: Phone number with country code
- website: Company website URL
- linkedin: LinkedIn company URL if known
- fit: "Engineering" (tablet tooling/manufacturing), "Scientific" (formulation/R&D), "Both", or "No Fit"
- reason: Brief reason for Natoli fit
- confidence: Score 1-100 on data accuracy

Return JSON array only, no markdown."""

                    response = client.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.7,
                        max_tokens=2000
                    )
                    
                    content = response.choices[0].message.content.strip()
                    if content.startswith("```"):
                        content = content.split("```")[1]
                        if content.startswith("json"):
                            content = content[4:]
                    
                    companies = json.loads(content)
                    for c in companies:
                        c["industry"] = industry
                        all_companies.append(c)
                    
                    scan.companies_json = json.dumps(all_companies)
                    scan.total_companies = len(all_companies)
                    db.session.commit()
                    
                except Exception as e:
                    print(f"Error scanning {industry}: {e}")
                    continue
            
            scan.status = "completed"
            scan.current_industry = len(industries)
            db.session.commit()
            
        except Exception as e:
            scan = MarketScan.query.get(scan_id)
            if scan:
                scan.status = "error"
                scan.error_message = str(e)
                db.session.commit()


@app.route("/api/market-scans/<int:scan_id>/start", methods=["POST"])
def api_start_background_scan(scan_id):
    """Start a background scan job."""
    import threading
    
    try:
        from models import MarketScan
        scan = MarketScan.query.get(scan_id)
        if not scan:
            return jsonify({"error": "Scan not found"}), 404
        
        if scan.status == "running":
            return jsonify({"error": "Scan already running"}), 400
        
        thread = threading.Thread(
            target=run_background_scan,
            args=(app.app_context(), scan_id)
        )
        thread.daemon = True
        thread.start()
        
        return jsonify({"status": "started", "id": scan_id})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-scans/running", methods=["GET"])
def api_get_running_scans():
    """Get all currently running scans."""
    try:
        from models import MarketScan
        scans = MarketScan.query.filter(MarketScan.status.in_(["pending", "running"])).all()
        
        return jsonify({
            "scans": [{
                "id": s.id,
                "name": s.name,
                "status": s.status,
                "current_industry": s.current_industry,
                "total_industries": s.total_industries,
                "total_companies": s.total_companies,
                "created_at": s.created_at.isoformat() if s.created_at else None
            } for s in scans]
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-scans/<int:scan_id>/stop", methods=["POST"])
def api_stop_scan(scan_id):
    """Stop a running market scan."""
    try:
        from models import MarketScan
        scan = MarketScan.query.get(scan_id)
        if not scan:
            return jsonify({"error": "Scan not found"}), 404
        
        if scan.status not in ["pending", "running"]:
            return jsonify({"error": "Scan is not running"}), 400
        
        scan.status = "stopped"
        db.session.commit()
        
        return jsonify({"success": True, "message": f"Scan '{scan.name}' stopped"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-discovery/ai-scan", methods=["POST"])
def api_ai_pressure_scan():
    """TonyOS LLM Brain - AI-powered pressure and decision analysis for companies."""
    data = request.get_json(force=True)
    company_text = data.get("text", "")
    company_name = data.get("name", "Unknown")
    industry = data.get("industry", "")
    
    if not company_text:
        return jsonify({"error": "Company text required"}), 400
    
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return jsonify({"error": "OpenAI API key not configured"}), 500
    
    system_prompt = """You are TonyOS Manufacturing Intelligence Engine for Natoli Engineering Company.

Think like a Natoli process expert evaluating tablet compression and powder compaction opportunities.

Evaluate the company for:
1. COMPRESSION RELEVANCE - Does this company work with tablets, pellets, compacts, or powder compression?
2. SCALING PRESSURE - Are they growing, scaling up, or facing throughput challenges?
3. TOOLING RISK - Do they have tooling wear, maintenance, or quality issues?
4. REGULATORY PRESSURE - FDA, GMP, validation, compliance concerns?
5. COMPETITOR GRAVITY - Are competitors like Fette, Korsch, Syntegon, IMA, Kikusui involved?
6. SILENCE SIGNALS - Is this company flying under the radar with hidden potential?

Division Routing:
- SCIENTIFIC: Formulation, R&D, scale-up, process development, QbD, bioavailability, dissolution, stability
- ENGINEERING: Tooling, punch/die, press maintenance, B/D tooling, turret, manufacturing equipment

Return ONLY valid JSON:
{
  "compression": "high|medium|low|none",
  "pressure": ["list of detected pressure signals"],
  "risk": ["list of risk factors"],
  "competitors": ["list of detected competitors"],
  "decision": "ENGAGE|EDUCATE|MONITOR|DEPRIORITIZE",
  "division": "Scientific|Engineering|Both|None",
  "confidence": 85,
  "summary": "One sentence summary of the opportunity"
}"""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Company: {company_name}\nIndustry: {industry}\n\nDescription:\n{company_text}"}
            ],
            temperature=0,
            max_tokens=500
        )
        
        result_text = response.choices[0].message.content.strip()
        if result_text.startswith("```"):
            result_text = result_text.split("```")[1]
            if result_text.startswith("json"):
                result_text = result_text[4:]
        
        result = json.loads(result_text)
        result["company"] = company_name
        result["industry"] = industry
        
        return jsonify(result)
    except json.JSONDecodeError as e:
        return jsonify({
            "error": "Failed to parse AI response",
            "raw": result_text if 'result_text' in dir() else str(e)
        }), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-discovery/ai-batch-scan", methods=["POST"])
def api_ai_batch_scan():
    """TonyOS LLM Brain - Batch AI analysis for multiple companies."""
    data = request.get_json(force=True)
    companies = data.get("companies", [])
    
    if not companies:
        return jsonify({"error": "Companies list required"}), 400
    
    if len(companies) > 20:
        return jsonify({"error": "Maximum 20 companies per batch"}), 400
    
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return jsonify({"error": "OpenAI API key not configured"}), 500
    
    results = []
    for company in companies:
        try:
            company_name = company.get('company', company.get('name', 'Unknown'))
            industry = company.get('industry', '')
            reason = company.get('reason', '')
            text = f"{company_name} - {industry} - {reason}"
            
            system_prompt = """You are TonyOS Manufacturing Intelligence Engine for Natoli. Evaluate tablet compression opportunities.
Return ONLY valid JSON with ALL fields:
{"compression":"high|medium|low|none","pressure":["list of pressure signals"],"risk":["list of risks"],"competitors":["Fette","Korsch","etc if detected"],"decision":"ENGAGE|EDUCATE|MONITOR|DEPRIORITIZE","division":"Scientific|Engineering|Both|None","confidence":85,"summary":"One sentence opportunity summary"}"""

            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": text}
                ],
                temperature=0,
                max_tokens=300
            )
            
            result_text = response.choices[0].message.content.strip()
            if result_text.startswith("```"):
                result_text = result_text.split("```")[1]
                if result_text.startswith("json"):
                    result_text = result_text[4:]
            
            result = json.loads(result_text)
            result["company"] = company_name
            result["original"] = company
            result.setdefault("pressure", [])
            result.setdefault("competitors", [])
            result.setdefault("risk", [])
            result.setdefault("summary", "")
            result.setdefault("confidence", 50)
            results.append(result)
        except Exception as e:
            results.append({
                "company": company.get('company', company.get('name', 'Unknown')),
                "error": str(e),
                "decision": "REVIEW",
                "division": "Unknown",
                "compression": "none",
                "pressure": [],
                "competitors": [],
                "confidence": 0,
                "summary": "Analysis failed"
            })
    
    return jsonify({"results": results, "count": len(results)})


@app.route("/api/scan-memory", methods=["GET"])
def api_get_scan_memory():
    """Get scan memory history with optional filters."""
    from models import ScanMemory
    territory = request.args.get("territory")
    decision = request.args.get("decision")
    limit = request.args.get("limit", 100, type=int)
    
    query = ScanMemory.query.order_by(ScanMemory.created_at.desc())
    if territory:
        query = query.filter(ScanMemory.territory == territory)
    if decision:
        query = query.filter(ScanMemory.final_decision == decision)
    
    memories = query.limit(limit).all()
    return jsonify([{
        "id": m.id,
        "company_name": m.company_name,
        "industry": m.industry,
        "territory": m.territory,
        "sales_rep": m.sales_rep,
        "ai_decision": m.ai_decision,
        "ai_compression": m.ai_compression,
        "ai_division": m.ai_division,
        "ai_confidence": m.ai_confidence,
        "ai_summary": m.ai_summary,
        "final_decision": m.final_decision,
        "notes": m.notes,
        "created_at": m.created_at.isoformat() if m.created_at else None
    } for m in memories])


@app.route("/api/scan-memory", methods=["POST"])
def api_save_scan_memory():
    """Save companies to scan memory."""
    from models import ScanMemory
    data = request.get_json(force=True)
    companies = data.get("companies", [])
    scan_id = data.get("scan_id")
    
    saved_records = []
    for c in companies:
        memory = ScanMemory(
            company_name=c.get("company", ""),
            industry=c.get("industry", ""),
            territory=c.get("territory", ""),
            sales_rep=c.get("salesRep", ""),
            ai_decision=c.get("aiDecision", ""),
            ai_compression=c.get("aiCompression", ""),
            ai_division=c.get("aiDivision", ""),
            ai_confidence=c.get("aiConfidence"),
            ai_summary=c.get("aiSummary", ""),
            pressure_signals=json.dumps(c.get("aiPressure", [])),
            competitors=json.dumps(c.get("aiCompetitors", [])),
            final_decision=c.get("aiDecision", ""),
            scan_id=scan_id
        )
        db.session.add(memory)
        db.session.flush()
        saved_records.append({"company": c.get("company", ""), "memory_id": memory.id})
    
    db.session.commit()
    return jsonify({"success": True, "saved": len(saved_records), "records": saved_records})


@app.route("/api/scan-memory/<int:memory_id>", methods=["PATCH"])
def api_update_scan_memory(memory_id):
    """Update a scan memory entry (for overrides and notes)."""
    from models import ScanMemory, OverrideLog
    data = request.get_json(force=True)
    
    memory = ScanMemory.query.get(memory_id)
    if not memory:
        return jsonify({"error": "Not found"}), 404
    
    new_decision = data.get("final_decision")
    if new_decision and new_decision != memory.final_decision:
        override = OverrideLog(
            company_name=memory.company_name,
            original_decision=memory.final_decision or memory.ai_decision or "",
            new_decision=new_decision,
            reason=data.get("override_reason", ""),
            changed_by=data.get("changed_by", "user"),
            scan_memory_id=memory_id
        )
        db.session.add(override)
        memory.final_decision = new_decision
    
    if "notes" in data:
        memory.notes = data["notes"]
    
    db.session.commit()
    return jsonify({"success": True})


@app.route("/api/override-logs", methods=["GET"])
def api_get_override_logs():
    """Get override logs."""
    from models import OverrideLog
    limit = request.args.get("limit", 50, type=int)
    logs = OverrideLog.query.order_by(OverrideLog.created_at.desc()).limit(limit).all()
    return jsonify([{
        "id": l.id,
        "company_name": l.company_name,
        "original_decision": l.original_decision,
        "new_decision": l.new_decision,
        "reason": l.reason,
        "changed_by": l.changed_by,
        "created_at": l.created_at.isoformat() if l.created_at else None
    } for l in logs])


@app.route("/api/decision-dashboard", methods=["GET"])
def api_decision_dashboard():
    """Get decision dashboard stats by territory and timeframe."""
    from models import ScanMemory
    from sqlalchemy import func
    from datetime import datetime, timedelta
    
    days = request.args.get("days", 30, type=int)
    cutoff = datetime.now() - timedelta(days=days)
    
    by_decision = db.session.query(
        ScanMemory.final_decision,
        func.count(ScanMemory.id)
    ).filter(ScanMemory.created_at >= cutoff).group_by(ScanMemory.final_decision).all()
    
    by_territory = db.session.query(
        ScanMemory.territory,
        ScanMemory.final_decision,
        func.count(ScanMemory.id)
    ).filter(ScanMemory.created_at >= cutoff).group_by(
        ScanMemory.territory, ScanMemory.final_decision
    ).all()
    
    by_division = db.session.query(
        ScanMemory.ai_division,
        func.count(ScanMemory.id)
    ).filter(ScanMemory.created_at >= cutoff).group_by(ScanMemory.ai_division).all()
    
    total = ScanMemory.query.filter(ScanMemory.created_at >= cutoff).count()
    
    territories = {}
    for territory, decision, count in by_territory:
        if territory not in territories:
            territories[territory] = {"ENGAGE": 0, "EDUCATE": 0, "MONITOR": 0, "DEPRIORITIZE": 0}
        if decision:
            territories[territory][decision] = count
    
    return jsonify({
        "total": total,
        "days": days,
        "by_decision": {d or "UNKNOWN": c for d, c in by_decision},
        "by_territory": territories,
        "by_division": {d or "UNKNOWN": c for d, c in by_division}
    })


@app.route("/api/pressure-signals", methods=["GET"])
def api_get_pressure_signals():
    """Get pressure signal library."""
    from models import PressureSignal
    signals = PressureSignal.query.filter_by(is_active=True).order_by(PressureSignal.category, PressureSignal.keyword).all()
    return jsonify([{
        "id": s.id,
        "keyword": s.keyword,
        "category": s.category,
        "weight": s.weight,
        "description": s.description
    } for s in signals])


@app.route("/api/pressure-signals", methods=["POST"])
def api_add_pressure_signal():
    """Add a new pressure signal."""
    from models import PressureSignal
    data = request.get_json(force=True)
    
    signal = PressureSignal(
        keyword=data.get("keyword", ""),
        category=data.get("category", "general"),
        weight=data.get("weight", 1),
        description=data.get("description", "")
    )
    db.session.add(signal)
    db.session.commit()
    return jsonify({"success": True, "id": signal.id})


@app.route("/api/pressure-signals/<int:signal_id>", methods=["DELETE"])
def api_delete_pressure_signal(signal_id):
    """Deactivate a pressure signal."""
    from models import PressureSignal
    signal = PressureSignal.query.get(signal_id)
    if not signal:
        return jsonify({"error": "Not found"}), 404
    signal.is_active = False
    db.session.commit()
    return jsonify({"success": True})


@app.route("/api/pressure-signals/seed", methods=["POST"])
def api_seed_pressure_signals():
    """Seed default pressure signals."""
    from models import PressureSignal
    
    defaults = [
        ("scale-up", "pressure", 3, "Company scaling production"),
        ("throughput", "pressure", 3, "Throughput constraints"),
        ("validation", "regulatory", 2, "Validation requirements"),
        ("FDA", "regulatory", 3, "FDA involvement"),
        ("GMP", "regulatory", 2, "GMP compliance"),
        ("tooling", "equipment", 2, "Tooling needs"),
        ("tablet", "compression", 3, "Tablet manufacturing"),
        ("compaction", "compression", 3, "Powder compaction"),
        ("formulation", "scientific", 2, "Formulation development"),
        ("R&D", "scientific", 2, "Research and development"),
        ("Fette", "competitor", 2, "Competitor detected"),
        ("Korsch", "competitor", 2, "Competitor detected"),
        ("Syntegon", "competitor", 2, "Competitor detected"),
        ("IMA", "competitor", 2, "Competitor detected"),
        ("Kikusui", "competitor", 2, "Competitor detected"),
        ("yield", "pressure", 2, "Yield improvement needs"),
        ("variability", "pressure", 2, "Process variability"),
        ("capacity", "pressure", 2, "Capacity constraints")
    ]
    
    added = 0
    for keyword, category, weight, desc in defaults:
        existing = PressureSignal.query.filter_by(keyword=keyword).first()
        if not existing:
            signal = PressureSignal(keyword=keyword, category=category, weight=weight, description=desc)
            db.session.add(signal)
            added += 1
    
    db.session.commit()
    return jsonify({"success": True, "added": added})


@app.route("/api/market-discovery/generate-message", methods=["POST"])
def api_generate_natoli_message():
    """Generate personalized Natoli outreach message using OpenAI."""
    data = request.get_json(force=True)
    company = data.get("company", "")
    industry = data.get("industry", "")
    fit = data.get("fit", "Engineering")
    state = data.get("state", "")
    reason = data.get("reason", "")
    relationship = data.get("relationship", "new")
    
    if not company:
        return jsonify({"error": "Company required"}), 400
    
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return jsonify({"error": "OpenAI API key not configured"}), 500
    
    relationship_context = {
        "new": "New Customer - First outreach, no prior relationship",
        "existing": "Existing Customer - Already buys from Natoli, looking to expand",
        "lapsed": "Lapsed Customer - Haven't ordered in 12+ months, re-engagement"
    }.get(relationship, "New Customer")
    
    prompt = f"""You are writing outreach on behalf of Natoli Engineering Company.

Natoli tone rules:
- Calm, credible, manufacturing-first
- No marketing hype or buzzwords
- No selling language or pressure
- Soft CTA only (e.g., "happy to discuss" or "let me know if useful")
- Written for experienced manufacturing professionals who know their craft

Context:
Company: {company}
Location: {state}
Industry: {industry}
Natoli Fit: {fit}
Why they may need Natoli: {reason}
Relationship Type: {relationship_context}

Write:
1) A subject line (boring, professional - not clickbait)
2) A short email body (90-130 words max)

The email must clearly answer:
- Who Natoli is (the industry's trusted source for tablet press tooling, parts, and compaction expertise)
- Why reliability, consistency, and compaction science matter
- Why this outreach is reasonable and relevant to them

Do NOT pitch specific products. Focus on credibility and relevance.

Format your response EXACTLY like this:
SUBJECT: [your subject line here]

BODY:
[your email body here]"""

    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an experienced manufacturing operator writing professional outreach for Natoli Engineering. Be direct, credible, and manufacturing-focused."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=500
        )
        
        content = response.choices[0].message.content.strip()
        
        # Parse the response
        subject = ""
        body = ""
        
        if "SUBJECT:" in content and "BODY:" in content:
            parts = content.split("BODY:")
            subject_part = parts[0].replace("SUBJECT:", "").strip()
            subject = subject_part.strip()
            body = parts[1].strip() if len(parts) > 1 else ""
        else:
            # Fallback parsing
            lines = content.split("\n")
            subject = lines[0].replace("Subject:", "").replace("SUBJECT:", "").strip()
            body = "\n".join(lines[1:]).strip()
        
        return jsonify({
            "subject": subject,
            "body": body,
            "company": company
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-discovery/enrich", methods=["POST"])
def api_company_enrichment():
    """Enrich company profiles using OpenAI"""
    data = request.get_json()
    prompt = data.get("prompt", "")
    
    if not prompt:
        return jsonify({"error": "Prompt required"}), 400
    
    try:
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a B2B company research assistant. Return only valid JSON arrays."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=2000
        )
        
        content = response.choices[0].message.content.strip()
        
        # Parse JSON from response
        import json
        json_match = re.search(r'\[[\s\S]*\]', content)
        if json_match:
            results = json.loads(json_match.group(0))
            return jsonify({"results": results})
        else:
            return jsonify({"error": "Failed to parse response", "raw": content}), 500
            
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/market-discovery/deep", methods=["POST"])
def api_market_discovery_deep():
    """Deep Discovery - Web search + robots.txt gate + page scraping + AI classification."""
    import urllib.robotparser as robotparser
    from urllib.parse import urlparse
    from bs4 import BeautifulSoup
    import requests as req
    
    data = request.get_json(force=True)
    keyword = data.get("keyword", "")
    count = min(data.get("count", 10), 20)
    
    if not keyword:
        return jsonify({"error": "Keyword required"}), 400
    
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return jsonify({"error": "OpenAI API key not configured"}), 500
    
    def can_scrape(url):
        """Check robots.txt permission"""
        try:
            parsed = urlparse(url)
            robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
            rp = robotparser.RobotFileParser()
            rp.set_url(robots_url)
            rp.read()
            return rp.can_fetch("MarketIntelBot/1.0", url)
        except:
            return False
    
    def fetch_page_text(url):
        """Safely fetch page content respecting robots.txt"""
        if not can_scrape(url):
            return None, "blocked_by_robots"
        
        try:
            r = req.get(url, timeout=10, headers={
                "User-Agent": "MarketIntelBot/1.0 (+https://natoli.com/market-research)"
            })
            soup = BeautifulSoup(r.text, "html.parser")
            
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            text = soup.get_text(separator=" ", strip=True)
            return text[:6000], "success"
        except Exception as e:
            return None, f"fetch_error: {str(e)}"
    
    def classify_company(url, text):
        """Use OpenAI to classify the company"""
        prompt = f"""You are an industrial classification engine for Natoli Engineering, which sells tablet presses, tooling, and formulation services.

Analyze this company website text and determine:
1. Company name (extract from content)
2. Industry category
3. Whether they manufacture compressed tablets, pellets, or solid dosage forms
4. Classification:
   - Engineering fit (needs tablet presses, tooling, manufacturing equipment)
   - Scientific fit (formulation, R&D, analytical services)
   - Both (manufacturing + formulation)
   - No Fit (doesn't use tablet compression)
5. Location (city, state if USA)
6. Key products or services
7. Confidence score (0-100)

Website URL: {url}
Website Text:
{text}

Return ONLY valid JSON with this format:
{{"company": "Name", "industry": "Category", "makes_tablets": true/false, "fit": "Engineering|Scientific|Both|No Fit", "city": "City", "state": "State", "products": "Brief description", "reason": "Why they need Natoli", "confidence": 85}}"""

        try:
            response = openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=500
            )
            content = response.choices[0].message.content.strip()
            if content.startswith("```"):
                content = content.split("```")[1]
                if content.startswith("json"):
                    content = content[4:]
            return json.loads(content)
        except:
            return None
    
    results = []
    errors = []
    verified = []
    
    try:
        url_prompt = f"""Find {count * 2} real US-based companies that are related to: "{keyword}"

Return ONLY a JSON array with company name and website URL. Focus on actual manufacturers, not retailers or distributors.
Format: [{{"company": "Name", "website": "https://..."}}]

Only include companies you're confident exist with valid websites."""

        url_response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": url_prompt}],
            temperature=0.3,
            max_tokens=1000
        )
        
        url_content = url_response.choices[0].message.content.strip()
        if "```" in url_content:
            url_content = url_content.split("```")[1].replace("json", "", 1)
        
        company_urls = json.loads(url_content)
        
        processed = 0
        seen_domains = set()
        
        for item in company_urls:
            if processed >= count:
                break
            
            url = item.get("website", "")
            company_hint = item.get("company", "")
            
            if not url or not url.startswith("http"):
                continue
            
            domain = urlparse(url).netloc
            if domain in seen_domains:
                continue
            seen_domains.add(domain)
            
            if any(skip in domain for skip in ["linkedin.com", "facebook.com", "twitter.com", "youtube.com", "wikipedia.org", "yelp.com", "indeed.com", "amazon.com"]):
                continue
            
            text, status = fetch_page_text(url)
            
            if text:
                classification = classify_company(url, text)
                if classification:
                    classification["website"] = url
                    classification["domain"] = domain
                    classification["scrape_status"] = "verified"
                    classification["company_hint"] = company_hint
                    results.append(classification)
                    processed += 1
                    verified.append(url)
            else:
                errors.append({"url": url, "company": company_hint, "status": status})
        
        return jsonify({
            "companies": results,
            "total": len(results),
            "verified": len(verified),
            "errors": errors,
            "keyword": keyword
        })
        
    except Exception as e:
        import traceback
        return jsonify({"error": str(e), "trace": traceback.format_exc()}), 500


@app.route("/api/market-discovery/analyze-url", methods=["POST"])
def api_analyze_single_url():
    """Analyze a single URL for Natoli fit"""
    import urllib.robotparser as robotparser
    from urllib.parse import urlparse
    from bs4 import BeautifulSoup
    import requests as req
    
    data = request.get_json(force=True)
    url = data.get("url", "")
    
    if not url:
        return jsonify({"error": "URL required"}), 400
    
    def can_scrape(target_url):
        try:
            parsed = urlparse(target_url)
            robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
            rp = robotparser.RobotFileParser()
            rp.set_url(robots_url)
            rp.read()
            return rp.can_fetch("MarketIntelBot/1.0", target_url)
        except:
            return False
    
    if not can_scrape(url):
        return jsonify({"error": "Website blocks automated access (robots.txt)", "blocked": True}), 403
    
    try:
        r = req.get(url, timeout=10, headers={
            "User-Agent": "MarketIntelBot/1.0 (+https://natoli.com/market-research)"
        })
        soup = BeautifulSoup(r.text, "html.parser")
        
        for script in soup(["script", "style", "nav", "footer"]):
            script.decompose()
        
        text = soup.get_text(separator=" ", strip=True)[:6000]
        
        prompt = f"""Analyze this company website for Natoli Engineering (tablet press & tooling manufacturer).

Website: {url}
Content: {text}

Return JSON:
{{"company": "Name", "industry": "Category", "fit": "Engineering|Scientific|Both|No Fit", "city": "City", "state": "State", "phone": "Phone if found", "products": "What they make", "reason": "Why they need Natoli", "confidence": 85}}"""
        
        response = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=500
        )
        
        content = response.choices[0].message.content.strip()
        if "```" in content:
            content = content.split("```")[1].replace("json", "", 1)
        
        result = json.loads(content)
        result["website"] = url
        result["scrape_status"] = "allowed"
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/lapsed-accounts/save", methods=["POST"])
def api_lapsed_accounts_save():
    """Save cleaned companies list to file"""
    data = request.get_json()
    companies = data.get("companies", [])
    csv_content = data.get("csv", "")
    filename = data.get("filename", "cleaned_companies.csv")
    
    if not companies:
        return jsonify({"error": "No companies provided"}), 400
    
    try:
        # Save to data directory
        import os
        data_dir = os.path.join(os.path.dirname(__file__), "data", "lapsed_accounts")
        os.makedirs(data_dir, exist_ok=True)
        
        filepath = os.path.join(data_dir, filename)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(csv_content)
        
        return jsonify({
            "success": True,
            "message": f"Saved {len(companies)} companies",
            "filepath": filepath
        })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/lapsed-accounts/list", methods=["GET"])
def api_lapsed_accounts_list():
    """List saved cleaned company files"""
    try:
        import os
        data_dir = os.path.join(os.path.dirname(__file__), "data", "lapsed_accounts")
        if not os.path.exists(data_dir):
            return jsonify({"files": []})
        
        files = []
        for f in os.listdir(data_dir):
            if f.endswith('.csv'):
                filepath = os.path.join(data_dir, f)
                files.append({
                    "name": f,
                    "size": os.path.getsize(filepath),
                    "modified": os.path.getmtime(filepath)
                })
        
        return jsonify({"files": sorted(files, key=lambda x: x['modified'], reverse=True)})
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# Natoli Reachout - Copy-paste LinkedIn/Email outreach generator
REACHOUT_PROMPT = """
You are acting as a senior HQ resource at Natoli Engineering.

TASK:
For each person provided, do the following:

1. Classify Tier:
- Tier 1: Director+ roles in MS&T, Engineering, Manufacturing, Ops with tablet press or scale exposure
- Tier 2: Manager-level technical roles influencing scale or tooling
- Tier 3: R&D roles without clear scale exposure
- Procurement or unclear roles: Not Ready to Contact

2. Identify Persona:
- Decision Maker
- Influencer
- Blocker

3. Determine Natoli Fit:
- Engineering
- Scientific
- Both
- Not Ready

4. Choose ONE sales angle:
- Speed / lead time
- Reliability / uptime
- Scale-up readiness
- Compliance / validation
- Not ready to contact

5. Generate outputs:
- LinkedIn connection message (1–2 sentences, no CTA, HQ resource tone)
- Email (short, clean spacing, no sales CTA, calm professional)
- C2 Notes (Tier, Persona, Why Fit, Angle, Routing)

Rules:
- Slightly vary wording per person
- Never sound salesy
- Never include emojis
- Never mention pricing
- Output MUST be valid JSON with these exact keys: tier, persona, natoli_fit, sales_angle, linkedin_message, email, c2_notes
- c2_notes should be an object with keys: tier, persona, why_fit, angle, routing

INPUT PERSON:
"""

@app.route("/reachout")
def reachout_page():
    return send_from_directory("static/internal/natoli-reachout/public", "index.html")

@app.route("/reachout/<path:filename>")
def reachout_static(filename):
    return send_from_directory("static/internal/natoli-reachout/public", filename)

@app.route("/api/reachout/generate", methods=["POST"])
def reachout_generate():
    try:
        data = request.get_json()
        lines = data.get("lines", [])
        
        if not lines:
            return jsonify({"error": "No lines provided"}), 400
        
        results = []
        
        for line in lines:
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": REACHOUT_PROMPT},
                    {"role": "user", "content": line}
                ],
                temperature=0.4
            )
            
            text = response.choices[0].message.content
            results.append({
                "input": line,
                "output": text
            })
        
        return jsonify({"results": results})
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


REACHOUT_SCORECARD_PROMPT = """You are the Natoli HQ Resource Reachout System. Generate a complete scorecard. Return JSON only (no markdown):

{"name":"","tier":"","tier_why":"","natoli_fit":"","lead_profile":"","why_fit":"","how_help":"","natoli_value":"","linkedin_message":"","email_subject":"","email_body":"","c2_notes":""}

=== TIER CLASSIFICATION (CRITICAL) ===

TIER 1 - THE ARCHITECTS (HIGHEST PRIORITY):
- C-Suite executives (CEO, COO, CFO, CTO, CSO, CMO) - ALWAYS TIER 1
- VP/SVP of Manufacturing, Operations, R&D, MS&T, Quality
- Director of R&D / Formulation Scientists ("Architect of Certainty")
- Director of MS&T / Tech Transfer
- Director of Engineering / Manufacturing
- Anyone with "Chief" in title

TIER 2 - THE BUILDERS:
- Senior Manager / Manager of Manufacturing, Operations, Engineering
- Process Engineering Manager
- Formulation Manager / Senior Scientist
- Production Manager with tablet/press exposure
- Manager-level technical roles influencing scale or tooling

TIER 3 - THE SUPPORTERS:
- R&D roles without clear scale-up exposure
- Lab Scientists without manufacturing connection
- Junior technical roles

NOT READY:
- Procurement / Purchasing (blocked category)
- HR / Finance / IT without manufacturing tie
- Unclear or unrelated roles

=== HERO PERSONAS BY TIER ===

TIER 1 HERO: "The Architect of Certainty"
- Core Belief: Nothing moves to production until certainty is achieved
- Emotional Driver: They want to bring order to the unpredictable - they want to KNOW, not guess
- Pain Points: Scale-up failures, weak compaction data, unpredictable powder behavior, delays that damage launch timelines
- Desired Outcomes: Clean compaction profiles, predictable scale-up, data-driven formulation optimization, regulatory alignment
- What Natoli Represents: The Architect's Blueprint - data, diagnostics, compaction analytics, and scale-up truth

TIER 2 HERO: "The Problem Solver"
- Core Belief: Every manufacturing challenge has a root cause that can be fixed
- Emotional Driver: They want to eliminate fire drills and create repeatable processes
- Pain Points: Tooling wear, press inconsistency, production bottlenecks, quality deviations
- Desired Outcomes: Reliable tooling performance, reduced downtime, consistent output

=== NATOLI FIT ===
- Engineering: Production scale, tooling, uptime, press performance, reliability
- Scientific: R&D, validation, scale-up, compaction analytics, formulation optimization
- Both: Combination - usually Tier 1 leaders overseeing both functions

=== LINKEDIN CONNECTION (EXACTLY 2 OPTIONS) ===
Generate exactly 2 options, separated by "---". BOTH must include "would love to connect":

Option 1 (Warm + Professional):
"Hi [Name], I came across your profile and would love to connect. I work with Natoli at our headquarters supporting teams on [Engineering: tooling and press performance / Scientific: formulation and scale-up challenges / Both: manufacturing and R&D challenges]."

Option 2 (Short + Human):
"Hi [Name], would love to connect. I'm with Natoli HQ and work with teams focused on [Engineering: tablet manufacturing / Scientific: scale-up and formulation / Both: solid dose manufacturing] - always good to connect with others in the space."

=== EMAIL SUBJECT LINES BY TIER ===

TIER 1 SUBJECTS (Choose one based on fit):
- "Scale-up visibility for [Company]" (Scientific)
- "Compaction data question" (Scientific)
- "Manufacturing reliability at [Company]" (Engineering)
- "Quick intro - Natoli HQ" (General)

TIER 2 SUBJECTS:
- "Tooling performance question" (Engineering)
- "Process consistency at [Company]" (Engineering)
- "Formulation support question" (Scientific)

TIER 3 SUBJECTS:
- "Quick intro - Natoli" (General)

=== EMAIL TEMPLATES BY FIT ===

FOR SCIENTIFIC FIT (Tier 1 - Architect of Certainty):
"Hi [First Name],

I came across your profile and wanted to reach out.

I work with Natoli at our headquarters supporting R&D and formulation teams working through scale-up uncertainty - unpredictable powder behavior, compaction profile gaps, and timelines that get pushed when confidence is not there.

If any of that resonates, I would be glad to share what we are seeing help similar teams move from guessing to knowing.

And if you are not the right person, would you mind pointing me toward whoever leads formulation or tech transfer?

Either way, appreciate your time.

Tony Beal"

FOR ENGINEERING FIT:
"Hi [First Name],

I came across your profile and wanted to reach out.

I work with Natoli at our headquarters and often connect with manufacturing leaders dealing with tooling wear, press downtime, or production consistency challenges.

If any of that sounds familiar, I would be happy to share what we are seeing work for similar teams.

And if you are not the right person, would you mind pointing me to whoever handles tooling or press operations?

Either way, appreciate your time.

Tony Beal"

FOR BOTH:
"Hi [First Name],

I came across your profile and wanted to reach out.

I work with Natoli at our headquarters supporting teams on both the tooling side and the formulation science side - from press reliability to scale-up data and compaction analytics.

If either area is something you are working through, happy to share what is helping other teams gain certainty before production.

And if you are not the right person, would you mind pointing me in the right direction?

Either way, appreciate your time.

Tony Beal"

=== RULES ===
- No em dashes (use regular dashes only)
- CRITICAL: Put each sentence on its own line with a blank line between paragraphs
- C-Suite is ALWAYS Tier 1 - no exceptions
- LinkedIn: EXACTLY 2 options, both must say "would love to connect"
- Warm, human, conversational tone
- Soft ask to pass along if not right person
- No hard CTA, just offering help
- End with just "Tony Beal" on its own line
- Use Tier 1 "Architect of Certainty" language for R&D/Formulation Directors
- Speak in terms of: data, predictability, reproducibility, certainty, truth
- Position Natoli as the partner that eliminates scale-up uncertainty"""

REACHOUT_FAST_PROMPT = """Natoli Engineering sales classifier. For this contact, return JSON only:

TIER: Tier 1 (Director+ MS&T/Eng/Mfg), Tier 2 (Manager technical), Tier 3 (R&D no scale), Not Ready (procurement/unclear)
PERSONA: Decision Maker, Influencer, or Blocker
FIT: Engineering, Scientific, Both, or Not Ready
ANGLE: Speed, Reliability, Scale-up, Compliance, or Not ready

OUTPUT FORMAT (JSON only, no markdown):
{"tier":"","persona":"","natoli_fit":"","sales_angle":"","linkedin_message":"","email":"","c2_notes":""}

linkedin_message: 1-2 sentences, HQ resource tone, no CTA
email: Subject + 2-3 sentence body, professional, no sales CTA
c2_notes: Brief notes on tier/persona/routing"""

@app.route("/api/reachout/generate-fast", methods=["POST"])
def reachout_generate_fast():
    try:
        data = request.get_json()
        
        name = data.get("name", "")
        degree = data.get("degree", "")
        title = data.get("title", "")
        country = data.get("country", "")
        state = data.get("state", "")
        timeInRole = data.get("timeInRole", "")
        company = data.get("company", "")
        
        contact_info = f"{name}"
        if degree:
            contact_info += f" ({degree} connection)"
        if title:
            contact_info += f", {title}"
        if company:
            contact_info += f" at {company}"
        if country or state:
            location = ", ".join(filter(None, [state, country]))
            contact_info += f" - {location}"
        if timeInRole:
            contact_info += f" ({timeInRole} in role)"
        
        if not contact_info.strip():
            return jsonify({"error": "No contact info provided"}), 400
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": REACHOUT_FAST_PROMPT},
                {"role": "user", "content": contact_info}
            ],
            temperature=0.3,
            max_tokens=500
        )
        
        text = response.choices[0].message.content
        
        try:
            json_match = text
            if "```" in text:
                import re
                code_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
                if code_match:
                    json_match = code_match.group(1)
            
            brace_match = json_match[json_match.find("{"):json_match.rfind("}")+1] if "{" in json_match else json_match
            parsed = json.loads(brace_match)
            
            return jsonify({
                "tier": parsed.get("tier", "Unknown"),
                "persona": parsed.get("persona", "Unknown"),
                "natoli_fit": parsed.get("natoli_fit", "Unknown"),
                "sales_angle": parsed.get("sales_angle", ""),
                "linkedin_message": parsed.get("linkedin_message", ""),
                "email": parsed.get("email", ""),
                "c2_notes": parsed.get("c2_notes", "")
            })
        except:
            return jsonify({
                "tier": "Unknown",
                "persona": "Unknown", 
                "natoli_fit": "Unknown",
                "linkedin_message": text,
                "email": "",
                "c2_notes": ""
            })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/reachout/scorecard", methods=["POST"])
def reachout_scorecard():
    try:
        data = request.get_json()
        contact = data.get("contact", "").strip()
        
        if not contact:
            return jsonify({"error": "No contact info provided"}), 400
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": REACHOUT_SCORECARD_PROMPT},
                {"role": "user", "content": contact}
            ],
            temperature=0.3,
            max_tokens=1000
        )
        
        text = response.choices[0].message.content
        
        try:
            json_match = text
            if "```" in text:
                import re
                code_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
                if code_match:
                    json_match = code_match.group(1)
            
            brace_match = json_match[json_match.find("{"):json_match.rfind("}")+1] if "{" in json_match else json_match
            parsed = json.loads(brace_match)
            
            return jsonify({
                "name": parsed.get("name", ""),
                "tier": parsed.get("tier", "Unknown"),
                "tier_why": parsed.get("tier_why", ""),
                "natoli_fit": parsed.get("natoli_fit", "Unknown"),
                "lead_profile": parsed.get("lead_profile", ""),
                "why_fit": parsed.get("why_fit", ""),
                "how_help": parsed.get("how_help", ""),
                "natoli_value": parsed.get("natoli_value", ""),
                "linkedin_message": parsed.get("linkedin_message", ""),
                "email_subject": parsed.get("email_subject", ""),
                "email_body": parsed.get("email_body", ""),
                "c2_notes": parsed.get("c2_notes", "")
            })
        except:
            return jsonify({
                "name": "Contact",
                "tier": "Unknown",
                "natoli_fit": "Unknown",
                "lead_profile": text,
                "linkedin_message": "",
                "email_subject": "",
                "email_body": "",
                "c2_notes": ""
            })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/fix-linkedin-post", methods=["POST"])
def fix_linkedin_post():
    """Use GPT-4 Vision to read text from a LinkedIn post image and return corrected version"""
    try:
        data = request.get_json()
        image_data = data.get("image", "")
        
        if not image_data:
            return jsonify({"error": "No image provided"}), 400
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": """You are a professional editor. Analyze this LinkedIn post image and extract all text.
Fix any spelling errors, spacing issues, or grammatical problems.

Return JSON only (no markdown):
{
  "header": "Small header text at top if any",
  "headline": "Main headline text (use \\n for line breaks)",
  "bullets": [
    "First point with <strong>key words</strong> bolded",
    "Second point with <strong>key words</strong> bolded",
    "Third point with <strong>key words</strong> bolded",
    "Fourth point with <strong>key words</strong> bolded"
  ],
  "tagline": "Tagline at bottom if any"
}

Rules:
- Fix all spacing issues (e.g. "thework" should be "the work")
- Fix all spelling errors
- Keep the same meaning and structure
- Bold 1-2 key words in each bullet using <strong> tags
- Make text professional and polished"""
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": image_data}
                        }
                    ]
                }
            ],
            max_tokens=1000
        )
        
        text = response.choices[0].message.content
        
        try:
            json_match = text
            if "```" in text:
                import re
                code_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
                if code_match:
                    json_match = code_match.group(1)
            
            brace_start = json_match.find("{")
            brace_end = json_match.rfind("}")
            if brace_start != -1 and brace_end != -1:
                json_match = json_match[brace_start:brace_end+1]
            
            parsed = json.loads(json_match)
            return jsonify(parsed)
            
        except json.JSONDecodeError:
            return jsonify({
                "header": "",
                "headline": "Your Headline Here",
                "bullets": [],
                "tagline": ""
            })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/generate-article", methods=["POST"])
def generate_article():
    """Transform text into an award-winning article format"""
    try:
        data = request.get_json()
        text = data.get("text", "").strip()
        author = data.get("author", "Tony Beal")
        category = data.get("category", "Leadership & Growth")
        tone = data.get("tone", "professional")
        
        if not text:
            return jsonify({"error": "No text provided"}), 400
        
        tone_instructions = {
            "professional": "Write in a polished, professional tone suitable for Harvard Business Review.",
            "conversational": "Write in a warm, conversational tone like talking to a trusted colleague.",
            "authoritative": "Write with authority and expertise, citing principles and best practices.",
            "inspirational": "Write with an inspiring, motivating tone that empowers the reader."
        }
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": f"""You are an award-winning editor at a premier business publication.
Transform the provided text into a polished, professional article.

{tone_instructions.get(tone, tone_instructions['professional'])}

Return JSON only (no markdown):
{{
  "title": "Compelling article title",
  "lead": "One powerful opening paragraph that hooks the reader (2-3 sentences)",
  "body": "<p>Full article body with proper HTML formatting.</p><h2>Use subheadings</h2><p>Include paragraphs, blockquotes where appropriate.</p><ul><li>Bullet points when helpful</li></ul>",
  "disclaimer": "Professional disclaimer about opinions and educational purposes"
}}

Rules:
- Fix all spelling and grammar errors
- Structure content with clear sections
- Add compelling subheadings (use <h2> tags)
- Include quotes/highlights in <blockquote> tags
- Make paragraphs scannable
- Keep the author's voice but elevate the writing
- End with a thought-provoking conclusion
- Include an appropriate disclaimer"""
                },
                {
                    "role": "user",
                    "content": text
                }
            ],
            max_tokens=3000
        )
        
        text_resp = response.choices[0].message.content
        
        try:
            json_match = text_resp
            if "```" in text_resp:
                import re
                code_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', text_resp)
                if code_match:
                    json_match = code_match.group(1)
            
            brace_start = json_match.find("{")
            brace_end = json_match.rfind("}")
            if brace_start != -1 and brace_end != -1:
                json_match = json_match[brace_start:brace_end+1]
            
            parsed = json.loads(json_match)
            return jsonify(parsed)
            
        except json.JSONDecodeError:
            return jsonify({
                "title": "Untitled Article",
                "lead": "",
                "body": f"<p>{text}</p>",
                "disclaimer": "The views expressed are those of the author and do not constitute professional advice."
            })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/fix-linkedin-cover", methods=["POST"])
def fix_linkedin_cover():
    """Use GPT-4 Vision to read text from an image and return corrected version"""
    try:
        data = request.get_json()
        image_data = data.get("image", "")
        
        if not image_data:
            return jsonify({"error": "No image provided"}), 400
        
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": """You are a professional editor. Analyze this LinkedIn article cover image and extract all text.
Fix any spelling errors, spacing issues, or grammatical problems.

Return JSON only (no markdown):
{
  "header": "The small caps header text at top",
  "line1": "First main headline line",
  "line2": "Second main headline line", 
  "line3": "Third main headline line",
  "bullets": [
    "First bullet point with <strong>key words</strong> bolded",
    "Second bullet point with <strong>key words</strong> bolded",
    "Third bullet point with <strong>key words</strong> bolded",
    "Fourth bullet point with <strong>key words</strong> bolded"
  ],
  "tagline": "The tagline at bottom"
}

Rules:
- Fix all spacing issues (e.g. "thework" should be "the work")
- Fix all spelling errors
- Keep the same meaning and structure
- Bold 1-2 key words in each bullet using <strong> tags
- Make text professional and polished"""
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": image_data}
                        }
                    ]
                }
            ],
            max_tokens=1000
        )
        
        text = response.choices[0].message.content
        
        try:
            json_match = text
            if "```" in text:
                import re
                code_match = re.search(r'```(?:json)?\s*([\s\S]*?)```', text)
                if code_match:
                    json_match = code_match.group(1)
            
            brace_start = json_match.find("{")
            brace_end = json_match.rfind("}")
            if brace_start != -1 and brace_end != -1:
                json_match = json_match[brace_start:brace_end+1]
            
            parsed = json.loads(json_match)
            return jsonify(parsed)
            
        except json.JSONDecodeError:
            return jsonify({
                "header": "DISCIPLINE - GROWTH - REAL WORK",
                "line1": "Showing Up Is Easy.",
                "line2": "Staying With It",
                "line3": "Is the Real Work.",
                "bullets": [
                    "Doing <strong>the work</strong> when no one is clapping",
                    "Going to <strong>therapy</strong> instead of pretending you're fine",
                    "Facing <strong>old patterns</strong> instead of blaming new situations",
                    "Choosing <strong>growth</strong> when avoiding would be easier"
                ],
                "tagline": "Quiet work compounds."
            })
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/validate-urls", methods=["POST"])
def validate_urls():
    """Validate a batch of URLs to check if they're accessible"""
    data = request.get_json()
    urls = data.get("urls", [])
    
    if not urls:
        return jsonify({"error": "No URLs provided"}), 400
    
    results = []
    for url in urls[:100]:  # Limit to 100 URLs per request
        if not url or not isinstance(url, str):
            results.append({"url": url, "valid": False, "status": "invalid_format"})
            continue
            
        # Ensure URL has protocol
        check_url = url
        if not check_url.startswith(('http://', 'https://')):
            check_url = 'https://' + check_url
        
        try:
            response = requests.head(check_url, timeout=5, allow_redirects=True,
                                    headers={"User-Agent": "Mozilla/5.0 (compatible; Natoli/1.0)"})
            valid = response.status_code < 400
            results.append({
                "url": url,
                "valid": valid,
                "status": response.status_code,
                "final_url": response.url if response.url != check_url else None
            })
        except requests.exceptions.Timeout:
            results.append({"url": url, "valid": False, "status": "timeout"})
        except requests.exceptions.SSLError:
            # Try HTTP if HTTPS fails
            try:
                http_url = check_url.replace('https://', 'http://')
                response = requests.head(http_url, timeout=5, allow_redirects=True,
                                        headers={"User-Agent": "Mozilla/5.0 (compatible; Natoli/1.0)"})
                valid = response.status_code < 400
                results.append({"url": url, "valid": valid, "status": response.status_code})
            except:
                results.append({"url": url, "valid": False, "status": "ssl_error"})
        except requests.exceptions.ConnectionError:
            results.append({"url": url, "valid": False, "status": "connection_error"})
        except Exception as e:
            results.append({"url": url, "valid": False, "status": str(e)[:50]})
    
    valid_count = sum(1 for r in results if r.get("valid"))
    return jsonify({
        "results": results,
        "total": len(results),
        "valid": valid_count,
        "invalid": len(results) - valid_count
    })


if __name__ == "__main__":
    init_journal_table()
    app.run(host="0.0.0.0", port=5000, debug=True)
